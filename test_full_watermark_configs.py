#!/usr/bin/env python3
"""
完整水印配置测试脚本
测试所有水印位置、角度、颜色等配置选项
"""

from watermark_engine import WatermarkEngine, WatermarkConfig, create_watermark_config
from docx import Document
import os

def create_test_document(title: str) -> Document:
    """创建标准测试文档"""
    doc = Document()
    
    # 设置A4页面
    from docx.shared import Inches
    from docx.enum.section import WD_ORIENTATION
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.orientation = WD_ORIENTATION.PORTRAIT
    
    # 添加内容
    doc.add_heading(title, 0)
    doc.add_paragraph("这是测试文档的内容段落。")
    doc.add_paragraph("用于验证水印效果的可见性和位置准确性。")
    doc.add_paragraph("请检查水印是否按照配置正确显示。")
    
    return doc

def test_all_positions():
    """测试所有位置选项"""
    print("测试1: 所有位置选项")
    
    positions = [
        ("center", "居中大水印"),
        ("repeat", "平铺多处"),
        ("background", "背景样式"),
        ("top-left", "左上角"),
        ("top-right", "右上角"),
        ("bottom-left", "左下角"),
        ("bottom-right", "右下角")
    ]
    
    for position, desc in positions:
        doc = create_test_document(f"位置测试 - {desc}")
        
        config = create_watermark_config(
            enable=True,
            text=f"位置:{desc}",
            font_size=24,
            color="#FF5722",
            position=position,
            opacity=0.5,
            rotation=-45
        )
        
        success = WatermarkEngine.apply_watermark(doc, config)
        if success:
            filename = f"test_position_{position}.docx"
            doc.save(filename)
            print(f"  ✓ {desc} - {filename}")
        else:
            print(f"  ✗ {desc} - 失败")

def test_rotation_angles():
    """测试不同倾斜角度"""
    print("\n测试2: 倾斜角度选项")
    
    angles = [-90, -45, -30, 0, 30, 45, 90]
    
    for angle in angles:
        doc = create_test_document(f"角度测试 - {angle}°")
        
        config = create_watermark_config(
            enable=True,
            text=f"角度{angle}°",
            font_size=28,
            color="#2196F3",
            position="center",
            opacity=0.6,
            rotation=angle
        )
        
        success = WatermarkEngine.apply_watermark(doc, config)
        if success:
            filename = f"test_angle_{angle}deg.docx"
            doc.save(filename)
            print(f"  ✓ {angle}° - {filename}")
        else:
            print(f"  ✗ {angle}° - 失败")

def test_font_sizes():
    """测试不同字体大小"""
    print("\n测试3: 字体大小选项")
    
    sizes = [12, 18, 24, 32, 48, 64]
    
    for size in sizes:
        doc = create_test_document(f"字体大小测试 - {size}pt")
        
        config = create_watermark_config(
            enable=True,
            text=f"{size}pt字体",
            font_size=size,
            color="#4CAF50",
            position="center",
            opacity=0.4,
            rotation=-45
        )
        
        success = WatermarkEngine.apply_watermark(doc, config)
        if success:
            filename = f"test_fontsize_{size}pt.docx"
            doc.save(filename)
            print(f"  ✓ {size}pt - {filename}")
        else:
            print(f"  ✗ {size}pt - 失败")

def test_colors():
    """测试不同颜色"""
    print("\n测试4: 颜色选项")
    
    colors = [
        ("#FF0000", "红色"),
        ("#00FF00", "绿色"),
        ("#0000FF", "蓝色"),
        ("#FF5722", "橙色"),
        ("#9C27B0", "紫色"),
        ("#607D8B", "蓝灰色"),
        ("#795548", "棕色")
    ]
    
    for color_code, color_name in colors:
        doc = create_test_document(f"颜色测试 - {color_name}")
        
        config = create_watermark_config(
            enable=True,
            text=f"{color_name}水印",
            font_size=26,
            color=color_code,
            position="center",
            opacity=0.5,
            rotation=-45
        )
        
        success = WatermarkEngine.apply_watermark(doc, config)
        if success:
            filename = f"test_color_{color_name}.docx"
            doc.save(filename)
            print(f"  ✓ {color_name} ({color_code}) - {filename}")
        else:
            print(f"  ✗ {color_name} - 失败")

def test_opacity_levels():
    """测试不同透明度"""
    print("\n测试5: 透明度选项")
    
    opacities = [0.1, 0.3, 0.5, 0.7, 0.9]
    
    for opacity in opacities:
        doc = create_test_document(f"透明度测试 - {int(opacity*100)}%")
        
        config = create_watermark_config(
            enable=True,
            text=f"透明度{int(opacity*100)}%",
            font_size=24,
            color="#FF9800",
            position="center",
            opacity=opacity,
            rotation=-45
        )
        
        success = WatermarkEngine.apply_watermark(doc, config)
        if success:
            filename = f"test_opacity_{int(opacity*100)}pct.docx"
            doc.save(filename)
            print(f"  ✓ {int(opacity*100)}% - {filename}")
        else:
            print(f"  ✗ {int(opacity*100)}% - 失败")

def test_comprehensive_combinations():
    """测试综合配置组合"""
    print("\n测试6: 综合配置组合")
    
    combinations = [
        {
            "name": "商务正式",
            "config": {
                "text": "机密文档",
                "font_size": 32,
                "color": "#1976D2",
                "position": "center",
                "opacity": 0.3,
                "rotation": 0
            }
        },
        {
            "name": "法律专用",
            "config": {
                "text": "法律文件 • 投标苦",
                "font_size": 24,
                "color": "#D32F2F",
                "position": "repeat",
                "opacity": 0.4,
                "rotation": -30
            }
        },
        {
            "name": "内部使用",
            "config": {
                "text": "INTERNAL USE ONLY",
                "font_size": 20,
                "color": "#666666",
                "position": "background",
                "opacity": 0.2,
                "rotation": 45
            }
        },
        {
            "name": "草稿标识",
            "config": {
                "text": "DRAFT",
                "font_size": 48,
                "color": "#FF5722",
                "position": "top-right",
                "opacity": 0.6,
                "rotation": -90
            }
        },
        {
            "name": "版权声明",
            "config": {
                "text": "© 2024 投标苦系统",
                "font_size": 16,
                "color": "#795548",
                "position": "bottom-left",
                "opacity": 0.5,
                "rotation": 0
            }
        }
    ]
    
    for combo in combinations:
        doc = create_test_document(f"综合测试 - {combo['name']}")
        
        config = create_watermark_config(
            enable=True,
            **combo['config']
        )
        
        success = WatermarkEngine.apply_watermark(doc, config)
        if success:
            filename = f"test_combo_{combo['name']}.docx"
            doc.save(filename)
            print(f"  ✓ {combo['name']} - {filename}")
        else:
            print(f"  ✗ {combo['name']} - 失败")

def main():
    """运行所有测试"""
    print("开始完整水印配置测试...")
    print("=" * 60)
    
    # 创建测试结果目录
    test_dir = "watermark_full_test_results"
    os.makedirs(test_dir, exist_ok=True)
    os.chdir(test_dir)
    
    try:
        # 运行所有测试
        test_all_positions()
        test_rotation_angles()
        test_font_sizes()
        test_colors()
        test_opacity_levels()
        test_comprehensive_combinations()
        
        print("\n" + "=" * 60)
        print("所有测试完成！")
        
        # 统计生成的文件
        docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
        print(f"\n生成了 {len(docx_files)} 个测试文档：")
        
        for i, filename in enumerate(sorted(docx_files), 1):
            file_size = os.path.getsize(filename)
            print(f"{i:2d}. {filename} ({file_size} 字节)")
        
        print(f"\n所有文件已保存在 {os.path.abspath('.')} 目录中")
        print("\n测试要点：")
        print("1. 检查每种位置的水印是否出现在正确位置")
        print("2. 验证倾斜角度的装饰效果是否符合预期")
        print("3. 确认不同字体大小的可读性")
        print("4. 检查颜色显示是否准确")
        print("5. 验证透明度效果的视觉层次")
        print("6. 综合配置的整体效果评估")
        
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
    finally:
        os.chdir('..')

if __name__ == "__main__":
    main() 