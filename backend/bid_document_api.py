from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form, Query, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy.orm import Session
from sqlalchemy import desc
from typing import List, Optional, Dict, Any
import os
import shutil
from datetime import datetime
import logging
import uuid
import json
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import get_db
from models import Project, ProjectSection, SectionDocument, Template, TemplateField, TemplateMapping, GeneratedDocument
from ai_service import AIService
from document_processor import document_processor
from document_generator import DocumentGenerator

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 创建路由器
router = APIRouter(prefix="/api/bid-documents", tags=["bid-documents"])

# 上传目录
UPLOAD_DIR = os.environ.get("UPLOAD_PATH", "/app/uploads")
BID_DOCS_DIR = os.path.join(UPLOAD_DIR, "bid_documents")
os.makedirs(BID_DOCS_DIR, exist_ok=True)

# 初始化服务
ai_service = AIService()
document_generator = DocumentGenerator()

# ============ 数据模型 ============

class TenderDocumentAnalysis(BaseModel):
    """招标文件分析结果"""
    tender_company: Optional[str] = None  # 招标人
    tender_agency: Optional[str] = None   # 招标代理机构
    bidder_name: Optional[str] = None     # 投标人
    legal_representative: Optional[str] = None  # 法定代表人
    authorized_representative: Optional[str] = None  # 授权代表
    bid_deadline: Optional[str] = None    # 投标截止时间
    bid_opening_time: Optional[str] = None  # 开标时间
    project_name: Optional[str] = None    # 项目名称
    project_description: Optional[str] = None  # 项目描述
    requirements: Optional[List[str]] = None  # 投标要求
    evaluation_criteria: Optional[List[str]] = None  # 评标标准
    confidence_score: Optional[float] = None  # 置信度

class BidDocumentTemplate(BaseModel):
    """投标文档模板"""
    id: int
    name: str
    type: str
    description: Optional[str] = None
    fields: List[Dict[str, Any]] = []

class BidDocumentSection(BaseModel):
    """投标文档章节"""
    id: int
    title: str
    description: Optional[str] = None
    order: int
    content_type: str  # template, custom, auto_generated
    template_id: Optional[int] = None
    custom_content: Optional[str] = None
    auto_generated_content: Optional[str] = None
    data_sources: Optional[Dict[str, Any]] = None  # 数据源配置

class BidDocumentGeneration(BaseModel):
    """投标文档生成配置"""
    project_id: int
    template_id: Optional[int] = None
    sections: List[BidDocumentSection]
    field_mappings: Dict[str, str]  # 字段映射
    include_attachments: bool = True
    auto_fill_data: bool = True

# ============ API端点 ============

@router.post("/analyze-tender-document")
async def analyze_tender_document(
    file: UploadFile = File(...),
    project_id: Optional[int] = Form(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """分析招标文件，提取关键信息"""
    try:
        # 保存上传的文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        ext = os.path.splitext(file.filename)[1]
        filename = f"tender_doc_{timestamp}_{unique_id}{ext}"
        file_path = os.path.join(BID_DOCS_DIR, filename)
        
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # 使用AI服务分析文档
        analysis_result = await ai_service.analyze_tender_document(file_path)
        
        # 如果指定了项目ID，更新项目信息
        if project_id:
            project = db.query(Project).filter(Project.id == project_id).first()
            if project:
                # 更新项目信息
                if analysis_result.get("tender_company"):
                    project.tender_company = analysis_result["tender_company"]
                if analysis_result.get("tender_agency"):
                    project.tender_agency = analysis_result["tender_agency"]
                if analysis_result.get("project_name"):
                    project.name = analysis_result["project_name"]
                if analysis_result.get("project_description"):
                    project.description = analysis_result["project_description"]
                
                db.commit()
        
        return {
            "success": True,
            "analysis_result": analysis_result,
            "file_path": file_path,
            "filename": filename
        }
        
    except Exception as e:
        logger.error(f"分析招标文件失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"分析招标文件失败: {str(e)}")

@router.get("/templates", response_model=List[BidDocumentTemplate])
async def get_bid_templates(
    template_type: Optional[str] = Query(None),
    db: Session = Depends(get_db)
):
    """获取投标文档模板列表"""
    try:
        query = db.query(Template)
        
        if template_type:
            query = query.filter(Template.type == template_type)
        else:
            # 默认获取投标相关模板
            query = query.filter(Template.type.in_(["bid_document", "cover_letter", "technical_proposal"]))
        
        templates = query.all()
        
        # 获取模板字段
        result = []
        for template in templates:
            fields = db.query(TemplateField).filter(TemplateField.template_id == template.id).all()
            field_list = []
            for field in fields:
                field_list.append({
                    "id": field.id,
                    "field_name": field.field_name,
                    "field_key": field.field_key,
                    "field_type": field.field_type,
                    "placeholder": field.placeholder,
                    "is_required": field.is_required
                })
            
            result.append({
                "id": template.id,
                "name": template.name,
                "type": template.type,
                "description": template.description,
                "fields": field_list
            })
        
        return result
        
    except Exception as e:
        logger.error(f"获取模板列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取模板列表失败: {str(e)}")

@router.post("/generate")
async def generate_bid_document(
    generation_config: BidDocumentGeneration,
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """生成投标文档"""
    try:
        # 验证项目是否存在
        project = db.query(Project).filter(Project.id == generation_config.project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        
        # 后台任务生成文档
        background_tasks.add_task(
            generate_bid_document_background,
            generation_config.dict(),
            str(db.bind.url)
        )
        
        return {
            "success": True,
            "message": "文档生成任务已启动，请稍后查看结果",
            "project_id": generation_config.project_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"启动文档生成任务失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"启动文档生成任务失败: {str(e)}")

@router.get("/projects/{project_id}/sections", response_model=List[BidDocumentSection])
async def get_project_sections(project_id: int, db: Session = Depends(get_db)):
    """获取项目的投标文档章节"""
    try:
        sections = db.query(ProjectSection).filter(
            ProjectSection.project_id == project_id
        ).order_by(ProjectSection.order).all()
        
        result = []
        for section in sections:
            # 获取章节的数据映射
            data_mappings = []
            for mapping in section.data_mappings:
                data_mappings.append({
                    "data_type": mapping.data_type,
                    "data_id": mapping.data_id,
                    "display_order": mapping.display_order,
                    "is_visible": mapping.is_visible,
                    "filter_conditions": mapping.filter_conditions
                })
            
            result.append({
                "id": section.id,
                "title": section.title,
                "description": section.description,
                "order": section.order,
                "content_type": "custom",  # 默认类型
                "template_id": None,
                "custom_content": None,
                "auto_generated_content": None,
                "data_sources": {
                    "mappings": data_mappings
                }
            })
        
        return result
        
    except Exception as e:
        logger.error(f"获取项目章节失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取项目章节失败: {str(e)}")

@router.post("/projects/{project_id}/sections")
async def create_bid_section(
    project_id: int,
    section_data: BidDocumentSection,
    db: Session = Depends(get_db)
):
    """创建投标文档章节"""
    try:
        # 验证项目是否存在
        project = db.query(Project).filter(Project.id == project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        
        # 创建章节
        section = ProjectSection(
            project_id=project_id,
            title=section_data.title,
            description=section_data.description,
            order=section_data.order
        )
        
        db.add(section)
        db.commit()
        db.refresh(section)
        
        # 如果有数据源配置，创建数据映射
        if section_data.data_sources and section_data.data_sources.get("mappings"):
            for mapping_data in section_data.data_sources["mappings"]:
                from models import SectionDataMapping
                mapping = SectionDataMapping(
                    section_id=section.id,
                    data_type=mapping_data["data_type"],
                    data_id=mapping_data["data_id"],
                    display_order=mapping_data.get("display_order", 0),
                    is_visible=mapping_data.get("is_visible", True),
                    filter_conditions=mapping_data.get("filter_conditions", {})
                )
                db.add(mapping)
            
            db.commit()
        
        return {
            "success": True,
            "section_id": section.id,
            "message": "章节创建成功"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建章节失败: {str(e)}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"创建章节失败: {str(e)}")

@router.get("/projects/{project_id}/preview")
async def preview_bid_document(project_id: int, db: Session = Depends(get_db)):
    """预览投标文档内容"""
    try:
        # 获取项目信息
        project = db.query(Project).filter(Project.id == project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        
        # 获取项目章节
        sections = db.query(ProjectSection).filter(
            ProjectSection.project_id == project_id
        ).order_by(ProjectSection.order).all()
        
        # 生成预览内容
        preview_content = []
        
        # 添加项目基本信息
        preview_content.append({
            "type": "project_info",
            "title": "项目基本信息",
            "content": {
                "项目名称": project.name,
                "招标人": project.tender_company,
                "招标代理机构": project.tender_agency,
                "投标人": project.bidder_name,
                "投标截止日期": project.deadline.strftime("%Y-%m-%d") if project.deadline else None,
                "项目描述": project.description
            }
        })
        
        # 添加章节内容
        for section in sections:
            section_content = {
                "type": "section",
                "title": section.title,
                "description": section.description,
                "documents": []
            }
            
            # 获取章节文档
            documents = db.query(SectionDocument).filter(
                SectionDocument.section_id == section.id
            ).order_by(SectionDocument.order).all()
            
            for doc in documents:
                section_content["documents"].append({
                    "filename": doc.original_filename,
                    "file_type": doc.file_type,
                    "status": doc.processing_status,
                    "page_count": doc.page_count
                })
            
            preview_content.append(section_content)
        
        return {
            "success": True,
            "project_id": project_id,
            "preview_content": preview_content
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成预览失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"生成预览失败: {str(e)}")

@router.get("/download/{document_id}")
async def download_bid_document(document_id: int, db: Session = Depends(get_db)):
    """下载生成的投标文档"""
    try:
        document = db.query(GeneratedDocument).filter(GeneratedDocument.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="文档不存在")
        
        if not os.path.exists(document.file_path):
            raise HTTPException(status_code=404, detail="文档文件不存在")
        
        return FileResponse(
            document.file_path,
            filename=document.filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"下载文档失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"下载文档失败: {str(e)}")

# ============ 后台任务函数 ============

async def generate_bid_document_background(generation_config: Dict[str, Any], db_url: str):
    """后台生成投标文档"""
    try:
        from database import get_db
        from sqlalchemy import create_engine
        from sqlalchemy.orm import sessionmaker
        
        # 创建数据库连接
        engine = create_engine(db_url)
        SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
        db = SessionLocal()
        
        try:
            project_id = generation_config["project_id"]
            
            # 获取项目信息
            project = db.query(Project).filter(Project.id == project_id).first()
            if not project:
                logger.error(f"项目 {project_id} 不存在")
                return
            
            # 生成文档
            output_path = await generate_complete_bid_document(
                project, generation_config, db
            )
            
            # 保存生成的文档记录
            document = GeneratedDocument(
                project_id=project_id,
                filename=os.path.basename(output_path),
                file_path=output_path,
                file_size=os.path.getsize(output_path),
                generation_time=0.0  # TODO: 记录实际生成时间
            )
            
            db.add(document)
            db.commit()
            
            logger.info(f"投标文档生成成功: {output_path}")
            
        finally:
            db.close()
            
    except Exception as e:
        logger.error(f"后台生成投标文档失败: {str(e)}")

async def generate_complete_bid_document(project: Project, config: Dict[str, Any], db: Session) -> str:
    """生成完整的投标文档"""
    try:
        # 创建新文档
        from docx import Document
        from docx.shared import Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)
        
        # 添加封面页
        add_cover_page(doc, project, config)
        
        # 添加目录页
        add_table_of_contents(doc)
        
        # 处理各个章节
        for section_config in config.get("sections", []):
            await add_section_content(doc, section_config, project, db)
        
        # 生成文件名和保存
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"bid_document_{project.name}_{timestamp}.docx"
        output_path = os.path.join(BID_DOCS_DIR, filename)
        
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        logger.error(f"生成投标文档失败: {str(e)}")
        raise e

def add_cover_page(doc: Document, project: Project, config: Dict[str, Any]):
    """添加封面页"""
    # 添加标题
    title = doc.add_heading('投标文件', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加项目信息
    doc.add_paragraph()  # 空行
    project_info = doc.add_paragraph()
    project_info.add_run(f"项目名称：{project.name}").bold = True
    project_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    tender_info = doc.add_paragraph()
    tender_info.add_run(f"招标人：{project.tender_company or '待填写'}")
    tender_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if project.tender_agency:
        doc.add_paragraph()
        agency_info = doc.add_paragraph()
        agency_info.add_run(f"招标代理机构：{project.tender_agency}")
        agency_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    bidder_info = doc.add_paragraph()
    bidder_info.add_run(f"投标人：{project.bidder_name or '待填写'}")
    bidder_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    doc.add_paragraph()
    date_info = doc.add_paragraph()
    date_info.add_run(f"投标日期：{datetime.now().strftime('%Y年%m月%d日')}")
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分页符
    doc.add_page_break()

def add_table_of_contents(doc: Document):
    """添加目录页"""
    toc_heading = doc.add_heading('目录', 1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # TODO: 实现自动目录生成
    doc.add_paragraph("目录内容将根据章节自动生成...")
    
    doc.add_page_break()

async def add_section_content(doc: Document, section_config: Dict[str, Any], project: Project, db: Session):
    """添加章节内容"""
    try:
        section_title = section_config.get("title", "未命名章节")
        
        # 添加章节标题
        heading = doc.add_heading(section_title, 1)
        
        # 根据内容类型处理
        content_type = section_config.get("content_type", "custom")
        
        if content_type == "template" and section_config.get("template_id"):
            # 使用模板生成内容
            await add_template_content(doc, section_config["template_id"], project, db)
        
        elif content_type == "auto_generated":
            # 自动生成内容
            await add_auto_generated_content(doc, section_config, project, db)
        
        else:
            # 自定义内容
            custom_content = section_config.get("custom_content", "")
            if custom_content:
                doc.add_paragraph(custom_content)
        
        # 添加分页符
        doc.add_page_break()
        
    except Exception as e:
        logger.error(f"添加章节内容失败: {str(e)}")
        doc.add_paragraph(f"章节内容生成失败: {str(e)}")

async def add_template_content(doc: Document, template_id: int, project: Project, db: Session):
    """添加模板内容"""
    try:
        template = db.query(Template).filter(Template.id == template_id).first()
        if not template:
            doc.add_paragraph("模板不存在")
            return
        
        # 获取模板字段
        fields = db.query(TemplateField).filter(TemplateField.template_id == template_id).all()
        
        # 获取字段映射值
        field_values = {}
        for field in fields:
            mapping = db.query(TemplateMapping).filter(
                TemplateMapping.project_id == project.id,
                TemplateMapping.field_id == field.id
            ).first()
            
            if mapping:
                field_values[field.field_key] = mapping.value
            else:
                field_values[field.field_key] = field.placeholder or ""
        
        # 根据模板类型生成内容
        if template.type == "cover_letter":
            add_cover_letter_content(doc, project, field_values)
        elif template.type == "technical_proposal":
            add_technical_proposal_content(doc, project, field_values)
        else:
            add_generic_template_content(doc, template, field_values)
            
    except Exception as e:
        logger.error(f"添加模板内容失败: {str(e)}")
        doc.add_paragraph(f"模板内容生成失败: {str(e)}")

async def add_auto_generated_content(doc: Document, section_config: Dict[str, Any], project: Project, db: Session):
    """添加自动生成的内容"""
    try:
        data_sources = section_config.get("data_sources", {})
        mappings = data_sources.get("mappings", [])
        
        for mapping in mappings:
            data_type = mapping.get("data_type")
            data_id = mapping.get("data_id")
            
            if data_type == "award":
                await add_award_content(doc, data_id, db)
            elif data_type == "performance":
                await add_performance_content(doc, data_id, db)
            elif data_type == "lawyer_certificate":
                await add_lawyer_content(doc, data_id, db)
                
    except Exception as e:
        logger.error(f"添加自动生成内容失败: {str(e)}")
        doc.add_paragraph(f"自动生成内容失败: {str(e)}")

def add_cover_letter_content(doc: Document, project: Project, field_values: Dict[str, str]):
    """添加封面函内容"""
    # 添加称谓
    doc.add_paragraph(f"{field_values.get('tender_company', project.tender_company or '招标人')}：")
    
    # 添加正文
    doc.add_paragraph()
    doc.add_paragraph(field_values.get('cover_letter_content', '我们郑重承诺...'))
    
    # 添加落款
    doc.add_paragraph()
    doc.add_paragraph(f"投标人：{field_values.get('bidder_name', project.bidder_name or '')}")
    doc.add_paragraph(f"法定代表人：{field_values.get('legal_representative', '')}")
    doc.add_paragraph(f"日期：{datetime.now().strftime('%Y年%m月%d日')}")

def add_technical_proposal_content(doc: Document, project: Project, field_values: Dict[str, str]):
    """添加技术方案内容"""
    # 添加技术方案概述
    doc.add_heading('技术方案概述', 2)
    doc.add_paragraph(field_values.get('technical_overview', '技术方案概述内容...'))
    
    # 添加技术方案详细内容
    doc.add_heading('技术方案详细内容', 2)
    doc.add_paragraph(field_values.get('technical_details', '技术方案详细内容...'))

def add_generic_template_content(doc: Document, template: Template, field_values: Dict[str, str]):
    """添加通用模板内容"""
    for field_key, value in field_values.items():
        if value and value != "待填写":
            doc.add_paragraph(f"{field_key}: {value}")

async def add_award_content(doc: Document, award_id: int, db: Session):
    """添加获奖信息内容"""
    try:
        from models import Award
        award = db.query(Award).filter(Award.id == award_id).first()
        if not award:
            return
        
        # 添加获奖信息
        award_para = doc.add_paragraph()
        award_para.add_run(f"• {award.title} - {award.brand} {award.year}年").bold = True
        
        if award.description:
            doc.add_paragraph(award.description)
            
    except Exception as e:
        logger.error(f"添加获奖信息失败: {str(e)}")

async def add_performance_content(doc: Document, performance_id: int, db: Session):
    """添加业绩信息内容"""
    try:
        from models import Performance
        performance = db.query(Performance).filter(Performance.id == performance_id).first()
        if not performance:
            return
        
        # 添加业绩信息
        performance_para = doc.add_paragraph()
        performance_para.add_run(f"• {performance.project_name} - {performance.client_name}").bold = True
        
        if performance.description:
            doc.add_paragraph(performance.description)
            
    except Exception as e:
        logger.error(f"添加业绩信息失败: {str(e)}")

async def add_lawyer_content(doc: Document, lawyer_id: int, db: Session):
    """添加律师信息内容"""
    try:
        from models import LawyerCertificate
        lawyer = db.query(LawyerCertificate).filter(LawyerCertificate.id == lawyer_id).first()
        if not lawyer:
            return
        
        # 添加律师信息
        lawyer_para = doc.add_paragraph()
        lawyer_para.add_run(f"• {lawyer.lawyer_name} - {lawyer.law_firm}").bold = True
        
        if lawyer.position:
            doc.add_paragraph(f"职位：{lawyer.position}")
            
    except Exception as e:
        logger.error(f"添加律师信息失败: {str(e)}")

def setup_router(app):
    """设置路由器"""
    app.include_router(router) 