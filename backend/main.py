from fastapi import FastAPI, File, UploadFile, Depends, HTTPException, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any, Union
import os
import sys
import shutil
import asyncio
import logging
from datetime import datetime, timezone, timedelta
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from sqlalchemy import desc
import pytz
from sqlalchemy.sql import func
from docxcompose.composer import Composer
from docx import Document as DocxDocument

# 添加当前目录到sys.path以确保模块可以被导入
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from database import get_db, init_db, SessionLocal, engine
from models import *
from ai_service import ai_service
from screenshot_service import screenshot_service
from document_generator import document_generator
import schemas
from schemas import (
    Award as AwardSchema, AwardCreate, AwardResponse,
    Performance as PerformanceSchema, PerformanceCreate, PerformanceResponse,
    Project as ProjectSchema, ProjectCreate, ProjectResponse,
    SystemSettingsResponse, BrandResponse, BusinessFieldResponse, BaseResponse
)

# 导入新的API模块（确保它们存在于同一目录下）
try:
    import project_api
    import template_api
    import section_api
    import search_api
    import ai_tools_api
    import file_management_api
    import bid_document_api
    from document_processor import docling_processor, format_heading_standalone
    IMPORT_SUCCESS = True
except ImportError as e:
    logging.error(f"导入API模块失败: {str(e)}")
    IMPORT_SUCCESS = False

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 创建FastAPI应用
APP_NAME = os.getenv("APP_NAME", "投标苦")
app = FastAPI(
    title=APP_NAME,
    description="法律行业投标资料管理和生成系统，包括投标文件自动组装功能",
    version="1.0.0"
)

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://frontend:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 挂载静态文件
app.mount("/screenshots", StaticFiles(directory="/app/screenshots"), name="screenshots")
app.mount("/uploads", StaticFiles(directory="/app/uploads"), name="uploads")
app.mount("/generated", StaticFiles(directory="/app/generated_docs"), name="generated")

# 创建必要的目录
os.makedirs("/app/uploads", exist_ok=True)
os.makedirs("/app/screenshots", exist_ok=True)
os.makedirs("/app/generated_docs", exist_ok=True)

MAX_UPLOAD_SIZE_MB = int(os.getenv('MAX_UPLOAD_SIZE_MB', '50'))
MAX_UPLOAD_SIZE = MAX_UPLOAD_SIZE_MB * 1024 * 1024
HISTORY_FILE = "/app/generated_docs/convert_history.json"

def format_heading(doc, text, level=1, center=False):
    """
    创建格式化的标题
    
    参数:
    - doc: Document对象
    - text: 标题文字
    - level: 标题级别 (0=主标题, 1=一级标题, 2=二级标题)
    - center: 是否居中
    
    返回:
    - 格式化的标题段落
    """
    from docx.oxml.shared import qn
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
    
    # 创建标题
    heading = doc.add_heading(text, level)
    
    # 设置对齐方式
    if center or level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # 清除段落的列表样式和项目符号，防止出现小黑点
        try:
            # 清除段落的编号和项目符号
            pPr = heading._element.get_or_add_pPr()
            
            # 移除编号属性
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            
            # 设置段落格式，明确禁用列表样式和分页控制
            if hasattr(heading, 'paragraph_format'):
                heading.paragraph_format.left_indent = None
                heading.paragraph_format.first_line_indent = None
                
                # 禁用分页控制选项，对应Word中的"分页"设置
                heading.paragraph_format.widow_control = False      # 孤行控制
                heading.paragraph_format.keep_with_next = False     # 与下段同页
                heading.paragraph_format.keep_together = False      # 段中不分页
                heading.paragraph_format.page_break_before = False  # 段前分页
                
        except Exception as style_error:
            logger.warning(f"清除标题样式时出错: {style_error}")
    
    # 设置字体
    if heading.runs:
        run = heading.runs[0]
        
        # 根据级别设置字体大小
        if level == 0:
            run.font.size = Pt(18)
            run.font.name = '楷体'
        elif level == 1:
            run.font.size = Pt(16)
            run.font.name = '楷体'
        elif level == 2:
            run.font.size = Pt(14)
            run.font.name = '楷体'
        else:
            run.font.size = Pt(12)
            run.font.name = '楷体'
        
        run.bold = True
        
        # 确保字体颜色是黑色
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
        
        # 设置中英文字体
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    return heading

@app.on_event("startup")
async def startup_event():
    """应用启动时初始化数据库"""
    try:
        init_db()
        logger.info("数据库初始化完成")
        
        # 初始化基础数据
        await init_base_data()
        
        # 初始化高级功能数据
        await init_advanced_data()
        
        # 启动时下载Docling模型
        await download_docling_models_on_startup()
        
        # 注册API路由
        if IMPORT_SUCCESS:
            # 注册项目API路由
            project_api.setup_router(app)
            logger.info("项目API路由注册成功")
            
            # 注册模板API路由
            template_api.setup_router(app)
            logger.info("模板API路由注册成功")
            
            # 注册智能章节管理API路由
            app.include_router(section_api.router)
            logger.info("智能章节管理API路由注册成功")
            
            # 注册AI自动检索API路由
            app.include_router(search_api.router)
            logger.info("AI自动检索API路由注册成功")
            
            # 注册投标文档制作API路由
            bid_document_api.setup_router(app)
            logger.info("投标文档制作API路由注册成功")
            
            # 注册AI工具API路由
            app.include_router(ai_tools_api.router)
            logger.info("AI工具API路由注册成功")
            
            # 注册文件管理API路由
            file_management_api.setup_router(app)
            logger.info("文件管理API路由注册成功")
            
            # 注册律师证管理API路由
            try:
                import lawyer_certificate_api
                lawyer_certificate_api.setup_router(app)
                logger.info("律师证管理API路由注册成功")
            except ImportError as e:
                logger.warning(f"律师证管理API路由注册失败: {str(e)}")
            
            # 注册业绩管理API路由
            try:
                import performance_api
                app.include_router(performance_api.router, prefix="/api/performances", tags=["performances"])
                logger.info("业绩管理API路由注册成功")
            except ImportError as e:
                logger.warning(f"业绩管理API路由注册失败: {str(e)}")
            
            # 注册奖项管理API路由
            try:
                import award_api
                app.include_router(award_api.router, prefix="/api/awards", tags=["awards"])
                logger.info("奖项管理API路由注册成功")
            except ImportError as e:
                logger.warning(f"奖项管理API路由注册失败: {str(e)}")
        else:
            logger.error("API路由注册失败，模块导入错误")
        
    except Exception as e:
        logger.error(f"应用启动失败: {str(e)}")

async def init_base_data():
    """初始化基础数据，如厂牌、业务领域等"""
    db = SessionLocal()
    try:
        # 检查是否已有厂牌数据
        if db.query(Brand).count() == 0:
            default_brands = [
                Brand(name="品牌A", full_name="品牌A全称", website="http://example.com/a"),
                Brand(name="品牌B", full_name="品牌B全称", website="http://example.com/b"),
                Brand(name="品牌C", full_name="品牌C全称", website="http://example.com/c", is_active=False),
                Brand(name="通用", full_name="通用/其他", website=""),
            ]
            db.add_all(default_brands)
            logger.info(f"已初始化 {len(default_brands)} 个默认厂牌")

        # 检查是否已有业务领域数据
        if db.query(BusinessField).count() == 0:
            default_fields = [
                BusinessField(name="领域X", description="业务领域X的描述"),
                BusinessField(name="领域Y", description="业务领域Y的描述"),
                BusinessField(name="领域Z", description="业务领域Z的描述"),
            ]
            db.add_all(default_fields)
            logger.info(f"已初始化 {len(default_fields)} 个默认业务领域")

        # 检查是否已有奖项数据
        if db.query(Award).count() == 0:
            default_awards = [
                Award(title="年度最佳法律顾问奖", brand="品牌A", year=2023, business_type="领域X", description="在公司并购领域表现卓越。", is_verified=True),
                Award(title="金融科技创新奖", brand="品牌B", year=2024, business_type="领域Y", description="为金融科技行业带来革命性创新。", is_verified=False),
            ]
            db.add_all(default_awards)
            logger.info(f"已初始化 {len(default_awards)} 个默认奖项")

        # 检查是否已有业绩数据 - 移除默认数据创建
        # if db.query(Performance).count() == 0:
        #     default_performances = [
        #         Performance(client_name="大型科技公司", project_name="Alpha项目-跨国并购案", project_type="并购重组", business_field="领域X", year=2023, contract_amount=1000000, description="一个复杂的跨国并购项目。"),
        #         Performance(client_name="初创金融公司", project_name="Beta项目-支付系统合规审查", project_type="合规与监管", business_field="领域Y", year=2024, contract_amount=500000, description="确保其支付系统符合所有相关法规。"),
        #     ]
        #     db.add_all(default_performances)
        #     logger.info(f"已初始化 {len(default_performances)} 个默认业绩")
        # 业绩数据现在从空开始，用户手动添加
        logger.info("业绩管理已准备就绪，数据从空开始")
            
        db.commit()
    except Exception as e:
        logger.error(f"基础数据初始化失败: {str(e)}")
        db.rollback()
    finally:
        db.close()

async def init_advanced_data():
    """初始化高级功能数据"""
    try:
        # 导入高级功能初始化模块
        from init_advanced_data import (
            init_section_types, init_data_sources, init_recommendation_rules,
            init_advanced_brands, init_advanced_business_fields
        )
        
        # 执行高级功能初始化
        init_section_types()
        init_data_sources()
        init_recommendation_rules()
        init_advanced_brands()
        init_advanced_business_fields()
        
        logger.info("高级功能数据初始化完成")
    except Exception as e:
        logger.error(f"高级功能数据初始化失败: {str(e)}")

async def download_docling_models_on_startup():
    """启动时下载Docling模型"""
    try:
        logger.info("🚀 开始启动时下载Docling模型...")
        
        # 检查Docling服务是否可用
        try:
            from docling_service import docling_service
            if docling_service and docling_service.is_initialized:
                logger.info("✅ Docling服务已初始化，模型已就绪")
                return
        except ImportError:
            logger.warning("⚠️ Docling服务不可用，跳过模型下载")
            return
        
        # 设置环境变量以优化下载
        import os
        os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
        os.environ["HUGGINGFACE_HUB_URL"] = "https://hf-mirror.com"
        os.environ["HF_HOME"] = "/root/.cache/huggingface"
        os.environ["TRANSFORMERS_CACHE"] = "/root/.cache/huggingface/transformers"
        os.environ["HF_HUB_CACHE"] = "/root/.cache/huggingface/hub"
        
        # 禁用hf_transfer以避免错误
        os.environ["HF_HUB_ENABLE_HF_TRANSFER"] = "0"
        
        # 创建缓存目录
        import subprocess
        cache_dirs = [
            "/root/.cache/huggingface",
            "/root/.cache/huggingface/transformers", 
            "/root/.cache/huggingface/hub",
            "/root/.cache/docling/models"
        ]
        
        for cache_dir in cache_dirs:
            os.makedirs(cache_dir, exist_ok=True)
        
        logger.info("📁 缓存目录已创建")
        
        # 使用docling-tools下载模型
        try:
            logger.info("⬇️ 开始下载Docling模型...")
            
            # 下载核心模型
            cmd = [
                "docling-tools", "models", "download",
                "--output-dir", "/root/.cache/docling/models",
                "--models", "layout,tableformer",
                "--force"
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300  # 5分钟超时
            )
            
            if result.returncode == 0:
                logger.info("✅ Docling模型下载成功")
                logger.info(f"📝 下载输出: {result.stdout}")
            else:
                logger.warning(f"⚠️ Docling模型下载失败: {result.stderr}")
                logger.info("🔄 尝试备用下载方法...")
                
                # 备用方法：直接初始化DoclingService
                try:
                    from docling_service import docling_service
                    if docling_service:
                        docling_service._download_models()
                        logger.info("✅ 备用下载方法成功")
                except Exception as backup_error:
                    logger.error(f"❌ 备用下载方法也失败: {backup_error}")
                    
        except subprocess.TimeoutExpired:
            logger.warning("⏰ 模型下载超时，将在后台继续")
        except Exception as e:
            logger.error(f"❌ 模型下载过程出错: {e}")
        
        # 设置OCR为开启状态
        try:
            from database import get_db
            from models import SystemSettings
            
            db = next(get_db())
            
            # 检查是否已有OCR设置
            ocr_setting = db.query(SystemSettings).filter(
                SystemSettings.setting_key == "docling_enable_ocr"
            ).first()
            
            if ocr_setting:
                ocr_setting.setting_value = "true"
                logger.info("🔧 已更新OCR设置为开启状态")
            else:
                # 创建新的OCR设置
                new_setting = SystemSettings(
                    setting_key="docling_enable_ocr",
                    setting_value="true",
                    setting_type="boolean",
                    category="ocr",
                    description="是否启用Docling OCR功能"
                )
                db.add(new_setting)
                logger.info("🔧 已创建OCR设置为开启状态")
            
            db.commit()
            logger.info("✅ OCR功能已设置为开启状态")
            
        except Exception as e:
            logger.error(f"❌ 设置OCR状态失败: {e}")
        finally:
            try:
                if 'db' in locals():
                    db.close()
            except:
                pass
        
        logger.info("🎉 Docling模型启动下载完成")
        
    except Exception as e:
        logger.error(f"❌ Docling模型启动下载失败: {e}")

# ==================== 文件上传和处理 ====================

@app.post("/api/upload/document")
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(...),  # "award" or "performance"
    db: Session = Depends(get_db)
):
    """上传文档并进行AI分析"""
    try:
        # 保存上传的文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{timestamp}_{file.filename}"
        filepath = os.path.join("/app/uploads", filename)
        
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # 根据文件类型进行处理
        if file.filename.lower().endswith('.pdf'):
            extracted_data = await ai_service.extract_text_from_pdf(filepath)
        elif file.filename.lower().endswith('.docx'):
            extracted_data = await ai_service.extract_text_from_docx(filepath)
        elif file.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            extracted_data = await ai_service.extract_text_from_image(filepath)
        else:
            raise HTTPException(status_code=400, detail="不支持的文件格式")
        
        if extracted_data.get("error"):
            raise HTTPException(status_code=500, detail=f"文件处理失败: {extracted_data['error']}")
        
        # AI分析文档内容
        if document_type == "award":
            text_content = extracted_data.get("full_text") or "\n".join([item.get("text", "") for item in extracted_data.get("content", [])])
            analysis_result = await ai_service.analyze_award_document(text_content)
        else:
            text_content = extracted_data.get("full_text") or "\n".join([item.get("text", "") for item in extracted_data.get("content", [])])
            analysis_result = await ai_service.analyze_performance_document(text_content)
        
        return {
            "success": True,
            "file_path": filepath,
            "extracted_data": extracted_data,
            "analysis_result": analysis_result,
            "document_type": document_type
        }
        
    except Exception as e:
        logger.error(f"文档上传处理失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/process/screenshot")
async def process_screenshot(
    urls: List[str],
    award_id: Optional[int] = None
):
    """处理网页截图"""
    try:
        if not urls:
            raise HTTPException(status_code=400, detail="URL列表不能为空")
        
        # 批量截图
        results = await screenshot_service.capture_award_pages(urls, award_id)
        
        return {
            "success": True,
            "results": results
        }
        
    except Exception as e:
        logger.error(f"网页截图处理失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# ==================== 获奖信息管理 ====================

@app.post("/api/awards")
async def create_award(award_data: Dict[str, Any], db: Session = Depends(get_db)):
    """创建获奖记录"""
    try:
        award = Award(
            title=award_data.get("title"),
            brand=award_data.get("brand"),
            year=award_data.get("year"),
            business_type=award_data.get("business_type"),
            description=award_data.get("description"),
            source_document=award_data.get("source_document"),
            source_url=award_data.get("source_url"),
            screenshot_path=award_data.get("screenshot_path"),
            screenshot_pages=award_data.get("screenshot_pages"),
            ai_analysis=award_data.get("ai_analysis"),
            confidence_score=award_data.get("confidence_score", 0.0),
            is_manual_input=award_data.get("is_manual_input", False)
        )
        
        db.add(award)
        db.commit()
        db.refresh(award)
        
        return {"success": True, "award_id": award.id}
        
    except Exception as e:
        logger.error(f"创建获奖记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/awards")
async def get_awards(
    brand: Optional[str] = Query(None),
    business_type: Optional[str] = Query(None),
    year: Optional[int] = Query(None),
    verified_only: bool = Query(False),
    skip: int = Query(0),
    limit: int = Query(100),
    db: Session = Depends(get_db)
):
    """获取获奖记录列表"""
    try:
        query = db.query(Award)
        
        if brand:
            query = query.filter(Award.brand == brand)
        if business_type:
            query = query.filter(Award.business_type == business_type)
        if year:
            query = query.filter(Award.year == year)
        if verified_only:
            query = query.filter(Award.is_verified == True)
        
        total = query.count()
        awards = query.offset(skip).limit(limit).all()
        
        return {
            "success": True,
            "total": total,
            "awards": [
                {
                    "id": award.id,
                    "title": award.title,
                    "brand": award.brand,
                    "year": award.year,
                    "business_type": award.business_type,
                    "description": award.description,
                    "is_verified": award.is_verified,
                    "confidence_score": award.confidence_score,
                    "screenshot_pages": award.screenshot_pages,
                    "created_at": award.created_at.isoformat()
                }
                for award in awards
            ]
        }
        
    except Exception as e:
        logger.error(f"获取获奖记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/awards/{award_id}")
async def update_award(
    award_id: int,
    award_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """更新获奖记录"""
    try:
        award = db.query(Award).filter(Award.id == award_id).first()
        if not award:
            raise HTTPException(status_code=404, detail="获奖记录不存在")
        
        # 更新字段
        for key, value in award_data.items():
            if hasattr(award, key):
                setattr(award, key, value)
        
        award.updated_at = datetime.utcnow()
        
        db.commit()
        db.refresh(award)
        
        return {"success": True, "message": "获奖记录更新成功"}
        
    except Exception as e:
        logger.error(f"更新获奖记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/awards/{award_id}")
async def delete_award(award_id: int, db: Session = Depends(get_db)):
    """删除获奖记录"""
    try:
        award = db.query(Award).filter(Award.id == award_id).first()
        if not award:
            raise HTTPException(status_code=404, detail="获奖记录不存在")
        
        db.delete(award)
        db.commit()
        
        return {"success": True, "message": "获奖记录删除成功"}
        
    except Exception as e:
        logger.error(f"删除获奖记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# ==================== 业绩信息管理 ====================

@app.post("/api/performances")
async def create_performance(performance_data: Dict[str, Any], db: Session = Depends(get_db)):
    """创建业绩记录"""
    try:
        performance = Performance(
            client_name=performance_data.get("client_name"),
            project_name=performance_data.get("project_name"),
            project_type=performance_data.get("project_type"),
            business_field=performance_data.get("business_field"),
            start_date=datetime.fromisoformat(performance_data["start_date"]) if performance_data.get("start_date") else None,
            end_date=datetime.fromisoformat(performance_data["end_date"]) if performance_data.get("end_date") else None,
            year=performance_data.get("year"),
            contract_amount=performance_data.get("contract_amount"),
            currency=performance_data.get("currency", "CNY"),
            description=performance_data.get("description"),
            source_document=performance_data.get("source_document"),
            contract_file=performance_data.get("contract_file"),
            ai_analysis=performance_data.get("ai_analysis"),
            confidence_score=performance_data.get("confidence_score", 0.0),
            is_manual_input=performance_data.get("is_manual_input", False)
        )
        
        db.add(performance)
        db.commit()
        db.refresh(performance)
        
        return {"success": True, "performance_id": performance.id}
        
    except Exception as e:
        logger.error(f"创建业绩记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/performances", response_model=PerformanceResponse)
def get_performances(
    business_field: Optional[str] = Query(None),
    project_type: Optional[str] = Query(None),
    year: Optional[int] = Query(None),
    verified_only: bool = Query(False),
    skip: int = Query(0),
    limit: int = Query(100),
    db: Session = Depends(get_db)
):
    """获取业绩数据列表"""
    try:
        query = db.query(Performance)

        if business_field:
            query = query.filter(Performance.business_field == business_field)

        if project_type:
            query = query.filter(Performance.project_type == project_type)

        if year:
            query = query.filter(Performance.year == year)

        if verified_only:
            query = query.filter(Performance.is_verified == True)

        # 按更新时间降序排列
        query = query.order_by(desc(Performance.updated_at))

        performances = query.offset(skip).limit(limit).all()
        
        # 将SQLAlchemy模型转换为Pydantic模型
        performance_schemas = [PerformanceSchema.from_orm(p) for p in performances]

        return {
            "success": True,
            "message": "业绩数据获取成功",
            "performances": performance_schemas
        }

    except Exception as e:
        logger.error(f"获取业绩数据失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取业绩数据失败: {str(e)}")

# ==================== 文档生成 ====================

@app.post("/api/generate/document")
async def generate_document(
    request_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """生成Word文档"""
    try:
        doc_type = request_data.get("type", "combined")  # "award", "performance", "combined"
        filters = request_data.get("filters", {})
        
        # 构建查询
        award_query = db.query(Award)
        performance_query = db.query(Performance)
        
        # 应用筛选条件
        if filters.get("brands"):
            award_query = award_query.filter(Award.brand.in_(filters["brands"]))
        
        if filters.get("business_fields"):
            award_query = award_query.filter(Award.business_type.in_(filters["business_fields"]))
            performance_query = performance_query.filter(Performance.business_field.in_(filters["business_fields"]))
        
        if filters.get("years"):
            award_query = award_query.filter(Award.year.in_(filters["years"]))
            performance_query = performance_query.filter(Performance.year.in_(filters["years"]))
        
        # 获取数据
        awards = []
        performances = []
        
        if doc_type in ["award", "combined"]:
            awards = award_query.all()
            awards = [
                {
                    "id": award.id,
                    "title": award.title,
                    "brand": award.brand,
                    "year": award.year,
                    "business_type": award.business_type,
                    "description": award.description,
                    "screenshot_pages": award.screenshot_pages or []
                }
                for award in awards
            ]
        
        if doc_type in ["performance", "combined"]:
            performances = performance_query.all()
            performances = [
                {
                    "id": perf.id,
                    "client_name": perf.client_name,
                    "project_name": perf.project_name,
                    "project_type": perf.project_type,
                    "business_field": perf.business_field,
                    "year": perf.year,
                    "contract_amount": perf.contract_amount,
                    "currency": perf.currency,
                    "description": perf.description,
                    "start_date": perf.start_date.strftime("%Y-%m-%d") if perf.start_date else None,
                    "end_date": perf.end_date.strftime("%Y-%m-%d") if perf.end_date else None,
                    "files": []  # TODO: 添加文件关联
                }
                for perf in performances
            ]
        
        # 生成文档
        if doc_type == "award":
            filepath = document_generator.generate_award_document(awards, filters)
        elif doc_type == "performance":
            filepath = document_generator.generate_performance_document(performances, filters)
        else:
            filepath = document_generator.generate_combined_document(awards, performances, filters)
        
        # 返回文件下载链接
        filename = os.path.basename(filepath)
        return {
            "success": True,
            "download_url": f"/api/download/{filename}",
            "filename": filename
        }
        
    except Exception as e:
        logger.error(f"文档生成失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    """下载生成的文档"""
    # 先在 generated_docs 目录查找（生成的文档）
    generated_filepath = os.path.join("/app/generated_docs", filename)
    if os.path.exists(generated_filepath):
        return FileResponse(
            path=generated_filepath,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    # 再在 uploads 目录查找（历史遗留文件）
    uploads_filepath = os.path.join("/app/uploads", filename)
    if os.path.exists(uploads_filepath):
        return FileResponse(
            path=uploads_filepath,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    raise HTTPException(status_code=404, detail="文件不存在")

# ==================== 配置管理 ====================

@app.get("/api/config/brands")
async def get_brands(db: Session = Depends(get_db)):
    """获取厂牌列表"""
    brands = db.query(Brand).filter(Brand.is_active == True).all()
    return {
        "success": True,
        "brands": [
            {
                "id": brand.id,
                "name": brand.name,
                "full_name": brand.full_name,
                "website": brand.website
            }
            for brand in brands
        ]
    }

@app.get("/api/config/business-fields")
async def get_business_fields(db: Session = Depends(get_db)):
    """获取业务领域列表"""
    fields = db.query(BusinessField).filter(BusinessField.is_active == True).all()
    return {
        "success": True,
        "business_fields": [
            {
                "id": field.id,
                "name": field.name,
                "description": field.description,
                "parent_id": field.parent_id
            }
            for field in fields
        ]
    }

# ==================== 健康检查 ====================

@app.get("/api/health")
async def health_check():
    """健康检查接口"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0"
    }

@app.get("/api/app-info")
async def get_app_info():
    """获取应用信息"""
    return {
        "app_name": APP_NAME,
        "version": "1.0.0",
        "description": "法律行业投标资料管理和生成系统",
        "max_upload_size_mb": MAX_UPLOAD_SIZE_MB
    }

# ==================== 系统设置管理 ====================

@app.get("/api/settings")
async def get_settings(category: Optional[str] = None, db: Session = Depends(get_db)):
    """获取系统设置"""
    try:
        query = db.query(SystemSettings)
        if category:
            query = query.filter(SystemSettings.category == category)
        
        settings = query.all()
        
        # 将设置转换为字典格式
        settings_dict = {}
        for setting in settings:
            # 敏感信息进行脱敏处理
            value = setting.setting_value
            if setting.is_sensitive and value:
                value = "******" if len(value) > 6 else "*" * len(value)
            
            settings_dict[setting.setting_key] = {
                "value": value,
                "type": setting.setting_type,
                "category": setting.category,
                "description": setting.description,
                "is_sensitive": setting.is_sensitive
            }
        
        return {
            "success": True,
            "settings": settings_dict
        }
        
    except Exception as e:
        logger.error(f"获取设置失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/settings")
async def update_settings(
    settings_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """更新系统设置"""
    try:
        updated_settings = []
        
        for key, value in settings_data.items():
            # 查找现有设置
            existing_setting = db.query(SystemSettings).filter(
                SystemSettings.setting_key == key
            ).first()
            
            if existing_setting:
                # 对于敏感信息，如果值为脱敏符号则跳过更新
                if existing_setting.is_sensitive and str(value) == "******":
                    logger.info(f"跳过敏感信息 {key} 的脱敏值更新")
                    continue
                # 更新现有设置
                existing_setting.setting_value = str(value)
                existing_setting.updated_at = datetime.utcnow()
            else:
                # 创建新设置（根据键名推断分类）
                category = "general"
                if key.startswith("ai_") or "model" in key.lower():
                    category = "ai"
                elif key.startswith("upload_"):
                    category = "upload"
                elif key.startswith("screenshot_"):
                    category = "screenshot"
                elif key.startswith("easyocr_") or key.startswith("docling_") or key == "enable_docling_ocr":
                    category = "ocr"
                
                new_setting = SystemSettings(
                    setting_key=key,
                    setting_value=str(value),
                    category=category,
                    is_sensitive="api_key" in key.lower() or "secret" in key.lower()
                )
                db.add(new_setting)
            
            updated_settings.append(key)
        
        db.commit()
        
        # 重新初始化AI服务以应用新的模型配置
        if any("model" in key.lower() or "ai_" in key for key in updated_settings):
            ai_service.reload_config()
        
        return {
            "success": True,
            "message": f"已更新 {len(updated_settings)} 个设置",
            "updated_settings": updated_settings
        }
        
    except Exception as e:
        logger.error(f"更新设置失败: {str(e)}")
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/settings/init-defaults")
async def init_default_settings(db: Session = Depends(get_db)):
    """初始化默认设置"""
    try:
        default_settings = [
            # AI模型设置
            {
                "key": "ai_provider",
                "value": "openai",
                "category": "ai",
                "description": "AI服务提供商 (openai/azure/custom)"
            },
            {
                "key": "ai_text_model",
                "value": "gpt-4",
                "category": "ai",
                "description": "文本分析模型"
            },
            {
                "key": "ai_vision_model",
                "value": "gpt-4-vision-preview",
                "category": "ai",
                "description": "图像分析模型"
            },
            {
                "key": "ai_api_key",
                "value": "",
                "category": "ai",
                "description": "AI API密钥",
                "is_sensitive": True
            },
            {
                "key": "ai_base_url",
                "value": "https://api.openai.com/v1",
                "category": "ai",
                "description": "AI API基础URL"
            },
            # 视觉模型设置
            {
                "key": "vision_provider",
                "value": "openai",
                "category": "vision",
                "description": "视觉模型提供商 (openai/ollama/azure/custom)"
            },
            {
                "key": "vision_base_url",
                "value": "",
                "category": "vision",
                "description": "视觉模型API基础URL（留空则使用ai_base_url）"
            },
            {
                "key": "vision_api_key",
                "value": "",
                "category": "vision",
                "description": "视觉模型API密钥（留空则使用ai_api_key）",
                "is_sensitive": True
            },
            {
                "key": "ollama_vision_base_url",
                "value": "http://localhost:11434/v1",
                "category": "vision",
                "description": "Ollama视觉模型服务地址"
            },

            # 上传设置
            {
                "key": "upload_max_file_size",
                "value": "524288000",
                "category": "upload",
                "description": "最大文件大小（字节）"
            },
            {
                "key": "upload_allowed_types",
                "value": "pdf,docx,doc,png,jpg,jpeg",
                "category": "upload",
                "description": "允许的文件类型"
            },
            # 截图设置
            {
                "key": "screenshot_timeout",
                "value": "30",
                "category": "screenshot",
                "description": "截图超时时间（秒）"
            },
            {
                "key": "screenshot_max_pages",
                "value": "20",
                "category": "screenshot",
                "description": "最大截图页数"
            },
            # OCR设置 - 默认启用EasyOCR
            {
                "key": "enable_docling_ocr",
                "value": "true",
                "category": "ocr",
                "description": "启用Docling OCR"
            },
            {
                "key": "docling_ocr_languages",
                "value": '["ch_sim", "en"]',
                "category": "ocr", 
                "description": "Docling OCR支持的语言"
            },
            {
                "key": "easyocr_enable",
                "value": "true",
                "category": "ocr",
                "description": "启用EasyOCR (默认开启)"
            },
            {
                "key": "easyocr_model_path",
                "value": "/easyocr_models",
                "category": "ocr",
                "description": "EasyOCR模型存储路径"
            },
            {
                "key": "easyocr_languages",
                "value": '["ch_sim", "en"]',
                "category": "ocr",
                "description": "EasyOCR支持的语言"
            },
            {
                "key": "easyocr_use_gpu",
                "value": "false",
                "category": "ocr",
                "description": "EasyOCR是否使用GPU"
            },
            {
                "key": "easyocr_download_proxy",
                "value": "",
                "category": "ocr",
                "description": "EasyOCR模型下载代理"
            },
            # AI分析高级设置
            {
                "key": "ai_analysis_enable_table_structure",
                "value": "false",
                "category": "ai_analysis",
                "description": "AI分析时是否启用表格结构分析"
            },
            {
                "key": "ai_analysis_enable_picture_classification",
                "value": "false",
                "category": "ai_analysis",
                "description": "AI分析时是否启用图片分类"
            },
            {
                "key": "ai_analysis_enable_picture_description",
                "value": "false",
                "category": "ai_analysis",
                "description": "AI分析时是否启用图片描述"
            },
            {
                "key": "ai_analysis_generate_page_images",
                "value": "false",
                "category": "ai_analysis",
                "description": "AI分析时是否生成页面图片"
            },
            {
                "key": "ai_analysis_generate_picture_images",
                "value": "false",
                "category": "ai_analysis",
                "description": "AI分析时是否生成图片文件"
            },
            # OCR精度优化设置
            {
                "key": "docling_confidence_threshold",
                "value": "0.5",
                "category": "ocr",
                "description": "OCR识别置信度阈值 (0.1-1.0)"
            },
            {
                "key": "docling_bitmap_area_threshold",
                "value": "0.05",
                "category": "ocr",
                "description": "OCR位图区域阈值 (0.01-0.1)"
            },
            {
                "key": "docling_force_full_page_ocr",
                "value": "false",
                "category": "ocr",
                "description": "是否强制全页OCR (可能提高漏行检测)"
            },
            {
                "key": "docling_recog_network",
                "value": "standard",
                "category": "ocr",
                "description": "OCR识别网络类型 (standard/fast)"
            },
            {
                "key": "docling_use_gpu",
                "value": "false",
                "category": "ocr",
                "description": "OCR是否使用GPU加速"
            },
            {
                "key": "docling_images_scale",
                "value": "2.0",
                "category": "ocr",
                "description": "OCR图像缩放比例 (1.0-3.0)"
            },
            # 图片描述配置
            {
                "key": "picture_description_prompt",
                "value": "请详细描述这张图片的内容，包括文字、图表、结构等信息。",
                "category": "ai_analysis",
                "description": "Docling图片描述的提示词"
            }
        ]
        
        # 检查SystemSettings表是否存在
        try:
            # 尝试创建表
            SystemSettings.__table__.create(bind=db.get_bind(), checkfirst=True)
            logger.info("创建SystemSettings表成功")
        except Exception as table_err:
            logger.warning(f"创建表失败或表已存在: {str(table_err)}")
        
        # 创建设置
        created_count = 0
        for setting_info in default_settings:
            # 检查设置是否已存在
            existing = db.query(SystemSettings).filter(
                SystemSettings.setting_key == setting_info["key"]
            ).first()
            
            if not existing:
                new_setting = SystemSettings(
                    setting_key=setting_info["key"],
                    setting_value=setting_info["value"],
                    category=setting_info["category"],
                    description=setting_info["description"],
                    is_sensitive=setting_info.get("is_sensitive", False)
                )
                db.add(new_setting)
                created_count += 1
        
        if created_count > 0:
            db.commit()
            logger.info(f"已初始化 {created_count} 个默认设置")
        
    except Exception as e:
        logger.error(f"初始化默认设置失败: {str(e)}")
        db.rollback()

@app.post("/api/convert-to-word")
async def convert_files_to_word(
    files: List[UploadFile] = File(default=[]),
    document_title: str = Form(default="转换文档"),
    show_main_title: bool = Form(default=True),
    show_file_titles: bool = Form(default=True),
    main_title_level: int = Form(default=1),
    file_title_level: int = Form(default=2),
    enable_numbering: bool = Form(default=True),
    enable_watermark: bool = Form(default=False),
    file_watermark_settings: str = Form(default="[]"),
    file_page_break_settings: str = Form(default="[]"),
    file_page_number_settings: str = Form(default="[]"),
    permanent_file_ids: str = Form(default="[]"),
    watermark_text: str = Form(default=""),
    watermark_font_size: int = Form(default=24),
    watermark_angle: int = Form(default=-45),
    watermark_opacity: int = Form(default=30),
    watermark_color: str = Form(default="#808080"),
    watermark_position: str = Form(default="center"),
    db: Session = Depends(get_db)
):
    import json
    import shutil
    from datetime import datetime, timedelta
    import os
    from docx import Document as DocxDocument
    from docxcompose.composer import Composer
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from PIL import Image
    from docx.oxml.shared import qn

    logger.info(f"收到转换请求 - 文档标题: {document_title}, 文件数量: {len(files)}")

    temp_dir = "temp_conversions"
    os.makedirs(temp_dir, exist_ok=True)
    file_paths = []
    file_names = []
    for file in files:
        file_path = os.path.join(temp_dir, file.filename)
        content = await file.read()
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        file_paths.append(file_path)
        file_names.append(file.filename)

    # 解析每个文件的页码设置
    try:
        file_page_number_list = json.loads(file_page_number_settings)
    except:
        file_page_number_list = []

    # 创建主文档
    main_doc = DocxDocument()
    composer = Composer(main_doc)
    # 设置A4页面
    section = main_doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # 主标题
    if show_main_title:
        para = main_doc.add_paragraph()
        run = para.add_run(document_title)
        run.bold = True
        run.font.size = Pt(18)
        # 设置中英文字体
        run.font.name = '楷体'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_files = len(file_paths)
    for i, file_path in enumerate(file_paths):
        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(filename)[1].lower()
        is_last_file = (i == total_files - 1)
        
        # 检查当前文件是否需要水印
        file_needs_watermark = False
        if enable_watermark and watermark_text:
            try:
                file_watermark_list = json.loads(file_watermark_settings)
                if i < len(file_watermark_list):
                    file_needs_watermark = file_watermark_list[i]
            except:
                file_needs_watermark = enable_watermark  # 如果解析失败，默认启用
        
        # 文件标题
        if show_file_titles:
            para = main_doc.add_paragraph()
            run = para.add_run(os.path.splitext(filename)[0])
            run.bold = True
            run.font.size = Pt(16)
            # 设置中英文字体
            run.font.name = '楷体'
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 页码文本
        file_needs_page_numbers = False
        if i < len(file_page_number_list):
            file_needs_page_numbers = file_page_number_list[i]
        if file_needs_page_numbers:
            # 直接插入页码文本
            para = main_doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(f"（第{i+1}页，共{total_files}页）")
            run.font.size = Pt(12)
            # 设置中英文字体
            run.font.name = '楷体'
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            run.italic = True
        
        # 文件内容
        if file_ext in ['.docx', '.doc']:
            # 用docxcompose合并
            sub_doc = DocxDocument(file_path)
            
            # 为合并的文档设置中英文字体
            for paragraph in sub_doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '楷体'
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            
            composer.append(sub_doc)
            
            # 如果当前文件需要水印，在合并后添加水印
            if file_needs_watermark:
                add_watermark_to_existing_document(main_doc, watermark_text, watermark_font_size, watermark_angle, watermark_opacity, watermark_color, watermark_position)
                logger.info(f"为文件 {filename} 添加水印: {watermark_text}")
                
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
            # 只插入图片
            try:
                image_width = Inches(5.5)
                img_para = main_doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(file_path, width=image_width)
                
                # 如果当前文件需要水印，在图片后添加水印
                if file_needs_watermark:
                    add_watermark_to_existing_document(main_doc, watermark_text, watermark_font_size, watermark_angle, watermark_opacity, watermark_color, watermark_position)
                    logger.info(f"为图片文件 {filename} 添加水印: {watermark_text}")
                    
            except Exception as e:
                main_doc.add_paragraph(f"图片插入失败: {str(e)}")
                
        elif file_ext == '.pdf':
            # PDF转图片插入
            import fitz
            pdf_document = fitz.open(file_path)
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                temp_image_path = f"{temp_dir}/temp_pdf_{i}_{page_num}.png"
                pix.save(temp_image_path)
                img_para = main_doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(temp_image_path, width=Inches(5.5))
                os.remove(temp_image_path)
                
                # 如果当前文件需要水印，在每页PDF图片后添加水印
                if file_needs_watermark:
                    add_watermark_to_existing_document(main_doc, watermark_text, watermark_font_size, watermark_angle, watermark_opacity, watermark_color, watermark_position)
                    logger.info(f"为PDF文件 {filename} 第{page_num+1}页添加水印: {watermark_text}")
                    
            pdf_document.close()
            
        # 分页符
        if not is_last_file:
            main_doc.add_page_break()

    # 保存
    output_filename = f"{document_title}.docx"
    output_path = os.path.join("/app/generated_docs", output_filename)
    os.makedirs("/app/generated_docs", exist_ok=True)
    composer.save(output_path)
    logger.info(f"文档保存成功: {output_path}")
    shutil.rmtree(temp_dir)
    
    # 返回下载URL
    download_url = f"/api/download/{output_filename}"
    
    return {
        "success": True, 
        "message": "转换成功", 
        "output_file": output_filename,
        "download_url": download_url,
        "processed_files": file_names
    }

@app.get("/api/convert-history")
async def get_convert_history():
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                history = json.load(f)
        else:
            history = []
        return {"success": True, "history": history}
    except Exception as e:
        return {"success": False, "message": f"读取历史失败: {str(e)}"}

@app.delete("/api/convert-history")
async def clear_convert_history(db: Session = Depends(get_db)):
    """清空转换历史并将生成的文件加入清理列表"""
    try:
        # 将现有生成的文档文件加入到数据库管理
        generated_docs_dir = "/app/generated_docs"
        if os.path.exists(generated_docs_dir):
            # 获取所有.docx文件
            docx_files = [f for f in os.listdir(generated_docs_dir) if f.endswith('.docx')]
            
            for filename in docx_files:
                file_path = os.path.join(generated_docs_dir, filename)
                if os.path.exists(file_path):
                    try:
                        # 检查文件是否已在数据库中
                        import hashlib
                        with open(file_path, 'rb') as f:
                            file_hash = hashlib.md5(f.read()).hexdigest()
                        
                        existing_file = db.query(ManagedFile).filter(
                            ManagedFile.file_hash == file_hash
                        ).first()
                        
                        if not existing_file:
                            # 添加到数据库管理中，标记为生成的临时文件
                            import mimetypes
                            file_size = os.path.getsize(file_path)
                            mime_type, _ = mimetypes.guess_type(filename)
                            if not mime_type:
                                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            
                            generated_file = ManagedFile(
                                original_filename=filename,
                                display_name=filename,
                                storage_path=file_path,
                                file_type='document',
                                mime_type=mime_type,
                                file_size=file_size,
                                file_hash=file_hash,
                                file_category="temporary_generated",  # 生成的临时文件
                                category="generated_document",
                                description="系统生成的转换文档",
                                expires_at=datetime.now() + timedelta(days=180),  # 180天过期
                                access_count=0,
                                last_accessed=datetime.now()
                            )
                            db.add(generated_file)
                            logger.info(f"将生成文档加入管理: {filename}")
                        
                    except Exception as db_err:
                        logger.warning(f"处理生成文档失败 {filename}: {db_err}")
        
        # 提交数据库更改
        try:
            db.commit()
        except Exception as commit_err:
            logger.warning(f"提交数据库更改失败: {commit_err}")
            db.rollback()
        
        # 清空历史文件
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, "w", encoding="utf-8") as f:
                json.dump([], f, ensure_ascii=False, indent=2)
        
        return {"success": True, "message": "转换历史已清空，生成的文档已加入180天清理列表"}
    except Exception as e:
        logger.error(f"清空历史失败: {str(e)}")
        return {"success": False, "message": f"清空历史失败: {str(e)}"}

# ==================
# Project API
# ==================
@app.post("/api/projects", response_model=BaseResponse)
def create_project(project: ProjectCreate, db: Session = Depends(get_db)):
    try:
        db_project = Project(**project.model_dump())
        db.add(db_project)
        db.commit()
        db.refresh(db_project)
        return {"success": True, "message": "项目创建成功", "data": {"id": db_project.id}}
    except Exception as e:
        db.rollback()
        logger.error(f"创建项目失败: {e}")
        raise HTTPException(status_code=500, detail="创建项目失败")

@app.get("/api/projects", response_model=ProjectResponse)
def get_projects(db: Session = Depends(get_db)):
    projects = db.query(Project).order_by(Project.created_at.desc()).all()
    return {
        "success": True,
        "message": "获取项目列表成功",
        "projects": projects
    }

def create_tables():
    try:
        models.Base.metadata.create_all(bind=engine)
        db = SessionLocal()
        
        # 检查并初始化SystemSettings表
        if db.query(models.SystemSettings).count() == 0:
            default_settings = models.SystemSettings(
                company_name="默认公司名称",
                company_address="默认公司地址",
                contact_person="默认联系人",
                contact_phone="1234567890",
                contact_email="default@example.com"
            )
            db.add(default_settings)
            db.commit()
            logger.info("创建SystemSettings表成功，并插入了默认设置")

        # 检查并初始化基础数据表
        if db.query(models.Brand).count() == 0 and \
           db.query(models.BusinessField).count() == 0 and \
           db.query(models.Award).count() == 0 and \
           db.query(models.Performance).count() == 0:
            init_base_data(db)
            logger.info("初始化基础数据表成功")
            
        if db.query(models.Project).count() == 0:
            logger.info("检测到Project表为空，跳过初始化数据。")

    except Exception as e:
        logger.error(f"创建表或初始化数据失败: {e}")
    finally:
        db.close()

def add_watermark_to_pdf(pdf_path: str, text: str, font_size: int, angle: int, opacity: int, color: str, position: str) -> str:
    """在PDF中添加水印，返回处理后的PDF路径"""
    try:
        import fitz  # PyMuPDF
        import math
        
        # 角度取反，保证与前端预览一致
        angle = -angle
        
        # 打开PDF
        doc = fitz.open(pdf_path)
        
        # 解析颜色
        color_hex = color.lstrip('#')
        if len(color_hex) == 6:
            r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        else:
            r, g, b = 128, 128, 128
        
        # 颜色值转换为0-1范围
        color_r = r / 255.0
        color_g = g / 255.0  
        color_b = b / 255.0
        
        # 透明度（0-1范围）
        opacity_ratio = opacity / 100.0
        
        logger.info(f"水印颜色处理: 接收到颜色 '{color}' -> 解析hex '{color_hex}' -> 原始RGB({r},{g},{b}) -> 透明度{opacity}% -> 最终RGB({color_r:.2f},{color_g:.2f},{color_b:.2f})")
        
        # 遍历每一页
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # 获取页面尺寸
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            
            # 计算文本尺寸以正确居中
            try:
                # 获取文本宽度（PyMuPDF返回像素宽度）
                text_width = fitz.get_text_length(text, fontname="china-ss", fontsize=font_size)
                # 更精确的文本高度计算
                text_height = font_size * 1.2  # 考虑字体的上升部和下降部
            except:
                # 如果测量失败，使用估算值
                text_width = len(text) * font_size * 0.6  # 估算：每个字符约0.6倍字体大小
                text_height = font_size * 1.2
            
            # 计算真正的页面中心点
            page_center_x = page_width / 2
            page_center_y = page_height / 2
            
            # 根据位置计算水印位置（确保真正居中）
            if position == "center":
                # 直接使用页面中心作为文本的中心点
                # 这样无论字体大小如何变化，水印都会真正居中
                x = page_center_x  # 文本中心点的x坐标
                y = page_center_y  # 文本中心点的y坐标
            elif position == "top-left":
                x = 100 + text_width / 2  # 让文本中心距离边缘100像素
                y = 100 + text_height / 2
            elif position == "top-right":
                x = page_width - 100 - text_width / 2
                y = 100 + text_height / 2
            elif position == "bottom-left":
                x = 100 + text_width / 2
                y = page_height - 100 - text_height / 2
            elif position == "bottom-right":
                x = page_width - 100 - text_width / 2
                y = page_height - 100 - text_height / 2
            else:  # 默认居中
                x = page_center_x
                y = page_center_y
            
            logger.info(f"水印位置计算: 页面({page_width:.1f}x{page_height:.1f}), 文本尺寸({text_width:.1f}x{text_height:.1f}), 中心点({x:.1f},{y:.1f})")
            
            try:
                if angle != 0:
                    # 尝试使用PyMuPDF的旋转功能
                    angle_rad = math.radians(angle)
                    
                    try:
                        # 方法1: 使用文本旋转矩阵
                        # 创建旋转变换矩阵
                        cos_a = math.cos(angle_rad)
                        sin_a = math.sin(angle_rad)
                        rotation_matrix = fitz.Matrix(cos_a, sin_a, -sin_a, cos_a, 0, 0)
                        
                        # 使用预设的中心点作为旋转中心
                        # x, y 已经是文本中心点，直接使用
                        transform_point = fitz.Point(x, y)
                        
                        # 使用变换矩阵插入文本，使用RGB颜色+fill_opacity参数
                        page.insert_text(
                            transform_point,
                            text,
                            fontsize=font_size,
                            color=(color_r, color_g, color_b),  # RGB颜色
                            fill_opacity=opacity_ratio,  # 透明度参数
                            fontname="china-ss",  # 修复：使用思源宋体支持中文
                            morph=(transform_point, rotation_matrix)  # 使用morph参数进行旋转
                        )
                        logger.info(f"页面 {page_num + 1} 角度水印添加成功（角度: {angle}°, 颜色: {color}）")
                        
                    except Exception as rotate_error:
                        logger.warning(f"旋转矩阵失败，尝试字符分布: {rotate_error}")
                        
                        # 方法2: 字符分布模拟角度
                        cos_a = math.cos(angle_rad)
                        sin_a = math.sin(angle_rad)
                        
                        # 计算字符分布的精确间距
                        char_spacing = text_width / len(text) if len(text) > 0 else font_size * 0.6
                        
                        # 计算起始位置，使整个文本以(x,y)为中心
                        start_offset_x = -text_width / 2
                        start_offset_y = 0  # 垂直居中，基线偏移为0
                        
                        # 分别放置每个字符
                        for i, char in enumerate(text):
                            # 计算字符在旋转前的相对位置
                            char_offset_x = start_offset_x + (i + 0.5) * char_spacing
                            char_offset_y = start_offset_y
                            
                            # 应用旋转变换
                            rotated_x = char_offset_x * cos_a - char_offset_y * sin_a
                            rotated_y = char_offset_x * sin_a + char_offset_y * cos_a
                            
                            # 计算最终位置（以预设的中心点为旋转中心）
                            char_x = x + rotated_x
                            char_y = y + rotated_y
                            
                            try:
                                page.insert_text(
                                    fitz.Point(char_x, char_y),
                                    char,
                                    fontsize=font_size,
                                    color=(color_r, color_g, color_b),  # RGB颜色
                                    fill_opacity=opacity_ratio,  # 透明度参数
                                    fontname="china-ss"  # 修复：使用思源宋体支持中文
                                )
                            except Exception as char_error:
                                logger.warning(f"字符 '{char}' 添加失败: {char_error}")
                                continue
                        
                        logger.info(f"页面 {page_num + 1} 使用字符分布模拟角度（角度: {angle}°, 颜色: {color}）")
                        
                else:
                    # 无旋转，需要计算文本起始位置（左下角）
                    # x, y是文本中心点，需要转换为文本起始点
                    text_start_x = x - text_width / 2
                    text_start_y = y + font_size / 3  # 调整基线位置
                    
                    # 直接添加文本，使用RGB颜色+fill_opacity参数
                    page.insert_text(
                        fitz.Point(text_start_x, text_start_y),
                        text,
                        fontsize=font_size,
                        color=(color_r, color_g, color_b),  # RGB颜色
                        fill_opacity=opacity_ratio,  # 透明度参数
                        fontname="china-ss"  # 修复：使用思源宋体支持中文
                    )
                    logger.info(f"页面 {page_num + 1} 水印添加成功（无旋转, 颜色: {color}）")
                    
            except Exception as text_error:
                logger.error(f"页面 {page_num + 1} 水印添加失败: {text_error}")
                # 降级处理：使用最简单的方法
                try:
                    # 降级处理：使用简单的居中计算
                    fallback_x = x - text_width / 2
                    fallback_y = y + font_size / 3
                    
                    # 降级处理也使用RGB颜色+fill_opacity参数
                    page.insert_text(
                        fitz.Point(fallback_x, fallback_y),
                        text,
                        fontsize=font_size,
                        color=(color_r, color_g, color_b),  # RGB颜色
                        fill_opacity=opacity_ratio,  # 透明度参数
                        fontname="china-ss"  # 修复：使用思源宋体支持中文
                    )
                    logger.info(f"页面 {page_num + 1} 使用降级方法添加水印")
                except Exception as fallback_error:
                    logger.error(f"页面 {page_num + 1} 所有水印方法都失败: {fallback_error}")
        
        # 保存处理后的PDF
        watermarked_pdf_path = pdf_path.replace('.pdf', '_watermarked.pdf')
        doc.save(watermarked_pdf_path)
        doc.close()
        
        logger.info(f"PDF水印添加成功: {watermarked_pdf_path}")
        return watermarked_pdf_path
        
    except Exception as e:
        logger.error(f"PDF水印添加失败: {e}")
        return pdf_path  # 如果失败，返回原PDF路径

def add_watermark_to_existing_document(doc, text, font_size, angle, opacity, color, position):
    """为Word文档添加水印"""
    if not text or not text.strip():
        return
    
    try:
        # 转换颜色
        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip('#')
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        
        rgb_color = hex_to_rgb(color)
        # 应用透明度
        opacity = max(0.0, min(1.0, opacity / 100.0))
        rgb_color = tuple(int(c * opacity + 255 * (1 - opacity)) for c in rgb_color)
        
        # 根据位置添加水印
        if position == "center":
            _add_center_watermark(doc, text, font_size, rgb_color, angle)
        elif position == "repeat":
            _add_repeat_watermark(doc, text, font_size, rgb_color, angle)
        elif position == "background":
            _add_background_watermark(doc, text, font_size, rgb_color, angle)
        elif position in ["top-left", "top-right", "bottom-left", "bottom-right"]:
            _add_corner_watermark(doc, text, font_size, rgb_color, angle, position)
        
        logger.info(f"水印添加成功: {text}")
        
    except Exception as e:
        logger.error(f"添加水印失败: {e}")

def _add_center_watermark(doc, text, font_size, rgb_color, angle):
    """添加居中大水印"""
    # 在文档中央位置插入水印
    for i in range(8):
        doc.add_paragraph()
    
    # 根据角度应用装饰
    if angle == 0:
        decorated_text = text
    elif angle > 0:
        decorated_text = f"╱ {text} ╱"
    else:
        decorated_text = f"╲ {text} ╲"
    
    # 上装饰边框
    border_para1 = doc.add_paragraph()
    border_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border1_run = border_para1.add_run("◆" * 20)
    border1_run.font.size = Pt(max(8, font_size // 3))
    border1_run.font.color.rgb = RGBColor(*rgb_color)
    border1_run.font.name = '楷体'
    
    # 主水印段落
    watermark_para = doc.add_paragraph()
    watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    main_run = watermark_para.add_run(f"『 {decorated_text} 』")
    main_run.font.size = Pt(font_size * 2)
    main_run.font.color.rgb = RGBColor(*rgb_color)
    main_run.font.bold = True
    main_run.font.name = '楷体'
    
    # 下装饰边框
    border_para2 = doc.add_paragraph()
    border_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border2_run = border_para2.add_run("◇" * 20)
    border2_run.font.size = Pt(max(8, font_size // 3))
    border2_run.font.color.rgb = RGBColor(*rgb_color)
    border2_run.font.name = '楷体'

def _add_repeat_watermark(doc, text, font_size, rgb_color, angle):
    """添加重复平铺水印"""
    current_length = len(doc.paragraphs)
    watermark_positions = []
    for i in range(5, max(50, current_length + 30), 12):
        watermark_positions.append(i)
    
    if angle == 0:
        decorated_text = text
    elif angle > 0:
        decorated_text = f"╱ {text} ╱"
    else:
        decorated_text = f"╲ {text} ╲"
    
    for pos in watermark_positions:
        while len(doc.paragraphs) < pos:
            doc.add_paragraph()
        
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        repeat_text = f"◆ {decorated_text} ◆   ◇ {decorated_text} ◇   ◆ {decorated_text} ◆"
        watermark_run = watermark_para.add_run(repeat_text)
        watermark_run.font.size = Pt(font_size)
        watermark_run.font.color.rgb = RGBColor(*rgb_color)
        watermark_run.font.bold = True
        watermark_run.italic = True
        watermark_run.font.name = '楷体'
        
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        deco_run = deco_para.add_run("~ " * 15)
        deco_run.font.size = Pt(max(8, font_size // 2))
        deco_run.font.color.rgb = RGBColor(*rgb_color)
        deco_run.font.name = '楷体'

def _add_background_watermark(doc, text, font_size, rgb_color, angle):
    """添加背景样式水印"""
    if angle == 0:
        decorated_text = text
    elif angle > 0:
        decorated_text = f"╱ {text} ╱"
    else:
        decorated_text = f"╲ {text} ╲"
    
    for row in range(12):
        bg_para = doc.add_paragraph()
        bg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if row == 5:  # 中心主水印
            bg_run = bg_para.add_run(f"【 {decorated_text} 】")
            bg_run.font.size = Pt(font_size * 2)
            bg_run.font.bold = True
        elif row == 2 or row == 8:  # 副水印行
            bg_run = bg_para.add_run(f"◆ {decorated_text} ◆")
            bg_run.font.size = Pt(font_size)
            bg_run.font.bold = True
        elif row % 2 == 0:  # 装饰行
            bg_run = bg_para.add_run("～ ～ ～ ～ ～ ～ ～ ～ ～ ～")
            bg_run.font.size = Pt(max(8, font_size // 2))
        else:  # 空行，只有少量装饰
            bg_run = bg_para.add_run("· · · · ·")
            bg_run.font.size = Pt(max(6, font_size // 3))
        
        bg_run.font.color.rgb = RGBColor(*rgb_color)
        bg_run.font.name = '楷体'

def _add_corner_watermark(doc, text, font_size, rgb_color, angle, position):
    """添加角落位置水印"""
    if angle == 0:
        decorated_text = text
    elif angle > 0:
        decorated_text = f"╱ {text} ╱"
    else:
        decorated_text = f"╲ {text} ╲"
    
    if position == "top-left":
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        prefix_lines = 2
    elif position == "top-right":
        alignment = WD_ALIGN_PARAGRAPH.RIGHT
        prefix_lines = 2
    elif position == "bottom-left":
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        prefix_lines = 15
    elif position == "bottom-right":
        alignment = WD_ALIGN_PARAGRAPH.RIGHT
        prefix_lines = 15
    else:
        alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix_lines = 8
    
    for i in range(prefix_lines):
        doc.add_paragraph()
    
    watermark_para = doc.add_paragraph()
    watermark_para.alignment = alignment
    
    watermark_run = watermark_para.add_run(f"◆ {decorated_text} ◆")
    watermark_run.font.size = Pt(font_size)
    watermark_run.font.color.rgb = RGBColor(*rgb_color)
    watermark_run.font.bold = True
    watermark_run.font.name = '楷体'
    
    deco_para = doc.add_paragraph()
    deco_para.alignment = alignment
    deco_run = deco_para.add_run("━" * 10)
    deco_run.font.size = Pt(max(8, font_size // 3))
    deco_run.font.color.rgb = RGBColor(*rgb_color)
    deco_run.font.name = '楷体'

# 添加分类提示词设置管理API
@app.get("/api/settings/classification-prompts")
async def get_classification_prompts(db: Session = Depends(get_db)):
    """获取分类提示词设置"""
    try:
        prompts = db.query(SystemSettings).filter(
            SystemSettings.category == "ai_classification"
        ).all()
        
        result = {}
        for prompt in prompts:
            result[prompt.setting_key] = {
                "value": prompt.setting_value,
                "description": prompt.description,
                "is_editable": prompt.is_editable,
                "requires_restart": prompt.requires_restart,
                "updated_at": prompt.updated_at.isoformat() if prompt.updated_at else None
            }
        
        return {"success": True, "prompts": result}
    except Exception as e:
        logger.error(f"获取分类提示词失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/settings/classification-prompts/{prompt_key}")
async def update_classification_prompt(
    prompt_key: str,
    prompt_data: dict,
    db: Session = Depends(get_db)
):
    """更新分类提示词"""
    try:
        # 验证prompt_key
        allowed_keys = ["classification_vision_prompt", "classification_text_prompt"]
        if prompt_key not in allowed_keys:
            raise HTTPException(status_code=400, detail=f"不支持的提示词类型: {prompt_key}")
        
        # 获取设置
        setting = db.query(SystemSettings).filter(
            SystemSettings.setting_key == prompt_key
        ).first()
        
        if not setting:
            raise HTTPException(status_code=404, detail="提示词设置不存在")
        
        if not setting.is_editable:
            raise HTTPException(status_code=400, detail="该设置不可编辑")
        
        # 更新值
        new_value = prompt_data.get("value", "")
        if not new_value.strip():
            raise HTTPException(status_code=400, detail="提示词内容不能为空")
        
        setting.setting_value = new_value
        setting.updated_at = func.now()
        
        db.commit()
        
        # 热重载AI服务配置
        ai_service.reload_config()
        
        return {
            "success": True,
            "message": f"提示词 {prompt_key} 更新成功",
            "updated_at": setting.updated_at.isoformat(),
            "reloaded": True
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"更新分类提示词失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/settings/classification-prompts/reset/{prompt_key}")
async def reset_classification_prompt(
    prompt_key: str,
    db: Session = Depends(get_db)
):
    """重置分类提示词为默认值"""
    try:
        # 验证prompt_key
        allowed_keys = ["classification_vision_prompt", "classification_text_prompt"]
        if prompt_key not in allowed_keys:
            raise HTTPException(status_code=400, detail=f"不支持的提示词类型: {prompt_key}")
        
        # 获取设置
        setting = db.query(SystemSettings).filter(
            SystemSettings.setting_key == prompt_key
        ).first()
        
        if not setting:
            raise HTTPException(status_code=404, detail="提示词设置不存在")
        
        # 获取默认提示词
        if prompt_key == "classification_vision_prompt":
            default_prompt = ai_service._get_default_vision_prompt()
        else:
            default_prompt = ai_service._get_default_text_prompt()
        
        # 更新为默认值
        setting.setting_value = default_prompt
        setting.updated_at = func.now()
        
        db.commit()
        
        # 热重载AI服务配置
        ai_service.reload_config()
        
        return {
            "success": True,
            "message": f"提示词 {prompt_key} 重置为默认值",
            "updated_at": setting.updated_at.isoformat(),
            "reloaded": True
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"重置分类提示词失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# 检查Docling是否可用
try:
    import docling
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

# Docling OCR管理API
@app.get("/api/ocr/docling/status")
async def get_docling_ocr_status():
    """获取Docling OCR状态"""
    try:
        # 使用DoclingService获取状态
        from docling_service import docling_service as ds
        status = ds.get_status()
        
        return {
            "success": True,
            "docling_available": status.get("docling_available", False),
            "ocr_enabled": status.get("config", {}).get("enable_ocr", False),
            "converter_initialized": status.get("initialized", False),
            "languages": status.get("config", {}).get("ocr_languages", [])
        }
    except Exception as e:
        logger.error(f"获取Docling OCR状态失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/ocr/docling/reload")
async def reload_docling_ocr():
    """重新加载Docling OCR配置"""
    try:
        ai_service.reload_config()
        
        return {
            "success": True,
            "message": "Docling OCR配置重新加载成功",
            "ocr_enabled": ai_service.enable_docling_ocr,
            "converter_initialized": ai_service.docling_converter is not None
        }
    except Exception as e:
        logger.error(f"重新加载Docling OCR配置失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/ocr/easyocr/download-models")
async def download_easyocr_models():
    """下载EasyOCR模型"""
    try:
        # 调用AI服务的模型下载方法
        result = await ai_service.download_easyocr_models()
        
        if result.get("success"):
            return {
                "success": True,
                "message": result.get("message", "EasyOCR模型下载成功"),
                "model_path": result.get("model_path")
            }
        else:
            return {
                "success": False,
                "message": result.get("error", "模型下载失败")
            }
    except Exception as e:
        logger.error(f"下载EasyOCR模型失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/ocr/easyocr/status")
async def get_easyocr_status():
    """获取EasyOCR模型状态"""
    try:
        import os
        # 使用DoclingService获取模型路径
        from docling_service import docling_service as ds
        status = ds.get_status()
        model_path = status.get("config", {}).get("easyocr_models_path", "")
        
        # 检查模型文件是否存在
        model_files = ['craft_mlt_25k.pth', 'zh_sim_g2.pth', 'english_g2.pth']
        existing_models = []
        
        if model_path and os.path.exists(model_path):
            for model_file in model_files:
                file_path = os.path.join(model_path, model_file)
                if os.path.exists(file_path):
                    existing_models.append(model_file)
        
        is_initialized = len(existing_models) >= 2  # 至少需要2个主要模型文件
        
        return {
            "success": True,
            "initialized": is_initialized,
            "model_path": model_path,
            "existing_models": existing_models,
            "total_models": len(model_files),
            "languages": status.get("config", {}).get("ocr_languages", [])
        }
    except Exception as e:
        logger.error(f"获取EasyOCR状态失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/ai-models/status")
async def get_ai_models_status():
    """获取AI模型状态"""
    try:
        from ai_service import ai_service
        status = ai_service.get_ai_models_status()
        return {"success": True, **status}
    except Exception as e:
        logger.error(f"获取AI模型状态失败: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/ai-models/download")
async def trigger_ai_models_download():
    """触发AI模型下载"""
    try:
        from ai_service import ai_service
        ai_service._start_ai_models_download()
        return {"success": True, "message": "AI模型下载已启动"}
    except Exception as e:
        logger.error(f"启动AI模型下载失败: {e}")
        return {"success": False, "error": str(e)}



if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)