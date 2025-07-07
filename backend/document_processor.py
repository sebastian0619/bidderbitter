import os
import logging
import time
from datetime import datetime
import uuid
import fitz  # PyMuPDF

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    logger.warning("python-magic未安装，将使用文件扩展名进行类型检测")
    MAGIC_AVAILABLE = False
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from PIL import Image, ImageDraw, ImageFont
import io
import asyncio
from typing import List, Dict, Any, Optional, Tuple
import shutil
import tempfile
from pathlib import Path
import math
# 导入统一的DoclingService
try:
    from docling_service import docling_service
    DOCLING_SERVICE_AVAILABLE = True
except ImportError as e:
    DOCLING_SERVICE_AVAILABLE = False
    docling_service = None

# 输出目录配置
UPLOAD_DIR = os.environ.get("UPLOAD_PATH", "./uploads")
CONVERTED_DIR = os.path.join(UPLOAD_DIR, "converted")
GENERATED_DIR = os.environ.get("GENERATED_DOCS_PATH", "./generated_docs")

# 确保目录存在
try:
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(CONVERTED_DIR, exist_ok=True)
    os.makedirs(GENERATED_DIR, exist_ok=True)
except OSError as e:
    logger.warning(f"无法创建目录: {e}")

def clean_heading_style(heading):
    """清除标题的项目符号和列表样式，防止出现小黑点"""
    try:
        from docx.oxml.shared import qn
        
        # 清除段落的编号和项目符号
        pPr = heading._element.get_or_add_pPr()
        
        # 移除编号属性
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
        
        # 设置段落格式，明确禁用列表样式和分页控制
        if hasattr(heading, 'paragraph_format'):
            heading.paragraph_format.left_indent = None
            heading.paragraph_format.first_line_indent = None
            
            # 禁用分页控制选项，对应Word中的"分页"设置
            heading.paragraph_format.widow_control = False      # 孤行控制
            heading.paragraph_format.keep_with_next = False     # 与下段同页
            heading.paragraph_format.keep_together = False      # 段中不分页
            heading.paragraph_format.page_break_before = False  # 段前分页
            
    except Exception as style_error:
        logger.warning(f"清除标题样式时出错: {style_error}")

def create_clean_heading(doc, text, level=1, center=False):
    """创建一个没有小黑点的干净标题"""
    heading = doc.add_heading(text, level)
    
    # 清除项目符号样式
    clean_heading_style(heading)
    
    # 设置对齐方式
    if center or level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return heading

class DoclingDocumentProcessor:
    """使用Docling进行文档处理的类"""
    
    def __init__(self):
        self.converter = self._create_converter()
    
    def _format_heading(self, doc, text, level=1, center=False):
        """
        创建格式化的标题
        
        参数:
        - doc: Document对象
        - text: 标题文字
        - level: 标题级别 (0=主标题, 1=一级标题, 2=二级标题)
        - center: 是否居中
        
        返回:
        - 格式化的标题段落
        """
        from docx.oxml.shared import qn
        from docx.oxml.ns import nsdecls, qn
        from docx.oxml import parse_xml
        
        # 创建标题
        heading = doc.add_heading(text, level)
        
        # 设置对齐方式
        if center or level == 0:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 清除段落的列表样式和项目符号，防止出现小黑点
        try:
            # 清除段落的编号和项目符号
            pPr = heading._element.get_or_add_pPr()
            
            # 移除编号属性
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            
            # 设置段落格式，明确禁用列表样式和分页控制
            if hasattr(heading, 'paragraph_format'):
                heading.paragraph_format.left_indent = None
                heading.paragraph_format.first_line_indent = None
                
                # 禁用分页控制选项，对应Word中的"分页"设置
                heading.paragraph_format.widow_control = False      # 孤行控制
                heading.paragraph_format.keep_with_next = False     # 与下段同页
                heading.paragraph_format.keep_together = False      # 段中不分页
                heading.paragraph_format.page_break_before = False  # 段前分页
                
        except Exception as style_error:
            logger.warning(f"清除标题样式时出错: {style_error}")
        
        # 设置字体
        if heading.runs:
            run = heading.runs[0]
            
            # 根据级别设置字体大小
            if level == 0:
                run.font.size = Pt(18)
                run.font.name = '楷体'
            elif level == 1:
                run.font.size = Pt(16)
                run.font.name = '楷体'
            elif level == 2:
                run.font.size = Pt(14)
                run.font.name = '楷体'
            else:
                run.font.size = Pt(12)
                run.font.name = '楷体'
            
            run.bold = True
    
            # 确保字体颜色是黑色
            run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
            
            # 设置中英文字体
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        
        return heading
    
    def _calculate_image_size_for_page(self, image_path: str, page_num: int, has_file_title: bool, max_height_inches: float = 8.0) -> Tuple[object, bool]:
        """
        智能计算图片在页面中的合适大小，检测是否需要分页
        """
        from docx.shared import Inches
        from PIL import Image
        try:
            # 根据页码设置不同的最大宽度
            if page_num == 0:
                # 第一页：设置较小宽度，为标题留空间
                max_width = 5.5  # 第一页统一使用5.5英寸
            else:
                # 非第一页：使用15cm = 5.91英寸（按用户要求）
                max_width = 5.91  # 15cm = 5.91英寸
                
            available_height = max_height_inches
            # 根据页面类型调整预留空间，非第一页极其紧凑
            if page_num == 0 and has_file_title:
                available_height -= 0.5  # 主标题+文件标题合计只预留0.5英寸
            elif page_num == 0:
                available_height -= 0.3  # 只有主标题，预留更少
            else:
                # 非第一页：几乎不预留空间，让图片占满页面
                available_height -= 0.1  # 非第一页只预留极少空间
            # 页码预留空间更少
            available_height -= 0.1
            with Image.open(image_path) as img:
                img_width, img_height = img.size
                aspect_ratio = img_height / img_width
                # 优先按最大宽度自适应，更大胆的尺寸策略
                scaled_height = max_width * aspect_ratio
                if scaled_height > available_height * 1.3:  # 只有更极端的长图才按高度缩放
                    target_width = available_height / aspect_ratio
                    target_width = min(target_width, max_width)
                    needs_page_break = scaled_height > available_height * 1.8
                else:
                    target_width = max_width  # 使用当前页码对应的最大宽度
                    needs_page_break = False
                # 确保宽度在合理范围内，调整最小宽度
                target_width = max(4.0, min(target_width, max_width))
                logger.info(f"图片尺寸分析(修正) - 页码: {page_num}, 原始: {img_width}x{img_height}, 最大宽度: {max_width:.2f}, 目标宽度: {target_width:.2f}英寸, 需要分页: {needs_page_break}")
                return Inches(target_width), needs_page_break
        except Exception as e:
            logger.warning(f"图片尺寸分析失败: {e}，使用默认尺寸")
            return self._calculate_image_width(page_num, has_file_title), False
    
    def _calculate_image_width(self, page_num: int, has_file_title: bool) -> object:
        """
        根据页码和是否有文件标题计算合适的图片宽度（保留原方法作为备选）
        """
        from docx.shared import Inches
        
        # 根据页面设置不同宽度
        if page_num == 0:
            # 第一页：统一使用5.5英寸
            return Inches(5.5)
        else:
            # 非第一页：使用15cm = 5.91英寸
            return Inches(5.91)
    
    def _create_converter(self):
        """使用统一的DoclingService"""
        if DOCLING_SERVICE_AVAILABLE and docling_service:
            logger.info("使用统一的DoclingService进行文档处理")
            return docling_service
        else:
            logger.warning("DoclingService不可用，将使用PyMuPDF作为备选方案")
            return None
    
    def _add_page_content_with_smart_sizing(self, doc: Document, image_path: str, page_num: int, 
                                          show_file_titles: bool, watermark_config: Dict[str, Any] = None,
                                          is_last_page: bool = False, total_pages: int = None):
        """
        智能添加页面内容，包括图片尺寸优化和分页控制
        """
        from PIL import Image
        from docx.shared import Inches
        try:
            # 智能计算图片大小（确保传递正确的has_file_title参数）
            image_width, needs_manual_break = self._calculate_image_size_for_page(
                image_path, page_num, True  # 明确传递True，确保第一页使用5.5英寸
            )
            # 判断极端长图（高宽比>2.8）
            with Image.open(image_path) as img:
                img_width, img_height = img.size
                aspect_ratio = img_height / img_width
                if aspect_ratio > 2.8:
                    # 自动分页裁切 - 根据页码使用不同宽度
                    if page_num == 0:
                        max_width = 5.5  # 第一页
                    else:
                        max_width = 5.91  # 非第一页：15cm
                    # 计算可用高度（与自适应算法一致）
                    available_height = 8.0  # 合理的基础高度
                    if page_num == 0 and show_file_titles:
                        available_height -= 1.0
                    elif page_num == 0:
                        available_height -= 0.7
                    else:
                        # 非第一页：几乎不减少高度，让图片占满页面
                        available_height -= 0.2
                    available_height -= 0.1
                    # 转换为像素高度（假设96dpi）
                    dpi = img.info.get('dpi', (96, 96))[1]
                    if not dpi or dpi < 10:
                        dpi = 96
                    max_page_height_px = int(available_height * dpi)
                    # 分割图片
                    num_slices = (img_height + max_page_height_px - 1) // max_page_height_px
                    for i in range(num_slices):
                        upper = i * max_page_height_px
                        lower = min((i + 1) * max_page_height_px, img_height)
                        box = (0, upper, img_width, lower)
                        slice_img = img.crop(box)
                        temp_path = f"temp_longimg_{page_num+1}_{i+1}.png"
                        slice_img.save(temp_path)
                        # 添加页码标题
                        para = doc.add_paragraph()
                        if total_pages:
                            run = para.add_run(f"第 {page_num + 1} 页(共 {total_pages} 页) - 第 {i + 1} 段")
                        else:
                            run = para.add_run(f"第 {page_num + 1}-{i + 1} 段")
                        run.font.size = Pt(10)
                        run.font.bold = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 插入图片（居中对齐）
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.add_run().add_picture(temp_path, width=Inches(max_width))
                        # 除最后一段外都分页
                        if not (is_last_page and i == num_slices - 1):
                            doc.add_page_break()
                        # 清理临时文件
                        import os
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    return
            # 普通图片按原逻辑
            para = doc.add_paragraph()
            if total_pages:
                run = para.add_run(f"第 {page_num + 1} 页(共 {total_pages} 页)")
            else:
                run = para.add_run(f"第 {page_num + 1} 页")
            run.font.size = Pt(10)  # 小字体
            run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                # 插入图片（居中对齐）
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(image_path, width=image_width)
                logger.info(f"成功插入图片，页码: {page_num + 1}, 宽度: {image_width}")
            except Exception as img_error:
                logger.error(f"图片插入失败: {img_error}")
                doc.add_paragraph(f"图片插入失败: {str(img_error)}")
            if not is_last_page:
                doc.add_page_break()
                logger.debug(f"添加分页符，页码: {page_num + 1}")
            else:
                logger.debug(f"跳过分页符（最后一页），页码: {page_num + 1}")
        except Exception as e:
            logger.error(f"智能页面内容添加失败: {e}")
            para = doc.add_paragraph()
            if total_pages:
                run = para.add_run(f"第 {page_num + 1} 页(共 {total_pages} 页)")
            else:
                run = para.add_run(f"第 {page_num + 1} 页")
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                image_width = self._calculate_image_width(page_num, show_file_titles)
                # 插入图片（居中对齐）
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(image_path, width=image_width)
                if not is_last_page:
                    doc.add_page_break()
            except Exception as fallback_error:
                doc.add_paragraph(f"图片处理失败: {str(fallback_error)}")
    
    async def process_pdf_with_docling(self, pdf_path: str, doc: Document, filename: str, watermark_config: Dict[str, Any] = None, show_file_titles: bool = True, file_title_level: int = 2, is_last_file: bool = False):
        """使用DoclingService处理PDF文件 - 充分利用Docling的文档理解能力"""
        try:
            if not self.converter:
                logger.warning("DoclingService不可用，使用PyMuPDF备选方案")
                return await self._process_pdf_fallback(pdf_path, doc, filename, watermark_config, show_file_titles, file_title_level, is_last_file)
            
            logger.info(f"开始使用DoclingService处理PDF: {filename}")
            
            # 首先检测PDF类型（扫描件 vs 非扫描件）
            pdf_type = await self._detect_pdf_type(pdf_path)
            logger.info(f"PDF类型检测结果: {pdf_type}")
            
            # 根据show_file_titles参数决定是否添加文件标题（直接显示文件名，不加前缀，去掉扩展名）
            if show_file_titles:
                filename_without_ext = os.path.splitext(filename)[0]
                self._format_heading(doc, filename_without_ext, level=file_title_level, center=False)
            
            if pdf_type == "scanned":
                # 扫描件PDF：使用Docling进行OCR和文本提取
                return await self._process_scanned_pdf(pdf_path, doc, filename, watermark_config, show_file_titles, file_title_level, is_last_file)
            else:
                # 非扫描件PDF：直接转换为图片，保持原始格式
                return await self._process_native_pdf(pdf_path, doc, filename, watermark_config, show_file_titles, file_title_level, is_last_file)
            
        except Exception as e:
            logger.error(f"DoclingService处理PDF失败: {e}")
            # 降级到PyMuPDF
            return await self._process_pdf_fallback(pdf_path, doc, filename, watermark_config, show_file_titles, file_title_level, is_last_file)
    
    async def _detect_pdf_type(self, pdf_path: str) -> str:
        """检测PDF类型：扫描件 vs 非扫描件"""
        try:
            import fitz
            
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            
            # 检查前3页（或全部页面，如果少于3页）
            pages_to_check = min(3, total_pages)
            text_ratio = 0
            image_ratio = 0
            
            for page_num in range(pages_to_check):
                page = pdf_document.load_page(page_num)
                
                # 获取文本
                text = page.get_text()
                text_length = len(text.strip())
                
                # 获取图片
                image_list = page.get_images()
                image_count = len(image_list)
                
                # 计算比例
                if text_length > 0:
                    text_ratio += 1
                if image_count > 0:
                    image_ratio += 1
            
            pdf_document.close()
            
            # 判断PDF类型
            text_percentage = text_ratio / pages_to_check
            image_percentage = image_ratio / pages_to_check
            
            logger.info(f"PDF类型分析 - 文本页面比例: {text_percentage:.2f}, 图片页面比例: {image_percentage:.2f}")
            
            # 如果大部分页面都有可提取文本，认为是非扫描件
            if text_percentage > 0.7:
                return "native"
            # 如果大部分页面都有图片且文本很少，认为是扫描件
            elif image_percentage > 0.7 and text_percentage < 0.3:
                return "scanned"
            else:
                # 混合类型，优先按非扫描件处理
                return "mixed"
                
        except Exception as e:
            logger.warning(f"PDF类型检测失败: {e}，默认按非扫描件处理")
            return "native"
    
    async def _process_scanned_pdf(self, pdf_path: str, doc: Document, filename: str, watermark_config: Dict[str, Any] = None, show_file_titles: bool = True, file_title_level: int = 2, is_last_file: bool = False):
        """处理扫描件PDF：使用Docling进行OCR和文本提取"""
        try:
            logger.info(f"处理扫描件PDF: {filename}")
            
            # 使用DoclingService转换PDF
            conv_result = await self.converter.convert_document(pdf_path)
            
            if not conv_result.get("success"):
                logger.warning(f"DoclingService转换失败: {conv_result.get('error')}")
                return await self._process_pdf_fallback(pdf_path, doc, filename, watermark_config, show_file_titles, file_title_level, is_last_file)
            
            # 添加文档内容
            text_content = conv_result.get("text", "")
            if text_content.strip():
                content_heading = create_clean_heading(doc, "OCR识别文本", level=2)
                # 将长文本分段
                paragraphs = text_content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            
            # 添加文档元数据
            metadata = conv_result.get("metadata", {})
            if metadata:
                stats_heading = create_clean_heading(doc, "文档信息", level=3)
                stats_para = doc.add_paragraph()
                if "pages" in metadata:
                    stats_para.add_run(f"总页数: {metadata['pages']} 页\n")
                if "file_type" in metadata:
                    stats_para.add_run(f"文件类型: {metadata['file_type']}\n")
                stats_para.add_run("处理方式: OCR文本识别\n")
            
            # 添加页面截图作为补充
            await self._add_page_screenshots(doc, pdf_path)
            
            # 在PDF处理完成后添加分页符（除非是最后一个文件）
            if not is_last_file:
                doc.add_page_break()
                logger.info(f"扫描件PDF处理完成，已添加分页符: {filename}")
            else:
                logger.info(f"扫描件PDF处理完成，跳过分页符（最后一个文件）: {filename}")
            
            return {
                "success": True,
                "message": "扫描件PDF处理成功（OCR文本识别）",
                "text_extracted": True,
                "text_length": len(text_content),
                "metadata": metadata,
                "pdf_type": "scanned"
            }
            
        except Exception as e:
            logger.error(f"扫描件PDF处理失败: {e}")
            return await self._process_pdf_fallback(pdf_path, doc, filename, watermark_config, show_file_titles, file_title_level, is_last_file)
    
    async def _process_native_pdf(self, pdf_path: str, doc: Document, filename: str, watermark_config: Dict[str, Any] = None, show_file_titles: bool = True, file_title_level: int = 2, is_last_file: bool = False):
        """处理非扫描件PDF：直接转换为图片，保持原始格式"""
        try:
            logger.info(f"处理非扫描件PDF: {filename}")
            
            # 添加处理说明
            info_heading = create_clean_heading(doc, "文档说明", level=2)
            info_para = doc.add_paragraph()
            info_para.add_run("此PDF文档包含可编辑文本，为保持原始格式，已转换为图片形式展示。")
            info_para.add_run("\n如需编辑文本内容，请使用原始PDF文件。")
            
            # 直接转换为图片并插入
            await self._add_page_screenshots_enhanced(doc, pdf_path, is_last_file)
            
            # 在PDF处理完成后添加分页符（除非是最后一个文件）
            if not is_last_file:
                doc.add_page_break()
                logger.info(f"非扫描件PDF处理完成，已添加分页符: {filename}")
            else:
                logger.info(f"非扫描件PDF处理完成，跳过分页符（最后一个文件）: {filename}")
            
            return {
                "success": True,
                "message": "非扫描件PDF处理成功（图片格式保持）",
                "text_extracted": False,
                "pdf_type": "native"
            }
            
        except Exception as e:
            logger.error(f"非扫描件PDF处理失败: {e}")
            return await self._process_pdf_fallback(pdf_path, doc, filename, watermark_config, show_file_titles, file_title_level, is_last_file)
    
    async def _add_page_screenshots_enhanced(self, doc: Document, pdf_path: str, is_last_file: bool = False):
        """增强的PDF页面截图功能，专门用于非扫描件PDF"""
        try:
            import fitz
            
            create_clean_heading(doc, "PDF页面内容", level=2)
            
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            
            for page_num in range(total_pages):
                page = pdf_document.load_page(page_num)
                
                # 使用更高分辨率确保文本清晰
                mat = fitz.Matrix(3.0, 3.0)  # 提高分辨率到3倍
                pix = page.get_pixmap(matrix=mat)
                
                temp_image_path = f"temp_native_page_{page_num + 1}_{uuid.uuid4().hex[:8]}.png"
                pix.save(temp_image_path)
                
                try:
                    # 添加页码标题
                    page_heading = create_clean_heading(doc, f"第 {page_num + 1} 页(共 {total_pages} 页)", level=4)
                    
                    # 智能计算图片大小 - 非扫描件使用更大的尺寸
                    image_width, _ = self._calculate_image_size_for_page(
                        temp_image_path, page_num, True, max_height_inches=9.0  # 增加最大高度
                    )
                    
                    # 添加图片（居中对齐）
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_para.add_run().add_picture(temp_image_path, width=image_width)
                    
                    # 只在不是最后一页时添加分页符
                    if page_num < total_pages - 1:
                        doc.add_page_break()
                    
                except Exception as img_error:
                    logger.warning(f"添加第{page_num + 1}页截图失败: {img_error}")
                    doc.add_paragraph(f"第{page_num + 1}页图像处理失败")
                
                # 清理临时文件
                if os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
            
            pdf_document.close()
                
        except Exception as e:
            logger.error(f"增强页面截图失败: {e}")
            doc.add_paragraph(f"PDF页面处理失败: {str(e)}")
    
    async def _add_page_screenshots(self, doc: Document, pdf_path: str):
        """添加PDF页面截图作为补充（用于扫描件PDF）"""
        try:
            import fitz
            
            create_clean_heading(doc, "PDF页面图像", level=2)
            
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            
            for page_num in range(min(total_pages, 10)):  # 最多处理10页，避免文档过大
                page = pdf_document.load_page(page_num)
                mat = fitz.Matrix(1.5, 1.5)  # 适中的分辨率
                pix = page.get_pixmap(matrix=mat)
                
                temp_image_path = f"temp_docling_page_{page_num + 1}_{uuid.uuid4().hex[:8]}.png"
                pix.save(temp_image_path)
                
                try:
                    # 添加页码标题
                    page_heading = create_clean_heading(doc, f"第 {page_num + 1} 页(共 {total_pages} 页)", level=4)
                    
                    # 智能计算图片大小 - 传递正确的has_file_title参数
                    image_width, _ = self._calculate_image_size_for_page(
                        temp_image_path, page_num, True, max_height_inches=6.0
                    )
                    
                    # 添加图片（居中对齐）
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_para.add_run().add_picture(temp_image_path, width=image_width)
                    
                    # 只在不是最后一页时添加分页符
                    if page_num < min(total_pages, 10) - 1:
                        doc.add_page_break()
                    
                except Exception as img_error:
                    logger.warning(f"添加第{page_num + 1}页截图失败: {img_error}")
                    doc.add_paragraph(f"第{page_num + 1}页图像处理失败")
                
                # 清理临时文件
                if os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
            
            pdf_document.close()
            
            if total_pages > 10:
                doc.add_paragraph(f"注意：文档共{total_pages}页，为节省空间仅显示前10页的图像。")
                
        except Exception as e:
            logger.error(f"添加页面截图失败: {e}")
    
    async def _process_pdf_fallback(self, pdf_path: str, doc: Document, filename: str, watermark_config: Dict[str, Any] = None, show_file_titles: bool = True, file_title_level: int = 2, is_last_file: bool = False):
        """使用PyMuPDF作为备选方案处理PDF"""
        try:
            logger.info(f"使用PyMuPDF处理PDF: {filename}")
            
            # 根据show_file_titles参数决定是否添加文件标题（直接显示文件名，不加前缀，去掉扩展名）
            if show_file_titles:
                filename_without_ext = os.path.splitext(filename)[0]
                self._format_heading(doc, filename_without_ext, level=file_title_level, center=False)
            
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            
            for page_num in range(total_pages):
                page = pdf_document.load_page(page_num)
                mat = fitz.Matrix(2.0, 2.0)  # 高质量渲染
                pix = page.get_pixmap(matrix=mat)
                
                temp_image_path = f"temp_page_{page_num + 1}_{uuid.uuid4().hex[:8]}.png"
                pix.save(temp_image_path)
                
                try:
                    # 使用智能分页逻辑，传递是否是最后一页的信息
                    # 只有当前是最后一页且是最后一个文件时，才算真正的最后一页
                    is_last_page_of_file = (page_num == total_pages - 1)
                    is_truly_last_page = is_last_page_of_file and is_last_file
                    self._add_page_content_with_smart_sizing(
                        doc, temp_image_path, page_num, show_file_titles, watermark_config, is_truly_last_page, total_pages
                    )
                    
                    if is_last_page_of_file:
                        logger.info(f"PDF处理完成，总页数: {total_pages}")
                        
                except Exception as e:
                    logger.error(f"页面处理失败: {e}")
                    doc.add_paragraph(f"第{page_num + 1}页处理失败: {str(e)}")
                
                # 清理临时文件
                if os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
            
            pdf_document.close()
            
            return {
                "success": True,
                "message": "PDF处理成功（PyMuPDF）",
                "text_extracted": False
            }
            
        except Exception as e:
            logger.error(f"PyMuPDF处理PDF失败: {e}")
            doc.add_paragraph(f"PDF处理失败 ({filename}): {str(e)}")
            return {
                "success": False,
                "message": f"PDF处理失败: {str(e)}"
            }
    
    async def process_image(self, image_path: str, doc: Document, filename: str, watermark_config: Dict[str, Any] = None, show_file_titles: bool = True, file_title_level: int = 2, is_last_file: bool = False):
        """处理图片文件 - 使用Docling进行智能处理"""
        try:
            logger.info(f"处理图片: {filename}")
            
            # 根据show_file_titles参数决定是否添加文件标题（直接显示文件名，不加前缀，去掉扩展名）
            if show_file_titles:
                filename_without_ext = os.path.splitext(filename)[0]
                self._format_heading(doc, filename_without_ext, level=file_title_level, center=False)
            
            # 尝试使用DoclingService处理图片
            if self.converter:
                try:
                    logger.info(f"使用DoclingService处理图片: {filename}")
                    
                    # 使用DoclingService转换图片
                    conv_result = await self.converter.convert_document(image_path)
                    
                    if conv_result.get("success"):
                        # 添加图片元数据信息
                        with Image.open(image_path) as img:
                            width, height = img.size
                            doc.add_paragraph(f"尺寸: {width} x {height} 像素")
                            doc.add_paragraph(f"格式: {img.format}")
                            doc.add_paragraph(f"模式: {img.mode}")
                        
                        # 如果提取到了文本内容，添加文本
                        text_content = conv_result.get("text", "")
                        if text_content.strip():
                            logger.info(f"DoclingService从图片中提取到文本内容")
                            text_heading = create_clean_heading(doc, "图片文本内容（OCR识别）", level=3)
                            doc.add_paragraph(text_content.strip())
                        else:
                            logger.info("DoclingService未从图片中提取到文本内容")
                        
                        # 使用智能图片大小计算（图片文件视为第0页）
                        try:
                            image_width, needs_page_break = self._calculate_image_size_for_page(
                                image_path, 0, True  # 明确传递True，确保单独图片使用5.5英寸
                            )
                        except Exception as e:
                            logger.warning(f"智能大小计算失败: {e}，使用默认方法")
                            image_width = self._calculate_image_width(0, True)
                        
                        # 添加原始图片（居中对齐）
                        image_heading = create_clean_heading(doc, "原始图片", level=3)
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.add_run().add_picture(image_path, width=image_width)
                        
                        # 只在不是最后一个文件时添加分页符
                        if not is_last_file:
                            doc.add_page_break()
                            logger.debug(f"图片处理完成（DoclingService），已添加分页符")
                        else:
                            logger.debug(f"图片处理完成（DoclingService），跳过分页符（最后一个文件）")
                        
                        return {
                            "success": True,
                            "message": "图片处理成功（DoclingService增强）",
                            "text_extracted": bool(text_content.strip())
                        }
                    else:
                        logger.warning(f"DoclingService转换失败: {conv_result.get('error')}")
                        # 降级到基础处理
                        pass
                        
                except Exception as docling_error:
                    logger.warning(f"DoclingService处理图片失败，降级到基础处理: {docling_error}")
                    # 降级到基础处理
                    pass
            
            # 基础图片处理（如果Docling不可用或失败）
            logger.info(f"使用基础方法处理图片: {filename}")
            
            with Image.open(image_path) as img:
                width, height = img.size
                doc.add_paragraph(f"尺寸: {width} x {height} 像素")
                doc.add_paragraph(f"格式: {img.format}")
                doc.add_paragraph(f"模式: {img.mode}")
            
            # 使用智能图片大小计算（图片文件视为第0页，明确传递True确保使用5.5英寸）
            try:
                image_width, needs_page_break = self._calculate_image_size_for_page(
                    image_path, 0, True  # 明确传递True，确保单独图片使用5.5英寸
                )
            except Exception as e:
                logger.warning(f"智能大小计算失败: {e}，使用默认方法")
                image_width = self._calculate_image_width(0, True)  # 图片文件视为第0页，明确传递True
            
            # 添加图片（居中对齐）
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.add_run().add_picture(image_path, width=image_width)
            
            # 只在不是最后一个文件时添加分页符
            if not is_last_file:
                doc.add_page_break()
                logger.debug(f"图片处理完成（基础），已添加分页符")
            else:
                logger.debug(f"图片处理完成（基础），跳过分页符（最后一个文件）")
            
            return {
                "success": True,
                "message": "图片处理成功（基础）"
            }
            
        except Exception as e:
            logger.error(f"图片处理失败: {e}")
            doc.add_paragraph(f"图片处理失败 ({filename}): {str(e)}")
            return {
                "success": False,
                "message": f"图片处理失败: {str(e)}"
            }
    
    async def convert_files_to_word(self, files: List[str], document_title: str):
        """将多个文件转换为Word文档"""
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置页面为A4大小
            from docx.shared import Inches
            from docx.enum.section import WD_ORIENTATION
            section = doc.sections[0]
            section.page_width = Inches(8.27)  # A4宽度 (210mm)
            section.page_height = Inches(11.69)  # A4高度 (297mm)
            section.orientation = WD_ORIENTATION.PORTRAIT  # 纵向
            section.left_margin = Inches(1.0)  # 1英寸页边距
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            
            self._format_heading(doc, document_title, level=0, center=True)
            
            processed_files = []
            results = []
            total_files = len(files)
            
            for file_index, file_path in enumerate(files):
                filename = os.path.basename(file_path)
                file_ext = os.path.splitext(filename)[1].lower()
                is_last_file = (file_index == total_files - 1)
                
                if file_ext == '.pdf':
                    # 使用Docling处理PDF
                    result = await self.process_pdf_with_docling(file_path, doc, filename)
                    processed_files.append(f"PDF: {filename}")
                    results.append(result)
                elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
                    # 处理图片
                    result = await self.process_image(file_path, doc, filename, None, True, 2, is_last_file)
                    processed_files.append(f"图片: {filename}")
                    results.append(result)
                else:
                    processed_files.append(f"不支持的格式: {filename}")
                    results.append({
                        "success": False,
                        "message": f"不支持的文件格式: {file_ext}"
                    })
            
            # 检查是否有处理失败的文件
            failed_count = sum(1 for r in results if not r.get("success", True))
            success_count = len(results) - failed_count
            
            return {
                "success": failed_count == 0,
                "message": f"处理完成，成功: {success_count}，失败: {failed_count}",
                "processed_files": processed_files,
                "results": results
            }
            
        except Exception as e:
            logger.error(f"文件转换失败: {e}")
            return {
                "success": False,
                "message": f"转换失败: {str(e)}"
            }

# 创建全局实例
docling_processor = DoclingDocumentProcessor()

def add_watermark_to_document(doc, watermark_config):
    """
    为Word文档添加可见的水印（直接在文档内容中插入）
    
    参数:
    - doc: Word文档对象
    - watermark_config: 水印配置字典
    """
    try:
        if not watermark_config.get("enabled") or not watermark_config.get("text"):
            logger.info("水印未启用或无水印文字，跳过水印添加")
            return
            
        watermark_text = watermark_config.get("text", "")
        font_size = watermark_config.get("font_size", 24)
        angle = watermark_config.get("angle", -45)
        opacity = watermark_config.get("opacity", 30)
        color_hex = watermark_config.get("color", "#808080")
        position = watermark_config.get("position", "center")
        
        logger.info(f"开始添加水印: '{watermark_text}', 大小: {font_size}, 位置: {position}")
        
        # 解析颜色 
        color_hex = color_hex.lstrip('#')
        if len(color_hex) == 6:
            r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        else:
            r, g, b = 128, 128, 128  # 默认灰色
        
        # 直接在文档内容中添加水印
        _add_visible_watermarks(doc, watermark_text, font_size, r, g, b, position)
        
        logger.info(f"水印添加完成: {watermark_text}")
                
    except Exception as e:
        logger.error(f"水印处理异常: {e}")

def _calculate_width(position):
    """根据位置计算水印宽度"""
    if position == "repeat":
        return "792pt"  # 全页宽度
    else:
        return "400pt"  # 居中水印宽度

def _calculate_height(position):
    """根据位置计算水印高度"""
    if position == "repeat":
        return "612pt"  # 全页高度
    else:
        return "100pt"  # 居中水印高度

def _add_visible_watermarks(doc, watermark_text, font_size, r, g, b, position):
    """在文档内容中添加可见的水印"""
    try:
        logger.info(f"添加可见水印: {watermark_text}, 位置: {position}")
        
        # 在文档开头添加水印说明
        watermark_intro = doc.add_paragraph()
        watermark_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
        intro_run = watermark_intro.add_run(f"【本文档包含水印: {watermark_text}】")
        intro_run.font.size = Pt(14)
        intro_run.font.color.rgb = RGBColor(r, g, b)
        intro_run.font.bold = True
        intro_run.font.italic = True
        
        # 添加分隔线
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator.add_run("=" * 50)
        sep_run.font.size = Pt(10)
        sep_run.font.color.rgb = RGBColor(r, g, b)
        
        if position == "repeat":
            # 平铺模式：在多个位置添加水印
            _add_repeat_watermarks(doc, watermark_text, font_size, r, g, b)
        else:
            # 单个大水印模式
            _add_center_watermark(doc, watermark_text, font_size, r, g, b)
            
        # 在文档末尾也添加水印
        doc.add_page_break()
        end_watermark = doc.add_paragraph()
        end_watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_run = end_watermark.add_run(f"【文档结束 - 水印: {watermark_text}】")
        end_run.font.size = Pt(16)
        end_run.font.color.rgb = RGBColor(r, g, b)
        end_run.font.bold = True
        
        logger.info("可见水印添加完成")
        
    except Exception as e:
        logger.error(f"可见水印添加失败: {e}")

def _add_repeat_watermarks(doc, watermark_text, font_size, r, g, b):
    """添加重复平铺水印"""
    try:
        # 添加3个明显的水印分布在文档中
        for i in range(3):
            # 添加分页符分隔
            if i > 0:
                doc.add_page_break()
            
            # 添加大号水印
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 创建醒目的水印文本
            watermark_run = watermark_para.add_run(f"★★★ {watermark_text} ★★★")
            watermark_run.font.size = Pt(font_size * 2)  # 2倍大小
            watermark_run.font.color.rgb = RGBColor(r, g, b)
            watermark_run.font.bold = True
            
            # 添加装饰行
            for j in range(3):
                deco_para = doc.add_paragraph()
                deco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if j == 1:
                    deco_text = f"◆◆◆ {watermark_text} ◆◆◆"
                    deco_size = Pt(font_size)
                else:
                    deco_text = "◇" * 20
                    deco_size = Pt(10)
                deco_run = deco_para.add_run(deco_text)
                deco_run.font.size = deco_size
                deco_run.font.color.rgb = RGBColor(r, g, b)
                deco_run.font.bold = True
            
        logger.info(f"平铺水印添加成功，共添加3个水印页面")
            
    except Exception as e:
        logger.error(f"平铺水印添加失败: {e}")

def _add_center_watermark(doc, watermark_text, font_size, r, g, b):
    """添加居中大水印"""
    try:
        # 添加一些空行定位到页面中央
        for _ in range(20):
            doc.add_paragraph()
        
        # 添加主水印
        main_watermark_para = doc.add_paragraph()
        main_watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 大号水印文字
        main_run = main_watermark_para.add_run(watermark_text)
        main_run.font.size = Pt(font_size * 3)  # 3倍大小
        main_run.font.color.rgb = RGBColor(r, g, b)
        main_run.font.bold = True
        
        # 添加装饰边框
        for i in range(3):
            border_para = doc.add_paragraph()
            border_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if i == 1:  # 中间行
                border_text = f"◆◆◆ {watermark_text} ◆◆◆"
                border_size = Pt(font_size)
            else:  # 上下边框
                border_text = "◇" * 30
                border_size = Pt(font_size // 2)
                
            border_run = border_para.add_run(border_text)
            border_run.font.size = border_size
            border_run.font.color.rgb = RGBColor(r, g, b)
            border_run.font.bold = True
        
        # 再添加一些空行
        for _ in range(20):
            doc.add_paragraph()
            
        logger.info(f"居中水印添加成功: {watermark_text}")
            
    except Exception as e:
        logger.error(f"居中水印添加失败: {e}")

def _add_content_watermarks(doc, watermark_text, font_size, r, g, b, position):
    """在文档内容中添加明显的水印（最终降级方案）"""
    try:
        # 在文档开头、中间、结尾添加大号水印
        watermark_positions = ["开头", "中间", "结尾"]
        
        for pos in watermark_positions:
            # 添加分页符
            if pos != "开头":
                doc.add_page_break()
            
            # 添加大号水印
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 使用特殊字符装饰水印
            run = watermark_para.add_run(f"◆◆◆ {watermark_text} ◆◆◆")
            run.font.size = Pt(font_size * 2)
            run.font.color.rgb = RGBColor(r, g, b)
            run.font.bold = True
            run.italic = True
            
            # 添加更多装饰行
            for i in range(3):
                deco_para = doc.add_paragraph()
                deco_run = deco_para.add_run("◇ " * 20)
                deco_run.font.size = Pt(8)
                deco_run.font.color.rgb = RGBColor(r, g, b)
                deco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        logger.info(f"内容水印添加成功: {watermark_text}")
        
    except Exception as e:
        logger.error(f"内容水印添加失败: {e}")



def format_heading_standalone(doc, text, level=1, center=False):
    """
    创建格式化的标题（独立函数版本）
    """
    from docx.oxml.shared import qn
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
    
    # 创建标题
    heading = doc.add_heading(text, level)
    
    # 设置对齐方式
    if center or level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 清除段落的列表样式和项目符号，防止出现小黑点
    try:
        # 清除段落的编号和项目符号
        pPr = heading._element.get_or_add_pPr()
        
        # 移除编号属性
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
        
        # 设置段落格式，明确禁用列表样式和分页控制
        if hasattr(heading, 'paragraph_format'):
            heading.paragraph_format.left_indent = None
            heading.paragraph_format.first_line_indent = None
            
            # 禁用分页控制选项，对应Word中的"分页"设置
            heading.paragraph_format.widow_control = False      # 孤行控制
            heading.paragraph_format.keep_with_next = False     # 与下段同页
            heading.paragraph_format.keep_together = False      # 段中不分页
            heading.paragraph_format.page_break_before = False  # 段前分页
            
    except Exception as style_error:
        logger.warning(f"清除标题样式时出错: {style_error}")
    
    # 设置字体
    if heading.runs:
        run = heading.runs[0]
        
        # 根据级别设置字体大小
        if level == 0:
            run.font.size = Pt(18)
            run.font.name = '楷体'
        elif level == 1:
            run.font.size = Pt(16)
            run.font.name = '楷体'
        elif level == 2:
            run.font.size = Pt(14)
            run.font.name = '楷体'
        else:
            run.font.size = Pt(12)
            run.font.name = '楷体'
        
        run.bold = True
        
        # 确保字体颜色是黑色
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
        
        # 设置中英文字体
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    return heading

class DocumentProcessor:
    """文档处理服务"""
    
    @staticmethod
    async def detect_file_type(file_path: str) -> str:
        """检测文件类型"""
        if MAGIC_AVAILABLE:
            try:
                mime = magic.Magic(mime=True)
                mime_type = mime.from_file(file_path)
                return mime_type
            except Exception as e:
                logger.warning(f"python-magic检测失败，使用文件扩展名: {e}")
        
        # 备选方案：根据文件扩展名推断MIME类型
        ext = file_path.split('.')[-1].lower()
        mime_mapping = {
            'pdf': 'application/pdf',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'png': 'image/png',
            'gif': 'image/gif',
            'bmp': 'image/bmp',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/msword',
            'txt': 'text/plain'
        }
        
        return mime_mapping.get(ext, f'application/{ext}')
    
    @staticmethod
    async def convert_to_word(file_path: str, filename: str = None) -> Tuple[str, int]:
        """
        将文件转换为Word格式
        
        参数:
        - file_path: 源文件路径
        - filename: 输出文件名(可选)
        
        返回:
        - 转换后的Word文档路径
        - 文档页数
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检测文件类型
        mime_type = await DocumentProcessor.detect_file_type(file_path)
        logger.info(f"文件类型: {mime_type}")
        
        # 生成输出文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not filename:
            filename = os.path.basename(file_path)
            basename, _ = os.path.splitext(filename)
        else:
            basename = filename
        
        output_filename = f"{basename}_{timestamp}.docx"
        output_path = os.path.join(CONVERTED_DIR, output_filename)
        
        # 根据文件类型调用相应的转换方法
        if "pdf" in mime_type.lower():
            return await DocumentProcessor.convert_pdf_to_word(file_path, output_path)
        elif "image" in mime_type.lower():
            return await DocumentProcessor.convert_image_to_word(file_path, output_path)
        elif "word" in mime_type.lower() or mime_type.endswith('document'):
            # 如果已经是Word文档，复制一份到转换目录
            shutil.copy2(file_path, output_path)
            
            # 尝试获取页数
            try:
                doc = Document(file_path)
                # Word文档页数计算不精确，这里估算一个值
                page_count = len(doc.paragraphs) // 40 + 1  # 假设每页40段落
            except:
                page_count = 1
            
            return output_path, page_count
        else:
            raise ValueError(f"不支持的文件类型: {mime_type}")

    @staticmethod
    async def convert_pdf_to_word(pdf_path: str, output_path: str) -> Tuple[str, int]:
        """
        将PDF转换为Word文档
        
        参数:
        - pdf_path: PDF文件路径
        - output_path: 输出Word文档路径
        
        返回:
        - 转换后的Word文档路径
        - PDF页数
        """
        # 创建Word文档
        doc = Document()
        
        try:
            # 打开PDF
            pdf_document = fitz.open(pdf_path)
            page_count = len(pdf_document)
            
            # 检测PDF类型
            pdf_type = await DocumentProcessor._detect_pdf_type_static(pdf_path)
            logger.info(f"PDF类型检测结果: {pdf_type}")
            
            # 添加文档标题
            if pdf_type == "scanned":
                format_heading_standalone(doc, "扫描件PDF转换结果", level=1, center=True)
                doc.add_paragraph("此PDF为扫描件，已通过OCR技术提取文本内容。")
            else:
                format_heading_standalone(doc, "PDF转换结果", level=1, center=True)
                doc.add_paragraph("此PDF包含可编辑文本，为保持原始格式已转换为图片形式。")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # 添加页面标题
                if page_num > 0:  # 第一页不添加页面分隔标题
                    format_heading_standalone(doc, f"第 {page_num + 1} 页(共 {page_count} 页)", level=2, center=True)
                
                # 对于扫描件PDF，尝试提取文本
                if pdf_type == "scanned":
                    text = page.get_text()
                    if text.strip():
                        doc.add_paragraph("提取的文本内容：")
                        doc.add_paragraph(text.strip())
                        doc.add_paragraph()  # 添加空行分隔
                
                # 将页面渲染为图片
                # 根据PDF类型使用不同的分辨率
                if pdf_type == "scanned":
                    mat = fitz.Matrix(2.0, 2.0)  # 扫描件使用标准分辨率
                else:
                    mat = fitz.Matrix(3.0, 3.0)  # 非扫描件使用高分辨率确保文本清晰
                
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # 创建临时图片文件
                img_stream = io.BytesIO(img_data)
                
                # 添加图片到Word（居中对齐）
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 根据页码和PDF类型确定图片大小
                if page_num == 0:
                    if pdf_type == "scanned":
                        img_width = Inches(5.5)  # 扫描件第一页适中大小
                    else:
                        img_width = Inches(6.0)  # 非扫描件第一页稍大
                else:
                    if pdf_type == "scanned":
                        img_width = Inches(6.23)  # 扫描件非第一页使用15.83厘米
                    else:
                        img_width = Inches(6.5)  # 非扫描件非第一页使用更大尺寸
                
                img_para.add_run().add_picture(img_stream, width=img_width)
                
                # 添加分页符（除最后一页）
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            
            # 保存Word文档
            doc.save(output_path)
            pdf_document.close()
            
            return output_path, page_count
            
        except Exception as e:
            logger.error(f"PDF转Word失败: {str(e)}")
            raise
    
    @staticmethod
    async def _detect_pdf_type_static(pdf_path: str) -> str:
        """静态方法：检测PDF类型：扫描件 vs 非扫描件"""
        try:
            import fitz
            
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            
            # 检查前3页（或全部页面，如果少于3页）
            pages_to_check = min(3, total_pages)
            text_ratio = 0
            image_ratio = 0
            
            for page_num in range(pages_to_check):
                page = pdf_document.load_page(page_num)
                
                # 获取文本
                text = page.get_text()
                text_length = len(text.strip())
                
                # 获取图片
                image_list = page.get_images()
                image_count = len(image_list)
                
                # 计算比例
                if text_length > 0:
                    text_ratio += 1
                if image_count > 0:
                    image_ratio += 1
            
            pdf_document.close()
            
            # 判断PDF类型
            text_percentage = text_ratio / pages_to_check
            image_percentage = image_ratio / pages_to_check
            
            logger.info(f"PDF类型分析 - 文本页面比例: {text_percentage:.2f}, 图片页面比例: {image_percentage:.2f}")
            
            # 如果大部分页面都有可提取文本，认为是非扫描件
            if text_percentage > 0.7:
                return "native"
            # 如果大部分页面都有图片且文本很少，认为是扫描件
            elif image_percentage > 0.7 and text_percentage < 0.3:
                return "scanned"
            else:
                # 混合类型，优先按非扫描件处理
                return "mixed"
                
        except Exception as e:
            logger.warning(f"PDF类型检测失败: {e}，默认按非扫描件处理")
            return "native"
    
    @staticmethod
    async def convert_image_to_word(image_path: str, output_path: str) -> Tuple[str, int]:
        """
        将图片转换为Word文档
        
        参数:
        - image_path: 图片文件路径
        - output_path: 输出Word文档路径
        
        返回:
        - 转换后的Word文档路径
        - 页数(图片为1页)
        """
        # 创建Word文档
        doc = Document()
        
        try:
            # 获取图片信息
            with Image.open(image_path) as img:
                width, height = img.size
                
                # 添加简要说明
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"图片尺寸: {width} x {height} 像素")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            # 添加图片（居中对齐）
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 单独图片使用第一页大小（因为通常是单页）
            img_para.add_run().add_picture(image_path, width=Inches(6.23))
            
            # 保存Word文档
            doc.save(output_path)
            
            return output_path, 1  # 图片算作1页
            
        except Exception as e:
            logger.error(f"图片转Word失败: {str(e)}")
            raise
    
    @staticmethod
    async def merge_documents(document_paths: List[str], output_filename: str = None) -> Tuple[str, int, float]:
        """
        合并多个Word文档为一个
        
        参数:
        - document_paths: 待合并Word文档路径列表
        - output_filename: 输出文件名(可选)
        
        返回:
        - 合并后的文档路径
        - 总页数
        - 处理时间(秒)
        """
        if not document_paths:
            raise ValueError("没有要合并的文档")
        
        start_time = time.time()
        
        # 生成输出文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not output_filename:
            output_filename = f"merged_{timestamp}.docx"
        
        output_path = os.path.join(GENERATED_DIR, output_filename)
        
        # 创建新文档作为主文档
        master_doc = Document()
        total_pages = 0
        
        try:
            # 遍历每个文档并合并
            for i, doc_path in enumerate(document_paths):
                # 打开子文档
                if not os.path.exists(doc_path):
                    logger.warning(f"文档不存在，已跳过: {doc_path}")
                    continue
                    
                try:
                    sub_doc = Document(doc_path)
                except Exception as e:
                    logger.error(f"无法打开文档 {doc_path}: {str(e)}")
                    continue
                
                # 添加分节符(除了第一个文档)
                if i > 0:
                    master_doc.add_section(WD_SECTION_START.NEW_PAGE)
                
                # 复制所有段落
                for para in sub_doc.paragraphs:
                    p = master_doc.add_paragraph()
                    p.paragraph_format.alignment = para.paragraph_format.alignment
                    p.paragraph_format.left_indent = para.paragraph_format.left_indent
                    p.paragraph_format.right_indent = para.paragraph_format.right_indent
                    p.paragraph_format.space_before = para.paragraph_format.space_before
                    p.paragraph_format.space_after = para.paragraph_format.space_after
                    p.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                    
                    for run in para.runs:
                        r = p.add_run(run.text)
                        r.bold = run.bold
                        r.italic = run.italic
                        r.underline = run.underline
                        if run.font.size:
                            r.font.size = run.font.size
                        if run.font.color.rgb:
                            r.font.color.rgb = run.font.color.rgb
                
                # 复制所有表格
                for table in sub_doc.tables:
                    tbl = master_doc.add_table(rows=0, cols=len(table.columns))
                    tbl.style = table.style
                    
                    for row in table.rows:
                        new_row = tbl.add_row()
                        for idx, cell in enumerate(row.cells):
                            if idx < len(new_row.cells):  # 防止列索引越界
                                new_cell = new_row.cells[idx]
                                new_cell.text = cell.text
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        for i, char in enumerate(run.text):
                                            if i < len(new_cell.paragraphs[0].runs):
                                                new_run = new_cell.paragraphs[0].runs[i]
                                                new_run.bold = run.bold
                                                new_run.italic = run.italic
                                                if run.font.size:
                                                    new_run.font.size = run.font.size
                
                # 复制图片和其他形状
                for shape in sub_doc.inline_shapes:
                    try:
                        if shape.type == 3:  # 图片
                            # 提取图片，保存为临时文件，然后添加到主文档
                            img_stream = io.BytesIO()
                            if hasattr(shape, '_inline') and hasattr(shape._inline, 'graphic') and \
                               hasattr(shape._inline.graphic, 'graphicData') and \
                               hasattr(shape._inline.graphic.graphicData, 'pic'):
                                
                                blob = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                                if hasattr(sub_doc.part, 'related_parts') and blob in sub_doc.part.related_parts:
                                    image_part = sub_doc.part.related_parts[blob]
                                    if hasattr(image_part, 'blob'):
                                        img_stream.write(image_part.blob)
                                        img_stream.seek(0)
                                        master_doc.add_picture(img_stream, width=shape.width, height=shape.height)
                    except Exception as img_err:
                        logger.error(f"复制图片失败: {str(img_err)}")
                
                # 估算页数(粗略估计)
                doc_pages = len(sub_doc.paragraphs) // 40 + 1
                total_pages += doc_pages
            
            # 保存合并后的文档
            master_doc.save(output_path)
            
            # 计算处理时间
            processing_time = time.time() - start_time
            
            return output_path, total_pages, processing_time
            
        except Exception as e:
            logger.error(f"合并文档失败: {str(e)}")
            raise
    
    @staticmethod
    async def fill_template(template_path: str, field_values: Dict[str, Any], output_filename: str = None) -> str:
        """
        填充Word模板
        
        参数:
        - template_path: 模板文件路径
        - field_values: 字段值字典 {field_key: value}
        - output_filename: 输出文件名(可选)
        
        返回:
        - 填充后的文档路径
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        # 生成输出文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not output_filename:
            basename = os.path.basename(template_path)
            name, _ = os.path.splitext(basename)
            output_filename = f"{name}_filled_{timestamp}.docx"
        
        output_path = os.path.join(GENERATED_DIR, output_filename)
        
        try:
            # 打开文档
            doc = Document(template_path)
            
            # 替换所有段落中的占位符
            for paragraph in doc.paragraphs:
                for key, value in field_values.items():
                    # 支持多种占位符格式
                    placeholders = [
                        f"{{${key}$}}",  # {$field_key$}
                        f"{{{key}}}",    # {field_key}
                        f"${key}$",      # $field_key$
                        f"[[{key}]]"     # [[field_key]]
                    ]
                    
                    for placeholder in placeholders:
                        if placeholder in paragraph.text:
                            str_value = str(value) if value is not None else ""
                            paragraph.text = paragraph.text.replace(placeholder, str_value)
            
            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in field_values.items():
                                placeholders = [
                                    f"{{${key}$}}",
                                    f"{{{key}}}",
                                    f"${key}$",
                                    f"[[{key}]]"
                                ]
                                
                                for placeholder in placeholders:
                                    if placeholder in paragraph.text:
                                        str_value = str(value) if value is not None else ""
                                        paragraph.text = paragraph.text.replace(placeholder, str_value)
            
            # 保存文档
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"填充模板失败: {str(e)}")
            raise
    
    @staticmethod
    async def analyze_template_fields(template_path: str) -> List[Dict[str, Any]]:
        """
        分析模板中的字段
        
        参数:
        - template_path: 模板文件路径
        
        返回:
        - 字段列表
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        try:
            doc = Document(template_path)
            fields = []
            
            # 定义可能的占位符正则模式
            # 这里使用简单字符串匹配而非正则，以便于演示
            placeholders_patterns = [
                ("{$", "$}"),   # {$field_key$}
                ("{", "}"),     # {field_key}
                ("$", "$"),     # $field_key$
                ("[[", "]]")    # [[field_key]]
            ]
            
            # 提取段落中的字段
            for paragraph in doc.paragraphs:
                text = paragraph.text
                
                for start_pat, end_pat in placeholders_patterns:
                    start_idx = 0
                    while True:
                        start_pos = text.find(start_pat, start_idx)
                        if start_pos == -1:
                            break
                            
                        end_pos = text.find(end_pat, start_pos + len(start_pat))
                        if end_pos == -1:
                            break
                            
                        # 提取字段名
                        field_key = text[start_pos + len(start_pat):end_pos]
                        
                        # 判断字段类型
                        field_type = "text"  # 默认为文本
                        if "date" in field_key.lower() or "日期" in field_key:
                            field_type = "date"
                        elif "number" in field_key.lower() or "金额" in field_key or "数量" in field_key:
                            field_type = "number"
                            
                        # 添加到结果
                        field = {
                            "field_key": field_key,
                            "field_name": field_key.replace("_", " ").title(),
                            "field_type": field_type,
                            "placeholder": text[start_pos:end_pos + len(end_pat)]
                        }
                        
                        # 检查是否已存在
                        if not any(f["field_key"] == field_key for f in fields):
                            fields.append(field)
                            
                        start_idx = end_pos + len(end_pat)
            
            # 提取表格中的字段
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            
                            for start_pat, end_pat in placeholders_patterns:
                                start_idx = 0
                                while True:
                                    start_pos = text.find(start_pat, start_idx)
                                    if start_pos == -1:
                                        break
                                        
                                    end_pos = text.find(end_pat, start_pos + len(start_pat))
                                    if end_pos == -1:
                                        break
                                        
                                    field_key = text[start_pos + len(start_pat):end_pos]
                                    
                                    # 判断字段类型
                                    field_type = "text"
                                    if "date" in field_key.lower() or "日期" in field_key:
                                        field_type = "date"
                                    elif "number" in field_key.lower() or "金额" in field_key or "数量" in field_key:
                                        field_type = "number"
                                        
                                    field = {
                                        "field_key": field_key,
                                        "field_name": field_key.replace("_", " ").title(),
                                        "field_type": field_type,
                                        "placeholder": text[start_pos:end_pos + len(end_pat)]
                                    }
                                    
                                    if not any(f["field_key"] == field_key for f in fields):
                                        fields.append(field)
                                        
                                    start_idx = end_pos + len(end_pat)
            
            return fields
            
        except Exception as e:
            logger.error(f"分析模板字段失败: {str(e)}")
            raise

    @staticmethod
    async def merge_word_documents(document_paths: List[str], output_path: str, 
                                  show_file_titles: bool = True, file_title_level: int = 2,
                                  add_page_breaks: bool = True) -> Dict[str, Any]:
        """
        拼接多个Word文档为一个文档
        
        参数:
        - document_paths: 待拼接Word文档路径列表
        - output_path: 输出文档路径
        - show_file_titles: 是否显示文件标题
        - file_title_level: 文件标题层级
        - add_page_breaks: 是否在文档间添加分页符
        
        返回:
        - 处理结果字典
        """
        if not document_paths:
            return {
                "success": False,
                "message": "没有要拼接的Word文档"
            }
        
        try:
            # 创建新文档作为主文档
            master_doc = Document()
            total_pages = 0
            processed_files = []
            
            # 遍历每个文档并拼接
            for i, doc_path in enumerate(document_paths):
                if not os.path.exists(doc_path):
                    logger.warning(f"Word文档不存在，已跳过: {doc_path}")
                    processed_files.append(f"跳过: {os.path.basename(doc_path)} (文件不存在)")
                    continue
                
                try:
                    # 打开子文档
                    sub_doc = Document(doc_path)
                    filename = os.path.basename(doc_path)
                    
                    # 如果不是第一个文档且需要分页符，添加分页符
                    if i > 0 and add_page_breaks:
                        master_doc.add_page_break()
                    
                    # 如果需要显示文件标题，添加标题
                    if show_file_titles:
                        # 移除文件扩展名
                        title_text = os.path.splitext(filename)[0]
                        format_heading_standalone(master_doc, title_text, level=file_title_level, center=False)
                    
                    # 复制子文档的所有段落到主文档
                    for paragraph in sub_doc.paragraphs:
                        # 创建新段落
                        new_para = master_doc.add_paragraph()
                        
                        # 复制段落格式
                        new_para.style = paragraph.style
                        new_para.alignment = paragraph.alignment
                        
                        # 复制段落内容
                        for run in paragraph.runs:
                            new_run = new_para.add_run(run.text)
                            # 复制运行格式
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.size = run.font.size
                            new_run.font.name = run.font.name
                            if run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                    
                    # 复制子文档的所有表格到主文档
                    for table in sub_doc.tables:
                        # 创建新表格
                        new_table = master_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = table.style
                        
                        # 复制表格内容
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_table.cell(i, j).text = cell.text
                    
                    # 复制子文档的所有图片到主文档
                    for rel in sub_doc.part.rels.values():
                        if "image" in rel.target_ref:
                            # 这里需要更复杂的图片处理逻辑
                            # 暂时跳过图片复制，因为需要处理关系引用
                            pass
                    
                    # 估算页数（简单估算）
                    estimated_pages = len(sub_doc.paragraphs) // 40 + 1
                    total_pages += estimated_pages
                    
                    processed_files.append(f"Word: {filename}")
                    logger.info(f"成功拼接Word文档: {filename}")
                    
                except Exception as e:
                    logger.error(f"处理Word文档失败 {doc_path}: {e}")
                    processed_files.append(f"失败: {os.path.basename(doc_path)}")
            
            # 保存拼接后的文档
            master_doc.save(output_path)
            
            return {
                "success": True,
                "message": f"成功拼接 {len(processed_files)} 个Word文档",
                "output_path": output_path,
                "total_pages": total_pages,
                "processed_files": processed_files
            }
            
        except Exception as e:
            logger.error(f"Word文档拼接失败: {e}")
            return {
                "success": False,
                "message": f"拼接失败: {str(e)}"
            }

# 单例实例
document_processor = DocumentProcessor() 