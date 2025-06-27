import os
import json
import time
import base64
import logging
from typing import Dict, List, Optional, Tuple, Any, Callable
from datetime import datetime
import openai
import pytesseract
import cv2
import numpy as np
from PIL import Image
import pdfplumber
import docx
from io import BytesIO
import httpx
import asyncio
from pathlib import Path

# 配置日志（需要在其他导入之前）
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Docling相关导入
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions
    from docling.backend.abstract_backend import DeclarativeDocumentBackend
    DOCLING_AVAILABLE = True
    logger.info("Docling模块可用")
except ImportError as e:
    DOCLING_AVAILABLE = False
    logger.warning(f"Docling模块不可用: {e}")

# 检查EasyOCR是否可用
try:
    import easyocr
    EASYOCR_AVAILABLE = True
    logger.info("EasyOCR模块可用")
except ImportError as e:
    EASYOCR_AVAILABLE = False
    logger.warning(f"EasyOCR模块不可用: {e}")

class AIService:
    """AI服务类，提供各种AI能力"""
    
    def __init__(self):
        self.ai_provider = os.getenv("AI_PROVIDER", "openai").lower()
        self.enable_ai = os.getenv("ENABLE_AI", "true").lower() == "true"
        
        # 初始化AI客户端
        self.ai_client = None
        if self.enable_ai:
            self._init_ai_client()
        
        # 初始化Docling转换器
        self.docling_converter = None
        if DOCLING_AVAILABLE:
            self._init_docling_converter()
        
        # 文档分类配置
        self.document_categories = {
            'performance_contract': '业绩合同',
            'award_certificate': '荣誉奖项',
            'qualification_certificate': '资质证照',
            'lawyer_certificate': '律师证',
            'other': '其他杂项'
        }
        
        # 业务领域分类
        self.business_fields = [
            '公司法律服务', '金融法律服务', '争议解决', '专业法律领域',
            '基础设施与能源', '房地产与土地', '国际贸易与海事', 
            '劳动与社会保障', '税务与财务', '新兴业务领域',
            '政府与公共事务', '跨境业务', '特殊行业'
        ]
        
        # 律师证标签配置
        self.lawyer_certificate_tags = {
            'position': ['合伙人', '律师'],  # 职位标签
            'business_fields': self.business_fields  # 业务领域标签
        }
        
        # Docling OCR配置（替代独立的EasyOCR）
        self.enable_docling_ocr = True
        self.docling_ocr_languages = ["ch_sim", "en"]
        
        # Ollama视觉配置
        self.ollama_vision_base_url = ""
        
        # 文档分类配置（重复定义，保持一致）
        self.document_categories = {
            "performance_contract": "业绩合同",
            "award_certificate": "荣誉奖项", 
            "qualification_certificate": "资质证照",
            "lawyer_certificate": "律师证",
            "other": "其他杂项"
        }
        
        # 业务领域列表
        self.business_fields = [
            "公司法律服务", "金融法律服务", "争议解决", "专业法律领域",
            "基础设施与能源", "房地产与土地", "国际贸易与海事", "劳动与社会保障",
            "税务与财务", "新兴业务领域", "政府与公共事务", "跨境业务"
        ]
        
        # 加载配置
        self.reload_config()
    
    def _get_setting_value(self, key: str, default: str = "") -> str:
        """从数据库获取设置值，如果不存在则返回环境变量或默认值"""
        try:
            from database import get_db
            from models import SystemSettings
            
            db = next(get_db())
            setting = db.query(SystemSettings).filter(
                SystemSettings.setting_key == key
            ).first()
            
            if setting and setting.setting_value:
                return setting.setting_value
            
            # 如果数据库中没有，尝试从环境变量获取
            env_mapping = {
                "ai_provider": "AI_PROVIDER",
                "ai_text_model": f"{self.ai_provider.upper()}_MODEL" if hasattr(self, 'ai_provider') else "OPENAI_MODEL",
                "ai_vision_model": f"{self.ai_provider.upper()}_VISION_MODEL" if hasattr(self, 'ai_provider') else "OPENAI_VISION_MODEL",
                "ai_api_key": f"{self.ai_provider.upper()}_API_KEY" if hasattr(self, 'ai_provider') else "OPENAI_API_KEY",
                "ai_base_url": f"{self.ai_provider.upper()}_BASE_URL" if hasattr(self, 'ai_provider') else "OPENAI_BASE_URL"
            }
            
            env_key = env_mapping.get(key, key.upper())
            return os.getenv(env_key, default)
            
        except Exception as e:
            logger.warning(f"无法从数据库获取设置 {key}: {str(e)}")
            # 回退到环境变量
            return os.getenv(key.upper(), default)
        finally:
            try:
                db.close()
            except:
                pass
    
    def reload_config(self):
        """重新加载配置并重新初始化AI客户端"""
        try:
            logger.info("重新加载AI服务配置...")
            
            # 重新读取AI提供商设置
            self.ai_provider = self._get_setting_value("ai_provider", "openai").lower()
            self.enable_ai = self._get_setting_value("enable_ai", "true").lower() == "true"
            
            # 重新初始化AI客户端
            self.ai_client = None
            if self.enable_ai:
                self._init_ai_client()
                
            # 重新初始化视觉模型客户端
            self._init_vision_client()
            
            # 重新初始化Docling转换器
            self._init_docling_converter()
            
            # 重新初始化Docling转换器
            self._init_docling_converter()
            
            logger.info(f"AI服务配置重载完成，当前提供商: {self.ai_provider}")
            
        except Exception as e:
            logger.error(f"重新加载AI配置失败: {str(e)}")
    
    def _init_ai_client(self):
        """初始化AI客户端"""
        try:
            if self.ai_provider == "openai":
                self.ai_client = openai.OpenAI(
                    api_key=self._get_setting_value("ai_api_key", os.getenv("OPENAI_API_KEY", "")),
                    base_url=self._get_setting_value("ai_base_url", os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1"))
                )
                self.model = self._get_setting_value("ai_text_model", os.getenv("OPENAI_MODEL", "gpt-4"))
                self.vision_model = self._get_setting_value("ai_vision_model", os.getenv("OPENAI_VISION_MODEL", "gpt-4-vision-preview"))
                
            elif self.ai_provider == "azure":
                from openai import AzureOpenAI
                self.ai_client = AzureOpenAI(
                    api_key=self._get_setting_value("ai_api_key", os.getenv("AZURE_OPENAI_API_KEY", "")),
                    api_version=self._get_setting_value("azure_api_version", os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01")),
                    azure_endpoint=self._get_setting_value("azure_endpoint", os.getenv("AZURE_OPENAI_ENDPOINT", ""))
                )
                self.model = self._get_setting_value("ai_text_model", os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4"))
                self.vision_model = self._get_setting_value("ai_vision_model", os.getenv("AZURE_OPENAI_VISION_DEPLOYMENT_NAME", "gpt-4-vision"))
                
            elif self.ai_provider == "ollama":
                # Ollama本地服务
                self.ai_client = openai.OpenAI(
                    api_key="ollama",  # Ollama不需要真实API key
                    base_url=self._get_setting_value("ai_base_url", os.getenv("OLLAMA_BASE_URL", "http://localhost:11434/v1"))
                )
                self.model = self._get_setting_value("ai_text_model", os.getenv("OLLAMA_MODEL", "llama3.2"))
                self.vision_model = self._get_setting_value("ai_vision_model", os.getenv("OLLAMA_VISION_MODEL", "llava:latest"))
                
            elif self.ai_provider == "custom":
                # 兼容OpenAI API的自定义服务
                self.ai_client = openai.OpenAI(
                    api_key=self._get_setting_value("ai_api_key", os.getenv("CUSTOM_AI_API_KEY", "")),
                    base_url=self._get_setting_value("ai_base_url", os.getenv("CUSTOM_AI_BASE_URL", ""))
                )
                self.model = self._get_setting_value("ai_text_model", os.getenv("CUSTOM_AI_MODEL", "gpt-4"))
                self.vision_model = self._get_setting_value("ai_vision_model", os.getenv("CUSTOM_AI_VISION_MODEL", "gpt-4-vision"))
                
            else:
                logger.warning(f"不支持的AI提供商: {self.ai_provider}，将禁用AI功能")
                self.enable_ai = False
            
            # 初始化视觉模型配置
            self.vision_provider = self._get_setting_value("vision_provider", self.ai_provider)
            self.vision_base_url = self._get_setting_value("vision_base_url", "")
            self.vision_api_key = self._get_setting_value("vision_api_key", "")
            
            # 如果视觉模型使用不同的提供商，创建独立的客户端
            if self.vision_provider != self.ai_provider:
                self._init_vision_client()
            else:
                self.vision_client = self.ai_client
            
            if self.ai_client:
                logger.info(f"AI客户端初始化成功 - 提供商: {self.ai_provider}, 文本模型: {self.model}")
                logger.info(f"视觉模型配置 - 提供商: {self.vision_provider}, 模型: {self.vision_model}")
                
        except Exception as e:
            logger.error(f"初始化AI客户端失败: {str(e)}")
            self.enable_ai = False
    
    def _init_vision_client(self):
        """初始化独立的视觉模型客户端"""
        try:
            if self.vision_provider == "ollama":
                self.vision_client = openai.OpenAI(
                    api_key="ollama",
                    base_url=self.ollama_vision_base_url or self._get_setting_value("ollama_vision_base_url", "http://localhost:11434/v1")
                )
                logger.info(f"Ollama视觉客户端初始化成功: {self.ollama_vision_base_url}")
                
            elif self.vision_provider == "openai":
                self.vision_client = openai.OpenAI(
                    api_key=self.vision_api_key or self._get_setting_value("vision_api_key", ""),
                    base_url=self.vision_base_url or "https://api.openai.com/v1"
                )
                
            elif self.vision_provider == "custom":
                self.vision_client = openai.OpenAI(
                    api_key=self.vision_api_key or self._get_setting_value("vision_api_key", ""),
                    base_url=self.vision_base_url or self._get_setting_value("vision_base_url", "")
                )
                
            else:
                logger.warning(f"不支持的视觉提供商: {self.vision_provider}")
                self.vision_client = self.ai_client
                
        except Exception as e:
            logger.error(f"初始化视觉客户端失败: {str(e)}")
            self.vision_client = self.ai_client
    
    def _init_docling_converter(self):
        """初始化Docling转换器"""
        try:
            # 配置PDF处理选项
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = False  # 禁用OCR避免性能问题
            pipeline_options.do_table_structure = True
            
            # 移除不兼容的配置项，避免backend错误
            try:
                pipeline_options.table_structure_options.do_cell_matching = True
            except AttributeError:
                # 如果该属性不存在，跳过
                pass
                
            # 简化OCR配置
            try:
                if hasattr(pipeline_options, 'ocr_options'):
                    pipeline_options.ocr_options.lang = ["chi_sim", "eng"]
                    pipeline_options.ocr_options.use_gpu = False
            except AttributeError:
                # 如果OCR选项不可用，跳过
                pass
            
            # 创建转换器（使用简化配置）
            self.docling_converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: pipeline_options
                }
            )
            logger.info("Docling转换器初始化成功（简化配置）")
        except Exception as e:
            logger.warning(f"Docling初始化失败，使用备用方案: {e}")
            # 使用最简单的配置作为备用
            try:
                self.docling_converter = DocumentConverter()
                logger.info("Docling转换器备用初始化成功")
            except Exception as fallback_e:
                logger.error(f"Docling备用初始化也失败: {fallback_e}")
                self.docling_converter = None
    
    def _init_easyocr_reader(self):
        """初始化EasyOCR读取器"""
        if not self.easyocr_enable or not EASYOCR_AVAILABLE:
            self.easyocr_reader = None
            return
            
        try:
            # 设置代理
            if self.easyocr_download_proxy:
                os.environ['http_proxy'] = self.easyocr_download_proxy
                os.environ['https_proxy'] = self.easyocr_download_proxy
                logger.info(f"设置EasyOCR下载代理: {self.easyocr_download_proxy}")
            
            # 设置模型路径
            if self.easyocr_model_path:
                os.environ['EASYOCR_MODULE_PATH'] = self.easyocr_model_path
                # 确保模型目录存在
                os.makedirs(self.easyocr_model_path, exist_ok=True)
                logger.info(f"EasyOCR模型路径: {self.easyocr_model_path}")
            
            # 初始化读取器
            logger.info("正在初始化EasyOCR读取器...")
            self.easyocr_reader = easyocr.Reader(
                self.easyocr_languages,
                gpu=self.easyocr_use_gpu,
                model_storage_directory=self.easyocr_model_path if self.easyocr_model_path else None,
                download_enabled=True
            )
            logger.info(f"EasyOCR读取器初始化成功，支持语言: {self.easyocr_languages}")
            
            # 清除代理环境变量
            if self.easyocr_download_proxy:
                os.environ.pop('http_proxy', None)
                os.environ.pop('https_proxy', None)
                
        except Exception as e:
            logger.error(f"初始化EasyOCR读取器失败: {str(e)}")
            self.easyocr_reader = None
            # 清除代理环境变量
            if self.easyocr_download_proxy:
                os.environ.pop('http_proxy', None)
                os.environ.pop('https_proxy', None)
    
    async def download_easyocr_models(self, progress_callback=None):
        """手动下载EasyOCR模型"""
        if not EASYOCR_AVAILABLE:
            return {"success": False, "error": "EasyOCR不可用"}
            
        try:
            # 设置代理
            if self.easyocr_download_proxy:
                os.environ['http_proxy'] = self.easyocr_download_proxy
                os.environ['https_proxy'] = self.easyocr_download_proxy
                logger.info(f"设置下载代理: {self.easyocr_download_proxy}")
            
            # 设置模型路径
            if self.easyocr_model_path:
                os.environ['EASYOCR_MODULE_PATH'] = self.easyocr_model_path
                os.makedirs(self.easyocr_model_path, exist_ok=True)
            
            if progress_callback:
                progress_callback({"status": "downloading", "progress": 10, "message": "开始下载模型..."})
            
            # 创建读取器（这会触发模型下载）
            reader = easyocr.Reader(
                self.easyocr_languages,
                gpu=self.easyocr_use_gpu,
                model_storage_directory=self.easyocr_model_path if self.easyocr_model_path else None,
                download_enabled=True
            )
            
            if progress_callback:
                progress_callback({"status": "testing", "progress": 80, "message": "测试模型..."})
            
            # 测试模型是否工作
            try:
                import numpy as np
                test_image = np.ones((100, 100, 3), dtype=np.uint8) * 255  # 白色图片
            except ImportError:
                # 如果numpy不可用，创建简单测试图片
                test_image = "/tmp/test_image.png"
                from PIL import Image
                img = Image.new('RGB', (100, 100), color='white')
                img.save(test_image)
            result = reader.readtext(test_image)
            
            if progress_callback:
                progress_callback({"status": "completed", "progress": 100, "message": "模型下载完成"})
            
            # 更新读取器
            self.easyocr_reader = reader
            
            # 清除代理环境变量
            if self.easyocr_download_proxy:
                os.environ.pop('http_proxy', None)
                os.environ.pop('https_proxy', None)
            
            return {
                "success": True,
                "message": f"EasyOCR模型下载成功，支持语言: {self.easyocr_languages}",
                "model_path": self.easyocr_model_path
            }
            
        except Exception as e:
            # 清除代理环境变量
            if self.easyocr_download_proxy:
                os.environ.pop('http_proxy', None)
                os.environ.pop('https_proxy', None)
            
            logger.error(f"EasyOCR模型下载失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def extract_text_from_image(self, image_path: str) -> Dict:
        """从图片中提取文字"""
        try:
            start_time = time.time()
            
            # 优先使用EasyOCR
            if self.easyocr_enable and self.easyocr_reader:
                try:
                    result = self.easyocr_reader.readtext(image_path)
                    text_lines = []
                    confidence_scores = []
                    
                    for (bbox, text, conf) in result:
                        if text.strip():
                            text_lines.append(text)
                            confidence_scores.append(conf)
                    
                    full_text = "\n".join(text_lines)
                    avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0
                    
                    processing_time = time.time() - start_time
                    
                    return {
                        "text": full_text,
                        "confidence": avg_confidence,
                        "processing_time": processing_time,
                        "method": "easyocr",
                        "details": {
                            "lines_detected": len(text_lines),
                            "confidence_scores": confidence_scores
                        }
                    }
                except Exception as e:
                    logger.warning(f"EasyOCR处理失败，降级到Tesseract: {str(e)}")
            
            # 降级到Tesseract OCR
            image = cv2.imread(image_path)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # 图像预处理
            gray = cv2.medianBlur(gray, 3)
            gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # OCR识别
            text = pytesseract.image_to_string(gray, lang='chi_sim+eng')
            
            processing_time = time.time() - start_time
            
            return {
                "text": text,
                "confidence": 0.8,  # 默认置信度
                "processing_time": processing_time,
                "method": "tesseract"
            }
            
        except Exception as e:
            logger.error(f"OCR处理失败: {str(e)}")
            return {
                "text": "",
                "confidence": 0.0,
                "processing_time": 0,
                "error": str(e)
            }
    
    async def extract_text_from_pdf(self, pdf_path: str) -> Dict:
        """从PDF中提取文字"""
        try:
            start_time = time.time()
            text_content = []
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        text_content.append({
                            "page": page_num + 1,
                            "text": text
                        })
            
            processing_time = time.time() - start_time
            
            return {
                "content": text_content,
                "total_pages": len(text_content),
                "processing_time": processing_time,
                "method": "pdfplumber"
            }
            
        except Exception as e:
            logger.error(f"PDF文字提取失败: {str(e)}")
            return {
                "content": [],
                "total_pages": 0,
                "processing_time": 0,
                "error": str(e)
            }
    
    async def extract_text_from_docx(self, docx_path: str) -> Dict:
        """从Word文档中提取文字"""
        try:
            start_time = time.time()
            
            doc = docx.Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            processing_time = time.time() - start_time
            
            return {
                "content": text_content,
                "full_text": "\n".join(text_content),
                "processing_time": processing_time,
                "method": "python-docx"
            }
            
        except Exception as e:
            logger.error(f"Word文档提取失败: {str(e)}")
            return {
                "content": [],
                "full_text": "",
                "processing_time": 0,
                "error": str(e)
            }
    
    async def analyze_award_document(self, text: str) -> Dict:
        """分析获奖文档内容"""
        if not self.enable_ai or not self.ai_client:
            return {
                "analysis": {"awards": [], "confidence": 0.0},
                "confidence": 0.0,
                "processing_time": 0,
                "message": "AI功能未启用或配置错误"
            }
            
        try:
            start_time = time.time()
            
            prompt = self._get_award_analysis_prompt()
            
            response = await self._call_ai_api(
                prompt=prompt,
                content=text,
                max_tokens=1500
            )
            
            processing_time = time.time() - start_time
            
            # 解析AI返回的结构化数据
            try:
                analysis_result = json.loads(response)
            except json.JSONDecodeError:
                # 如果返回的不是JSON，尝试提取关键信息
                analysis_result = self._parse_award_response(response)
            
            return {
                "analysis": analysis_result,
                "confidence": analysis_result.get("confidence", 0.7),
                "processing_time": processing_time,
                "raw_response": response
            }
            
        except Exception as e:
            logger.error(f"获奖文档分析失败: {str(e)}")
            return {
                "analysis": {},
                "confidence": 0.0,
                "processing_time": 0,
                "error": str(e)
            }
    
    async def analyze_performance_document(self, text: str) -> Dict:
        """分析业绩文档内容"""
        if not self.enable_ai or not self.ai_client:
            return {
                "analysis": {"performances": [], "confidence": 0.0},
                "confidence": 0.0,
                "processing_time": 0,
                "message": "AI功能未启用或配置错误"
            }
            
        try:
            start_time = time.time()
            
            prompt = self._get_performance_analysis_prompt()
            
            response = await self._call_ai_api(
                prompt=prompt,
                content=text,
                max_tokens=1500
            )
            
            processing_time = time.time() - start_time
            
            # 解析AI返回的结构化数据
            try:
                analysis_result = json.loads(response)
            except json.JSONDecodeError:
                analysis_result = self._parse_performance_response(response)
            
            return {
                "analysis": analysis_result,
                "confidence": analysis_result.get("confidence", 0.7),
                "processing_time": processing_time,
                "raw_response": response
            }
            
        except Exception as e:
            logger.error(f"业绩文档分析失败: {str(e)}")
            return {
                "analysis": {},
                "confidence": 0.0,
                "processing_time": 0,
                "error": str(e)
            }
    
    async def analyze_image_content(self, image_path: str) -> Dict:
        """分析图像内容"""
        if not self.enable_ai or not self.ai_client:
            return {
                "description": "AI功能未启用或配置错误",
                "processing_time": 0,
                "method": "disabled"
            }
            
        try:
            start_time = time.time()
            
            # 转换图像为base64
            with open(image_path, "rb") as image_file:
                image_base64 = base64.b64encode(image_file.read()).decode('utf-8')
            
            # 使用Vision API分析图像
            response = await self._call_vision_api(image_base64)
            
            processing_time = time.time() - start_time
            
            return {
                "description": response,
                "processing_time": processing_time,
                "method": f"{self.ai_provider}-vision"
            }
            
        except Exception as e:
            logger.error(f"图像分析失败: {str(e)}")
            return {
                "description": "",
                "processing_time": 0,
                "error": str(e)
            }
    
    async def _call_ai_api(self, prompt: str, content: str, max_tokens: int = 1000) -> str:
        """调用AI API"""
        if not self.ai_client:
            raise Exception("AI客户端未初始化")
            
        try:
            response = self.ai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": content}
                ],
                max_tokens=max_tokens,
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"AI API调用失败: {str(e)}")
            raise e
    
    async def _call_vision_api(self, image_base64: str) -> str:
        """调用Vision API"""
        if not hasattr(self, 'vision_client') or not self.vision_client:
            raise Exception("视觉模型客户端未初始化")
            
        try:
            # 构建消息内容
            if self.vision_provider == "ollama":
                # Ollama的消息格式
                response = await self._call_ollama_vision_api(image_base64)
            else:
                # OpenAI兼容格式
                response = await self._call_openai_vision_api(image_base64)
            
            return response
        except Exception as e:
            logger.error(f"Vision API调用失败: {str(e)}")
            raise e
    
    async def _call_ollama_vision_api(self, image_base64: str) -> str:
        """调用Ollama视觉API"""
        try:
            # Ollama的视觉模型调用方式
            response = self.vision_client.chat.completions.create(
                model=self.vision_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "请分析这张图片的内容，如果是获奖证书或法律行业相关文档，请详细描述奖项信息、获奖方、年份等关键信息。请用中文回答。"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_base64}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=1000,
                temperature=0.1
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Ollama Vision API调用失败: {str(e)}")
            # 如果Ollama调用失败，可以降级到OCR+文本分析
            raise e
    
    async def _call_openai_vision_api(self, image_base64: str) -> str:
        """调用OpenAI兼容的视觉API"""
        try:
            response = self.vision_client.chat.completions.create(
                model=self.vision_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "请分析这张图片的内容，如果是获奖证书或法律行业相关文档，请详细描述奖项信息、获奖方、年份等关键信息。请用中文回答。"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_base64}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=1000
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"OpenAI Vision API调用失败: {str(e)}")
            raise e
    
    def _get_award_analysis_prompt(self) -> str:
        """获取获奖分析提示词"""
        return """
你是一个专业的法律行业获奖信息分析专家。请分析以下文档内容，提取获奖相关信息。

支持的厂牌包括（请使用标准名称）：
- Chambers (Chambers and Partners)
- Legal 500 (The Legal 500)
- IFLR (International Financial Law Review)
- Who's Who Legal
- Best Lawyers
- ALB (Asian Legal Business)
- Legal Band
- Asialaw
- China Law & Practice
- IAM (Intellectual Asset Management)
- Global Arbitration Review
- Global Competition Review
- Benchmark Litigation
- China Business Law Journal
等知名法律评级机构

业务类型分类（请使用以下标准分类）：
- 公司法律服务：公司业务、并购重组、外商投资、私募股权/风险投资、公司治理、合规与监管
- 金融法律服务：银行与金融、资本市场、债券发行、资产证券化、基金与资产管理、保险法
- 争议解决：争议解决、国际仲裁、商事仲裁、诉讼代理、跨境争议
- 专业法律领域：知识产权、专利申请与保护、商标注册与维权、反不正当竞争
- 基础设施与能源：建设工程、基础设施、能源法、环境法、新能源
- 房地产：房地产、土地使用权、房地产开发、城市更新
- 国际贸易：国际贸易、航运海事、海关与贸易合规
- 税务财务：税法、税务争议、税务筹划、国际税务
- 新兴业务：医疗健康、互联网、人工智能、金融科技
- 跨境业务：跨境投资、境外上市、一带一路

请按照以下JSON格式返回结果：
{
    "awards": [
        {
            "title": "奖项名称",
            "brand": "厂牌名称（使用上述标准名称）",
            "year": 年份,
            "business_type": "业务类型（使用上述标准分类）",
            "description": "获奖描述",
            "url": "相关链接（如果有）",
            "confidence": 置信度(0-1)
        }
    ],
    "confidence": 整体置信度(0-1)
}

注意：
1. 如果文档中包含多个奖项，请全部提取
2. 厂牌名称必须使用上述标准名称，不要使用缩写或别名
3. 业务类型必须从上述分类中选择最匹配的类别
4. 如果某些信息不确定，请在confidence中体现
5. 年份应为4位数字格式
"""
    
    def _get_performance_analysis_prompt(self) -> str:
        """获取业绩分析提示词"""
        return """
你是一个专业的法律服务业绩信息分析专家。请分析以下文档内容，提取业绩相关信息。

业务领域标准分类（请使用以下分类）：
- 公司法律服务：公司业务、并购重组、外商投资、私募股权/风险投资、公司治理、合规与监管、反垄断与竞争法、数据保护与隐私、公司重组与破产
- 金融法律服务：银行与金融、资本市场、债券发行、资产证券化、基金与资产管理、保险法、融资租赁、金融科技、绿色金融、REITs
- 争议解决：争议解决、国际仲裁、商事仲裁、诉讼代理、调解服务、执行与保全、跨境争议、建设工程争议、金融争议
- 专业法律领域：知识产权、专利申请与保护、商标注册与维权、版权保护、商业秘密、反不正当竞争、娱乐法、体育法、网络法、电子商务法
- 基础设施与能源：建设工程、基础设施、能源法、石油天然气、电力法、新能源、环境法、碳中和与ESG、矿业法
- 房地产与土地：房地产、土地使用权、房地产开发、房地产投资、物业管理、城市更新、特色小镇、产业园区
- 国际贸易与海事：国际贸易、海关与贸易合规、反倾销与反补贴、自贸区业务、航运海事、船舶金融、货物运输、海事保险、港口法务
- 劳动与社会保障：劳动法、劳动争议、人力资源、社会保险、工伤赔偿、劳动合规、外籍员工、高管激励
- 税务与财务：税法、税务争议、税务筹划、国际税务、转让定价、关税与进出口税、增值税、企业所得税、个人所得税
- 新兴业务领域：医疗健康、生物医药、医疗器械、互联网、人工智能、区块链、虚拟货币、游戏法、教育法、食品药品
- 政府与公共事务：政府法律顾问、PPP项目、政府采购、行政法、刑事辩护、反腐败与职务犯罪、监察法、国家安全法
- 跨境业务：跨境投资、境外上市、QFII/QDII、外汇管理、国际制裁、一带一路、中美贸易、跨境数据传输、跨境电商
- 特殊行业：航空航天、汽车制造、化工医药、电信通信、传媒娱乐、金融科技、新零售、供应链金融、农业法、旅游法

项目类型分类：
- 长期顾问：为客户提供长期法律顾问服务
- 重大个案：为客户处理特定重大法律事务

请按照以下JSON格式返回结果：
{
    "performances": [
        {
            "client_name": "客户名称",
            "project_name": "项目名称",
            "project_type": "项目类型（长期顾问/重大个案）",
            "business_field": "业务领域（使用上述标准分类）",
            "start_date": "开始日期（YYYY-MM-DD格式）",
            "end_date": "结束日期（YYYY-MM-DD格式）",
            "year": 年份,
            "contract_amount": 合同金额,
            "currency": "货币单位（CNY/USD/EUR/HKD等）",
            "description": "项目描述",
            "confidence": 置信度(0-1)
        }
    ],
    "confidence": 整体置信度(0-1)
}

注意：
1. 如果是合同文档，重点提取合同金额、服务期限等信息
2. 如果是业绩汇总文档，提取所有项目信息
3. 日期格式请统一为YYYY-MM-DD
4. 金额请提取数字部分，货币单位单独标注
5. 业务领域必须从上述分类中选择最匹配的类别
6. 项目类型只能是"长期顾问"或"重大个案"
7. 年份应为4位数字格式
"""
    
    def _parse_award_response(self, response: str) -> Dict:
        """解析获奖分析响应（备用方法）"""
        # 简单的文本解析逻辑
        return {
            "awards": [],
            "confidence": 0.5,
            "raw_text": response
        }
    
    def _parse_performance_response(self, response: str) -> Dict:
        """解析业绩分析响应（备用方法）"""
        # 简单的文本解析逻辑
        return {
            "performances": [],
            "confidence": 0.5,
            "raw_text": response
        }

    async def chat_with_tools(
        self, 
        user_message: str, 
        system_prompt: str = "",
        tools: List[Dict[str, Any]] = None,
        tool_executor: Callable = None,
        max_tokens: int = 2000
    ) -> Dict[str, Any]:
        """支持工具调用的AI对话"""
        try:
            if not self.enable_ai or not self.ai_client:
                return {
                    "success": False,
                    "error": "AI服务未启用或未初始化"
                }
            
            # 构建消息历史
            messages = []
            if system_prompt:
                messages.append({"role": "system", "content": system_prompt})
            
            messages.append({"role": "user", "content": user_message})
            
            # 构建工具定义
            tool_definitions = []
            if tools:
                for tool in tools:
                    tool_definitions.append({
                        "type": "function",
                        "function": {
                            "name": tool["name"],
                            "description": tool["description"],
                            "parameters": {
                                "type": "object",
                                "properties": tool["parameters"],
                                "required": [k for k, v in tool["parameters"].items() if "default" not in v]
                            }
                        }
                    })
            
            # 调用AI API
            response = await self._call_ai_api_with_tools(
                messages=messages,
                tools=tool_definitions,
                max_tokens=max_tokens
            )
            
            # 处理工具调用
            tools_used = []
            if response.get("choices") and response["choices"][0].get("message", {}).get("tool_calls"):
                tool_calls = response["choices"][0]["message"]["tool_calls"]
                
                for tool_call in tool_calls:
                    tool_name = tool_call["function"]["name"]
                    tool_args = json.loads(tool_call["function"]["arguments"])
                    
                    # 执行工具
                    if tool_executor:
                        try:
                            tool_result = await tool_executor(tool_name, tool_args)
                            tools_used.append({
                                "tool_name": tool_name,
                                "arguments": tool_args,
                                "result": tool_result
                            })
                        except Exception as e:
                            logger.error(f"工具执行失败 {tool_name}: {str(e)}")
                            tools_used.append({
                                "tool_name": tool_name,
                                "arguments": tool_args,
                                "error": str(e)
                            })
            
            # 如果有工具调用结果，再次调用AI进行总结
            if tools_used:
                # 构建包含工具结果的用户消息
                tool_results_message = f"""
用户问题：{user_message}

已执行的工具和结果：
{json.dumps(tools_used, ensure_ascii=False, indent=2)}

请根据工具执行结果，为用户提供完整的回答。
"""
                
                messages.append({"role": "user", "content": tool_results_message})
                
                # 再次调用AI
                final_response = await self._call_ai_api_with_tools(
                    messages=messages,
                    tools=[],  # 不再需要工具调用
                    max_tokens=max_tokens
                )
                
                return {
                    "success": True,
                    "response": final_response["choices"][0]["message"]["content"],
                    "tools_used": tools_used,
                    "raw_response": final_response
                }
            else:
                return {
                    "success": True,
                    "response": response["choices"][0]["message"]["content"],
                    "tools_used": [],
                    "raw_response": response
                }
                
        except Exception as e:
            logger.error(f"AI工具对话失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def analyze_text(self, prompt: str, max_tokens: int = 2000) -> Dict[str, Any]:
        """分析文本内容"""
        try:
            if not self.enable_ai or not self.ai_client:
                return {
                    "success": False,
                    "error": "AI服务未启用或未初始化"
                }
            
            messages = [
                {"role": "system", "content": "你是一个专业的法律文档分析助手，请根据用户的要求分析文档内容。"},
                {"role": "user", "content": prompt}
            ]
            
            response = await self._call_ai_api_with_tools(
                messages=messages,
                tools=[],
                max_tokens=max_tokens
            )
            
            return {
                "success": True,
                "analysis": response["choices"][0]["message"]["content"],
                "raw_response": response
            }
            
        except Exception as e:
            logger.error(f"文本分析失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _call_ai_api_with_tools(
        self, 
        messages: List[Dict[str, Any]], 
        tools: List[Dict[str, Any]] = None,
        max_tokens: int = 2000
    ) -> Dict[str, Any]:
        """调用支持工具的AI API"""
        try:
            if self.ai_provider == "openai":
                kwargs = {
                    "model": self.model,
                    "messages": messages,
                    "max_tokens": max_tokens,
                    "temperature": 0.7
                }
                
                if tools:
                    kwargs["tools"] = tools
                    kwargs["tool_choice"] = "auto"
                
                response = self.ai_client.chat.completions.create(**kwargs)
                return response.model_dump()
                
            elif self.ai_provider == "azure":
                kwargs = {
                    "model": self.model,
                    "messages": messages,
                    "max_tokens": max_tokens,
                    "temperature": 0.7
                }
                
                if tools:
                    kwargs["tools"] = tools
                    kwargs["tool_choice"] = "auto"
                
                response = self.ai_client.chat.completions.create(**kwargs)
                return response.model_dump()
                
            elif self.ai_provider == "custom":
                kwargs = {
                    "model": self.model,
                    "messages": messages,
                    "max_tokens": max_tokens,
                    "temperature": 0.7
                }
                
                if tools:
                    kwargs["tools"] = tools
                    kwargs["tool_choice"] = "auto"
                
                response = self.ai_client.chat.completions.create(**kwargs)
                return response.model_dump()
                
            else:
                raise ValueError(f"不支持的AI提供商: {self.ai_provider}")
                
        except Exception as e:
            logger.error(f"AI API调用失败: {str(e)}")
            raise e
    
    async def classify_document_with_docling(self, file_path: str) -> Dict[str, Any]:
        """使用Docling分析文档并进行分类"""
        try:
            if not self.docling_converter:
                return await self._fallback_document_classification(file_path)
            
            start_time = time.time()
            
            # 使用Docling处理文档
            conv_result = self.docling_converter.convert(file_path)
            document = conv_result.document
            
            # 提取文本内容
            full_text = document.export_to_text()
            
            # 提取表格信息
            tables_info = []
            try:
                for table in document.tables:
                    table_text = table.export_to_text() if hasattr(table, 'export_to_text') else str(table)
                    tables_info.append(table_text)
            except:
                pass
            
            # 提取图片信息
            images_info = []
            try:
                for image in document.pictures:
                    image_info = {
                        'caption': getattr(image, 'caption', ''),
                        'text': getattr(image, 'text', '')
                    }
                    images_info.append(image_info)
            except:
                pass
            
            processing_time = time.time() - start_time
            
            # 使用AI进行分类
            classification_result = await self._classify_document_content(
                full_text=full_text,
                tables=tables_info,
                images=images_info
            )
            
            return {
                "success": True,
                "method": "docling",
                "processing_time": processing_time,
                "extracted_content": {
                    "text": full_text,
                    "tables": tables_info,
                    "images": images_info
                },
                "classification": classification_result
            }
            
        except Exception as e:
            logger.error(f"Docling文档分类失败: {str(e)}")
            # 降级到备用方案
            return await self._fallback_document_classification(file_path)
    
    async def classify_document_with_vision(self, file_path: str) -> Dict[str, Any]:
        """使用视觉模型分析文档"""
        try:
            if not self.enable_ai or not self.ai_client:
                return {
                    "success": False,
                    "error": "AI服务未启用"
                }
            
            start_time = time.time()
            
            # 如果是PDF，转换第一页为图片
            if file_path.lower().endswith('.pdf'):
                image_path = await self._convert_pdf_to_image(file_path)
                if not image_path:
                    return {"success": False, "error": "PDF转图片失败"}
            else:
                image_path = file_path
            
            # 转换图像为base64
            with open(image_path, "rb") as image_file:
                image_base64 = base64.b64encode(image_file.read()).decode('utf-8')
            
            # 构建分类提示词
            classification_prompt = self._get_document_classification_prompt()
            
            # 调用视觉API
            response = await self._call_vision_api_for_classification(
                image_base64, 
                classification_prompt
            )
            
            processing_time = time.time() - start_time
            
            # 解析分类结果
            try:
                classification_result = json.loads(response)
            except json.JSONDecodeError:
                classification_result = self._parse_classification_response(response)
            
            # 清理临时图片文件
            if file_path.lower().endswith('.pdf') and image_path != file_path:
                try:
                    os.remove(image_path)
                except:
                    pass
            
            return {
                "success": True,
                "method": "vision_model",
                "processing_time": processing_time,
                "classification": classification_result,
                "raw_response": response
            }
            
        except Exception as e:
            logger.error(f"视觉模型文档分类失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def smart_document_analysis(self, file_path: str, enable_vision: bool = True, enable_ocr: bool = True) -> Dict[str, Any]:
        """智能文档分析 - 综合使用文本提取、OCR和视觉模型"""
        try:
            start_time = time.time()
            
            results = {
                "file_path": file_path,
                "text_extraction_result": None,
                "ocr_result": None,
                "vision_result": None,
                "text_analysis_result": None,
                "final_classification": None,
                "confidence": 0.0
            }
            
            # 1. 文本提取（Docling或传统方法）
            logger.info("开始文本提取...")
            if DOCLING_AVAILABLE and self.docling_converter:
                logger.info("使用Docling进行文档分析...")
                docling_result = await self.classify_document_with_docling(file_path)
                results["text_extraction_result"] = docling_result
            else:
                # 使用备用文本提取
                logger.info("使用备用文本提取方法...")
                fallback_result = await self._fallback_document_classification(file_path)
                results["text_extraction_result"] = fallback_result
            
            # 2. OCR文本识别（针对图片或PDF文档）
            if enable_ocr:
                logger.info("开始OCR文本识别...")
                ocr_result = await self._perform_ocr_analysis(file_path)
                results["ocr_result"] = ocr_result
            
            # 3. 视觉模型分析
            if enable_vision and self.enable_ai:
                logger.info("使用视觉模型进行文档分析...")
                vision_result = await self.classify_document_with_vision(file_path)
                results["vision_result"] = vision_result
            
            # 4. 文本模型分析（基于提取的文本内容）
            if self.enable_ai:
                logger.info("使用文本模型进行内容分析...")
                text_analysis_result = await self._perform_text_analysis(results)
                results["text_analysis_result"] = text_analysis_result
            
            # 5. 综合分析结果（多重验证）
            final_classification = self._merge_multi_source_results(
                text_extraction=results.get("text_extraction_result", {}).get("classification"),
                ocr_analysis=results.get("ocr_result", {}).get("classification"),
                vision_analysis=results.get("vision_result", {}).get("classification"),
                text_analysis=results.get("text_analysis_result", {}).get("classification")
            )
            
            results["final_classification"] = final_classification
            results["confidence"] = final_classification.get("confidence", 0.0)
            results["processing_time"] = time.time() - start_time
            
            logger.info(f"智能文档分析完成，最终分类：{final_classification.get('category')}，置信度：{final_classification.get('confidence')}")
            
            return {
                "success": True,
                "results": results
            }
            
        except Exception as e:
            logger.error(f"智能文档分析失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _perform_ocr_analysis(self, file_path: str) -> Dict[str, Any]:
        """执行OCR文本识别分析"""
        try:
            ocr_result = {
                "success": False,
                "text": "",
                "confidence": 0.0,
                "method": "none",
                "classification": None
            }
            
            # 判断文件类型
            if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
                # 图片文件直接OCR
                text_result = await self.extract_text_from_image(file_path)
                ocr_result.update({
                    "success": True,
                    "text": text_result.get("text", ""),
                    "confidence": text_result.get("confidence", 0.0),
                    "method": text_result.get("method", "tesseract")
                })
                
            elif file_path.lower().endswith('.pdf'):
                # PDF文件先转换为图片再OCR
                temp_image_path = await self._convert_pdf_to_image(file_path, 0)
                if temp_image_path:
                    text_result = await self.extract_text_from_image(temp_image_path)
                    ocr_result.update({
                        "success": True,
                        "text": text_result.get("text", ""),
                        "confidence": text_result.get("confidence", 0.0),
                        "method": text_result.get("method", "tesseract")
                    })
                    # 清理临时文件
                    try:
                        os.remove(temp_image_path)
                    except:
                        pass
            
            # 如果OCR成功提取到文本，进行分类分析
            if ocr_result["success"] and ocr_result["text"].strip():
                logger.info(f"OCR提取文本成功，长度：{len(ocr_result['text'])}字符")
                classification = await self._classify_document_content(ocr_result["text"])
                ocr_result["classification"] = classification
            
            return ocr_result
            
        except Exception as e:
            logger.error(f"OCR分析失败: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "classification": {"category": "other", "confidence": 0.0}
            }
    
    async def _perform_text_analysis(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """基于已提取的文本进行AI分析"""
        try:
            # 收集所有可用的文本内容
            all_texts = []
            
            # 从文本提取结果中获取文本
            text_extraction = results.get("text_extraction_result")
            if text_extraction and text_extraction.get("success"):
                extracted_content = text_extraction.get("extracted_content", {})
                if isinstance(extracted_content, dict):
                    if extracted_content.get("text"):
                        all_texts.append(f"文档提取文本：\n{extracted_content['text']}")
                    if extracted_content.get("tables"):
                        all_texts.append(f"表格内容：\n{chr(10).join(extracted_content['tables'])}")
                
            # 从OCR结果中获取文本
            ocr_result = results.get("ocr_result")
            if ocr_result and ocr_result.get("success") and ocr_result.get("text"):
                all_texts.append(f"OCR识别文本：\n{ocr_result['text']}")
            
            if not all_texts:
                return {
                    "success": False,
                    "error": "没有找到可分析的文本内容",
                    "classification": {"category": "other", "confidence": 0.0}
                }
            
            # 合并所有文本进行分析
            combined_text = "\n\n".join(all_texts)
            logger.info(f"准备分析合并文本，总长度：{len(combined_text)}字符")
            
            # 调用AI进行文本分析
            classification = await self._classify_document_content(combined_text)
            
            return {
                "success": True,
                "analyzed_text_length": len(combined_text),
                "classification": classification
            }
            
        except Exception as e:
            logger.error(f"文本分析失败: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "classification": {"category": "other", "confidence": 0.0}
            }
    
    def _merge_multi_source_results(self, text_extraction=None, ocr_analysis=None, vision_analysis=None, text_analysis=None) -> Dict[str, Any]:
        """合并多种来源的分析结果"""
        try:
            # 收集所有结果
            all_results = []
            result_sources = []
            
            if text_extraction and text_extraction.get("category"):
                all_results.append(text_extraction)
                result_sources.append("text_extraction")
            
            if ocr_analysis and ocr_analysis.get("category"):
                all_results.append(ocr_analysis)
                result_sources.append("ocr")
            
            if vision_analysis and vision_analysis.get("category"):
                all_results.append(vision_analysis)
                result_sources.append("vision")
            
            if text_analysis and text_analysis.get("category"):
                all_results.append(text_analysis)
                result_sources.append("text_ai")
            
            if not all_results:
                return {
                    "category": "other",
                    "confidence": 0.0,
                    "source": "none",
                    "reason": "没有有效的分析结果"
                }
            
            # 专门处理律师证识别
            lawyer_cert_results = [r for r in all_results if self._is_lawyer_certificate(r)]
            if lawyer_cert_results:
                # 如果有律师证识别结果，优先使用
                best_lawyer_result = max(lawyer_cert_results, key=lambda x: x.get("confidence", 0))
                best_lawyer_result["category"] = "lawyer_certificate"  # 确保分类正确
                best_lawyer_result["source"] = f"lawyer_cert_detection_{result_sources[all_results.index(best_lawyer_result)]}"
                
                # 添加律师证相关的标签建议
                suggested_tags = []
                if "律师" in str(best_lawyer_result.get("description", "")):
                    suggested_tags.append("律师")
                if "合伙人" in str(best_lawyer_result.get("description", "")):
                    suggested_tags.append("合伙人")
                
                best_lawyer_result["suggested_tags"] = suggested_tags
                return best_lawyer_result
            
            # 按置信度排序，选择最佳结果
            all_results.sort(key=lambda x: x.get("confidence", 0), reverse=True)
            best_result = all_results[0]
            
            # 如果最高置信度的结果置信度较低，进行一致性验证
            if best_result.get("confidence", 0) < 0.7:
                # 检查是否有多个结果一致
                category_counts = {}
                for result in all_results:
                    category = result.get("category", "other")
                    category_counts[category] = category_counts.get(category, 0) + 1
                
                # 如果有多个结果一致，提高置信度
                most_common_category = max(category_counts.items(), key=lambda x: x[1])
                if most_common_category[1] > 1:  # 有至少2个结果一致
                    for result in all_results:
                        if result.get("category") == most_common_category[0]:
                            result["confidence"] = min(result.get("confidence", 0) + 0.2, 0.9)
                            result["source"] = f"consensus_{len(result_sources)}_sources"
                            return result
            
            # 记录分析来源
            best_result["source"] = f"best_of_{len(result_sources)}_sources"
            best_result["analysis_sources"] = result_sources
            
            return best_result
            
        except Exception as e:
            logger.error(f"合并多源分析结果失败: {str(e)}")
            return {
                "category": "other",
                "confidence": 0.0,
                "source": "error",
                "error": str(e)
            }
    
    def _is_lawyer_certificate(self, result: Dict[str, Any]) -> bool:
        """判断是否为律师证"""
        if not result:
            return False
        
        # 检查关键词
        text_content = str(result.get("description", "")) + " " + str(result.get("keywords", []))
        text_content = text_content.lower()
        
        lawyer_cert_keywords = [
            "律师执业证", "执业证书", "律师证", "执业证号", 
            "律师执业", "执业律师", "司法局", "司法厅",
            "律师姓名", "执业机构", "证书编号"
        ]
        
        return any(keyword in text_content for keyword in lawyer_cert_keywords)

    async def extract_document_tags(self, file_path: str, existing_tags: List[str] = None) -> Dict[str, Any]:
        """从文档中提取标签"""
        try:
            # 先进行智能分析
            analysis_result = await self.smart_document_analysis(file_path)
            
            if not analysis_result.get("success"):
                return analysis_result
            
            classification = analysis_result["results"]["final_classification"]
            
            # 基于分类结果生成标签
            suggested_tags = []
            
            # 基于文档类型添加标签
            if classification.get("category"):
                category_name = self.document_categories.get(
                    classification["category"], 
                    classification["category"]
                )
                suggested_tags.append(category_name)
            
            # 基于业务领域添加标签
            if classification.get("business_field"):
                suggested_tags.append(classification["business_field"])
            
            # 基于年份添加标签
            if classification.get("year"):
                suggested_tags.append(f"{classification['year']}年")
            
            # 基于关键信息添加标签
            if classification.get("keywords"):
                suggested_tags.extend(classification["keywords"])
            
            # 去重并过滤
            suggested_tags = list(set([tag for tag in suggested_tags if tag and len(tag.strip()) > 0]))
            
            # 合并现有标签
            if existing_tags:
                all_tags = list(set(existing_tags + suggested_tags))
            else:
                all_tags = suggested_tags
            
            return {
                "success": True,
                "suggested_tags": suggested_tags,
                "all_tags": all_tags,
                "classification": classification
            }
            
        except Exception as e:
            logger.error(f"提取文档标签失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _classify_document_content(self, full_text: str, tables: List[str] = None, images: List[Dict] = None) -> Dict[str, Any]:
        """基于文档内容进行AI分类"""
        try:
            if not self.enable_ai or not self.ai_client:
                return {"category": "other", "confidence": 0.0}
            
            # 构建分析内容
            content_parts = [f"文档文本内容：\n{full_text}"]
            
            if tables:
                content_parts.append(f"表格信息：\n{chr(10).join(tables)}")
            
            if images:
                image_texts = [img.get('text', '') + ' ' + img.get('caption', '') for img in images if img.get('text') or img.get('caption')]
                if image_texts:
                    content_parts.append(f"图片相关文本：\n{chr(10).join(image_texts)}")
            
            analysis_content = "\n\n".join(content_parts)
            
            # 构建分类提示词
            prompt = self._get_comprehensive_classification_prompt()
            
            # 调用AI进行分类
            response = await self._call_ai_api(
                prompt=prompt,
                content=analysis_content,
                max_tokens=1000
            )
            
            # 解析结果
            try:
                classification_result = json.loads(response)
            except json.JSONDecodeError:
                classification_result = self._parse_classification_response(response)
            
            return classification_result
            
        except Exception as e:
            logger.error(f"AI内容分类失败: {str(e)}")
            return {"category": "other", "confidence": 0.0}
    
    async def _fallback_document_classification(self, file_path: str) -> Dict[str, Any]:
        """备用文档分类方案"""
        try:
            start_time = time.time()
            
            # 使用传统方法提取文本
            if file_path.lower().endswith('.pdf'):
                text_result = await self.extract_text_from_pdf(file_path)
                full_text = "\n".join([page["text"] for page in text_result.get("content", [])])
            elif file_path.lower().endswith('.docx'):
                text_result = await self.extract_text_from_docx(file_path)
                full_text = text_result.get("full_text", "")
            else:
                # 图片文件使用OCR
                text_result = await self.extract_text_from_image(file_path)
                full_text = text_result.get("text", "")
            
            processing_time = time.time() - start_time
            
            # 基于文本进行分类
            classification_result = await self._classify_document_content(full_text)
            
            return {
                "success": True,
                "method": "fallback",
                "processing_time": processing_time,
                "extracted_content": {"text": full_text},
                "classification": classification_result
            }
            
        except Exception as e:
            logger.error(f"备用文档分类失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _convert_pdf_to_image(self, pdf_path: str, page_num: int = 0) -> Optional[str]:
        """将PDF的指定页面转换为图片"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(pdf_path)
            page = doc[page_num]
            
            # 转换为图片
            mat = fitz.Matrix(2, 2)  # 缩放因子
            pix = page.get_pixmap(matrix=mat)
            
            # 保存临时图片
            temp_image_path = pdf_path.replace('.pdf', f'_page_{page_num}.png')
            pix.save(temp_image_path)
            
            doc.close()
            return temp_image_path
            
        except Exception as e:
            logger.error(f"PDF转图片失败: {str(e)}")
            return None
    
    async def _call_vision_api_for_classification(self, image_base64: str, prompt: str) -> str:
        """调用视觉API进行文档分类"""
        try:
            if self.vision_provider == "ollama":
                # 使用Ollama进行分类
                response = self.vision_client.chat.completions.create(
                    model=self.vision_model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{image_base64}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=1000,
                    temperature=0.1
                )
                return response.choices[0].message.content
            else:
                # 使用其他提供商（OpenAI兼容）
                response = await self._call_ai_api_with_tools(
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{image_base64}"
                                    }
                                }
                            ]
                        }
                    ],
                    tools=[],
                    max_tokens=1000
                )
                return response["choices"][0]["message"]["content"]
            
        except Exception as e:
            logger.error(f"视觉分类API调用失败: {str(e)}")
            raise e
    
    def _merge_classification_results(self, docling_result: Dict = None, vision_result: Dict = None) -> Dict[str, Any]:
        """合并Docling和视觉模型的分类结果"""
        try:
            merged_result = {
                "category": "other",
                "business_field": None,
                "confidence": 0.0,
                "keywords": [],
                "year": None,
                "source": "unknown"
            }
            
            # 如果有Docling结果，优先使用
            if docling_result and docling_result.get("category"):
                merged_result.update(docling_result)
                merged_result["source"] = "docling"
                
                # 如果还有视觉结果，进行验证和补充
                if vision_result and vision_result.get("confidence", 0) > merged_result.get("confidence", 0):
                    # 如果视觉结果置信度更高，使用视觉结果
                    if vision_result.get("category") and vision_result["confidence"] > 0.8:
                        merged_result["category"] = vision_result["category"]
                        merged_result["confidence"] = vision_result["confidence"]
                        merged_result["source"] = "vision_override"
                    
                    # 补充业务领域信息
                    if not merged_result.get("business_field") and vision_result.get("business_field"):
                        merged_result["business_field"] = vision_result["business_field"]
                    
                    # 合并关键词
                    vision_keywords = vision_result.get("keywords", [])
                    if vision_keywords:
                        merged_result["keywords"] = list(set(merged_result.get("keywords", []) + vision_keywords))
            
            # 如果只有视觉结果
            elif vision_result and vision_result.get("category"):
                merged_result.update(vision_result)
                merged_result["source"] = "vision_only"
            
            return merged_result
            
        except Exception as e:
            logger.error(f"合并分类结果失败: {str(e)}")
            return {
                "category": "other",
                "confidence": 0.0,
                "source": "error"
            }
    
    def _get_document_classification_prompt(self) -> str:
        """获取文档分类提示词（从数据库读取）"""
        try:
            prompt = self._get_setting_value("classification_vision_prompt", "")
            if not prompt:
                # 如果数据库中没有，使用默认提示词
                prompt = self._get_default_vision_prompt()
                logger.warning("使用默认视觉分类提示词")
            
            # 替换业务领域占位符
            if "{business_fields}" in prompt:
                prompt = prompt.replace("{business_fields}", ', '.join(self.business_fields))
            
            return prompt
        except Exception as e:
            logger.error(f"获取分类提示词失败: {str(e)}")
            return self._get_default_vision_prompt()
    
    def _get_default_vision_prompt(self) -> str:
        """获取默认视觉分类提示词"""
        return f"""
你是一个专业的法律文档分类专家。请分析这份文档图片，判断文档类型并提取关键信息。

文档类型分类（请仔细判断）：
1. performance_contract - 业绩合同：法律服务合同、委托协议、顾问协议等
2. award_certificate - 荣誉奖项：Chambers、Legal 500、Best Lawyers等法律行业奖项证书、排名认证等
3. qualification_certificate - 资质证照：律师事务所执业许可证、营业执照、组织机构代码证等机构资质
4. lawyer_certificate - 律师证：个人律师执业证书（包含律师姓名、执业证号、执业机构等信息）
5. other - 其他杂项：不属于以上类别的其他文档

特别注意分类规则：
- 个人律师执业证书 → lawyer_certificate（律师证）
- 律师事务所执业许可证 → qualification_certificate（资质证照）
- 营业执照等机构证照 → qualification_certificate（资质证照）

律师证特有字段（如果是律师证，请提取）：
- 律师姓名
- 执业证号
- 执业机构
- 发证机关（通常是司法局/司法厅）
- 年龄/身份证号（如有显示）

业务领域分类（如果适用）：
{', '.join(self.business_fields)}

请按照以下JSON格式返回结果：
{{
    "category": "文档类型代码",
    "category_name": "文档类型中文名称",
    "business_field": "业务领域（如果适用）",
    "year": 年份,
    "keywords": ["关键词1", "关键词2", "关键词3"],
    "confidence": 置信度(0-1),
    "description": "文档描述和分类依据",
    "specific_type": "具体类型（如：律师执业证、事务所许可证、Chambers排名等）",
    "lawyer_info": {{
        "name": "律师姓名（如果是律师证）",
        "certificate_number": "执业证号（如果是律师证）",
        "law_firm": "执业机构（如果是律师证）",
        "issuing_authority": "发证机关（如果是律师证）",
        "age": "年龄（如果显示）",
        "id_number": "身份证号（如果显示）"
    }}
}}
"""
    
    def _get_comprehensive_classification_prompt(self) -> str:
        """获取综合分类提示词（从数据库读取）"""
        try:
            prompt = self._get_setting_value("classification_text_prompt", "")
            if not prompt:
                # 如果数据库中没有，使用默认提示词
                prompt = self._get_default_text_prompt()
                logger.warning("使用默认文本分类提示词")
            
            # 替换业务领域占位符
            if "{business_fields}" in prompt:
                prompt = prompt.replace("{business_fields}", ', '.join(self.business_fields))
            
            return prompt
        except Exception as e:
            logger.error(f"获取综合分类提示词失败: {str(e)}")
            return self._get_default_text_prompt()
    
    def _get_default_text_prompt(self) -> str:
        """获取默认文本分类提示词"""
        return f"""
你是一个专业的法律文档分类专家。请基于提供的文档内容进行分类分析。

文档类型分类（请仔细判断）：
1. performance_contract - 业绩合同：法律服务合同、委托协议、顾问协议、服务协议等
2. award_certificate - 荣誉奖项：Chambers、Legal 500、Best Lawyers、IFLR、Who's Who Legal等法律行业奖项证书、排名认证等
3. qualification_certificate - 资质证照：律师事务所执业许可证、营业执照、组织机构代码证、统一社会信用代码证等机构资质
4. lawyer_certificate - 律师证：个人律师执业证书（包含律师姓名、执业证号、执业机构等信息）
5. other - 其他杂项：不属于以上类别的其他文档

特别注意分类规则：
- 个人律师执业证书 → lawyer_certificate（律师证）  
- 律师事务所执业许可证 → qualification_certificate（资质证照）
- 营业执照等机构证照 → qualification_certificate（资质证照）

律师证关键词：律师执业证、执业证书、律师证、执业证号、律师姓名、执业机构、司法局、司法厅

业务领域分类：
{', '.join(self.business_fields)}

请按照以下JSON格式返回结果：
{{
    "category": "文档类型代码",
    "category_name": "文档类型中文名称", 
    "business_field": "业务领域（如果适用）",
    "year": 年份,
    "keywords": ["关键词1", "关键词2", "关键词3"],
    "confidence": 置信度(0-1),
    "description": "分类依据和文档描述",
    "specific_type": "具体类型（如：律师执业证、事务所许可证、Chambers排名等）",
    "key_entities": {{
        "holder_name": "持证人/获奖方名称",
        "issuer": "颁发/发证机构",
        "certificate_number": "证书/许可证号（如果是资质证照）",
        "award_name": "奖项名称（如果是获奖）",
        "client_name": "客户名称（如果是合同）",
        "amount": "合同金额（如果是合同）",
        "date_issued": "颁发/签署日期"
    }},
    "lawyer_info": {{
        "name": "律师姓名（如果是律师证）",
        "certificate_number": "执业证号（如果是律师证）",
        "law_firm": "执业机构（如果是律师证）",
        "issuing_authority": "发证机关（如果是律师证）",
        "age": "年龄（如果显示）",
        "id_number": "身份证号（如果显示）"
    }}
}}
"""
    
    def _parse_classification_response(self, response: str) -> Dict[str, Any]:
        """解析分类响应（备用方法）"""
        try:
            # 改进的关键词匹配分类（按优先级）
            response_lower = response.lower()
            
            # 优先级1：资质证照（最高优先级）
            qualification_keywords = [
                '律师执业证', '执业证书', '执业许可证', '律师事务所执业许可证',
                '营业执照', '统一社会信用代码', '组织机构代码', 
                '资质证明', '认证证书', '从业资格', '会员证书',
                '司法局', '司法厅', '司法部', '市场监督管理局',
                '证号', '许可证号', '执业证号', '信用代码'
            ]
            
            # 优先级2：荣誉奖项
            award_keywords = [
                'chambers', 'legal 500', 'best lawyers', 'iflr', 'who\'s who legal',
                '排名', '评级', '奖项', '荣誉', '推荐律师',
                'ranked', 'recommended', 'leading', 'notable', 'rising star'
            ]
            
            # 优先级3：业绩合同
            contract_keywords = [
                '合同', '协议', '委托书', '服务协议', '顾问协议',
                '甲方', '乙方', '委托方', '受托方',
                '服务费', '律师费', '顾问费', '合同金额',
                '服务期限', '合同期限', '委托期间'
            ]
            
            # 按优先级进行匹配
            category = 'other'
            confidence = 0.3
            specific_type = None
            
            # 检查资质证照关键词（最高优先级）
            qualification_matches = [kw for kw in qualification_keywords if kw in response_lower]
            if qualification_matches:
                category = 'qualification_certificate'
                confidence = 0.7 + min(len(qualification_matches) * 0.1, 0.2)  # 0.7-0.9
                
                # 确定具体类型
                if any(kw in response_lower for kw in ['律师执业证', '执业证书']):
                    specific_type = '律师执业证'
                elif any(kw in response_lower for kw in ['执业许可证', '律师事务所']):
                    specific_type = '律师事务所执业许可证'
                elif '营业执照' in response_lower:
                    specific_type = '营业执照'
                else:
                    specific_type = '资质证照'
            
            # 检查荣誉奖项关键词
            elif any(kw in response_lower for kw in award_keywords):
                category = 'award_certificate'
                award_matches = [kw for kw in award_keywords if kw in response_lower]
                confidence = 0.6 + min(len(award_matches) * 0.1, 0.3)  # 0.6-0.9
                
                # 确定具体类型
                if 'chambers' in response_lower:
                    specific_type = 'Chambers排名'
                elif 'legal 500' in response_lower:
                    specific_type = 'Legal 500排名'
                elif 'best lawyers' in response_lower:
                    specific_type = 'Best Lawyers排名'
                else:
                    specific_type = '法律行业奖项'
            
            # 检查合同关键词
            elif any(kw in response_lower for kw in contract_keywords):
                category = 'performance_contract'
                contract_matches = [kw for kw in contract_keywords if kw in response_lower]
                confidence = 0.6 + min(len(contract_matches) * 0.1, 0.2)  # 0.6-0.8
                specific_type = '法律服务合同'
            
            # 尝试提取年份
            import re
            year_match = re.search(r'20\d{2}', response)
            year = int(year_match.group()) if year_match else None
            
            # 提取关键词用于标签
            all_keywords = qualification_keywords + award_keywords + contract_keywords
            found_keywords = [kw for kw in all_keywords if kw in response_lower][:3]
            
            return {
                "category": category,
                "category_name": self.document_categories.get(category, category),
                "business_field": None,
                "year": year,
                "keywords": found_keywords,
                "confidence": confidence,
                "description": f"基于关键词匹配的分类结果，匹配到：{', '.join(found_keywords) if found_keywords else '无明确关键词'}",
                "specific_type": specific_type,
                "raw_text": response
            }
            
        except Exception as e:
            logger.error(f"解析分类响应失败: {str(e)}")
            return {
                "category": "other",
                "category_name": "其他杂项",
                "confidence": 0.0,
                "description": f"分类解析失败: {str(e)}",
                "error": str(e)
            }

# 创建全局AI服务实例
ai_service = AIService() 