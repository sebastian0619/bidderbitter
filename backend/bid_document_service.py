"""投标文档生成服务"""
import os
import logging
from typing import Dict, List, Optional
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

# 输出目录
OUTPUT_DIR = Path("generated_docs")
OUTPUT_DIR.mkdir(exist_ok=True)


class BidDocumentService:
    """投标文档生成服务"""
    
    def __init__(self):
        self.output_dir = OUTPUT_DIR
    
    def generate_bid_document(self, project, sections: List, files: List, db=None) -> Dict:
        """生成投标文档
        
        Args:
            project: 项目对象
            sections: 章节列表
            files: 关联文件列表
            db: 数据库会话
        
        Returns:
            生成结果
        """
        if not HAS_DOCX:
            return {"success": False, "error": "python-docx 未安装，请运行: apt install python3-docx"}
        
        try:
            
            logger.info(f"开始生成投标文档: {project.name}")
            
            # 创建 Word 文档
            doc = Document()
            
            # 添加封面
            self._add_cover_page(doc, project)
            
            # 添加目录页
            self._add_toc_page(doc, sections)
            
            # 按章节添加内容
            for section in sorted(sections, key=lambda s: s.order):
                self._add_section(doc, section, files)
            
            # 保存文档
            filename = f"{project.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.output_dir / filename
            doc.save(str(filepath))
            
            logger.info(f"投标文档生成成功: {filepath}")
            
            return {
                "success": True,
                "filepath": str(filepath),
                "filename": filename,
                "file_size": os.path.getsize(str(filepath))
            }
            
        except ImportError:
            logger.error("python-docx 未安装")
            return {"success": False, "error": "python-docx 未安装，请运行: apt install python3-docx"}
        except Exception as e:
            logger.error(f"生成投标文档失败: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _add_cover_page(self, doc, project):
        """添加封面页"""
        # 添加空行
        for _ in range(6):
            doc.add_paragraph()
        
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("投标文件")
        run.font.size = Pt(36)
        run.bold = True
        
        doc.add_paragraph()
        
        # 项目信息
        info_items = [
            ("项目名称", project.name),
            ("招标人", project.tender_company or "待填写"),
            ("招标代理机构", project.tender_agency or "待填写"),
            ("投标人", project.bidder_name or "待填写"),
            ("截止日期", project.deadline.strftime("%Y年%m月%d日") if project.deadline else "待填写"),
            ("编制日期", datetime.now().strftime("%Y年%m月%d日")),
        ]
        
        for label, value in info_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}：{value}")
            run.font.size = Pt(14)
        
        doc.add_page_break()
    
    def _add_toc_page(self, doc, sections):
        """添加目录页"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("目  录")
        run.bold = True
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        for i, section in enumerate(sorted(sections, key=lambda s: s.order), 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {section.title}")
            run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_section(self, doc, section, files):
        """添加章节内容"""
        # 章节标题
        heading = doc.add_heading(section.title, level=1)
        
        if section.description:
            p = doc.add_paragraph(section.description)
            p.style.font.size = Pt(11)
        
        # 获取该章节的文件
        section_files = [f for f in files if self._file_matches_section(f, section)]
        
        if section_files:
            for file in section_files:
                self._add_file_to_doc(doc, file)
        else:
            doc.add_paragraph("（暂无相关材料）")
        
        doc.add_page_break()
    
    def _file_matches_section(self, file, section) -> bool:
        """判断文件是否属于该章节"""
        if not section.section_type:
            return True
        
        # 根据文件分类匹配章节类型
        category = file.ai_category or file.category or ""
        section_type = section.section_type.lower()
        
        mapping = {
            "qualification": ["资质证照", "资质", "证书"],
            "performance": ["业绩", "合同"],
            "award": ["奖项荣誉", "奖项"],
            "finance": ["财务资料", "财务"],
            "team": ["团队", "律师"],
        }
        
        keywords = mapping.get(section_type, [])
        return any(kw in category for kw in keywords) if keywords else True
    
    def _add_file_to_doc(self, doc, file):
        """将文件添加到文档"""
        from docx.shared import Pt, Cm
        
        # 添加文件名作为子标题
        doc.add_heading(file.display_name, level=2)
        
        # 如果是图片，插入图片
        if file.file_type == "image":
            try:
                doc.add_picture(file.storage_path, width=Cm(15))
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {str(e)}]")
        
        # 如果是 PDF，添加说明
        elif file.file_type == "pdf":
            doc.add_paragraph(f"[PDF 文件: {file.original_filename}]")
        
        # 其他类型
        else:
            doc.add_paragraph(f"[文件: {file.original_filename}]")


# 全局实例
bid_document_service = BidDocumentService()
