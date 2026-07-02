"""招标文件分析服务

使用 lexitool / python-docx 读取招标文件，提取：
1. 项目基本信息（项目名称、招标人、截止日期等）
2. 文档结构（章节标题）
3. 评分标准、资质要求等关键信息
"""
import re
import logging
from typing import Dict, List, Optional
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


class TenderAnalyzer:
    """招标文件分析器"""
    
    # 常见招标信息关键词
    FIELD_PATTERNS = {
        "project_name": [
            r"项目名称[：:]\s*(.+)",
            r"采购项目[：:]\s*(.+)",
            r"工程名称[：:]\s*(.+)",
            r"招标项目[：:]\s*(.+)",
        ],
        "tender_company": [
            r"招标人[：:]\s*(.+)",
            r"采购人[：:]\s*(.+)",
            r"业主[：:]\s*(.+)",
            r"甲方[：:]\s*(.+)",
        ],
        "tender_agency": [
            r"招标代理[机构]*[：:]\s*(.+)",
            r"代理机构[：:]\s*(.+)",
        ],
        "bidder_name": [
            r"投标人[：:]\s*(.+)",
            r"供应商[：:]\s*(.+)",
            r"乙方[：:]\s*(.+)",
        ],
        "deadline": [
            r"投标截止[时间日期]*[：:]\s*(\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?)",
            r"截止时间[：:]\s*(\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?)",
            r"开标时间[：:]\s*(\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?)",
        ],
        "budget": [
            r"预算[金额]*[：:]\s*([\d,.]+)\s*[万]?元",
            r"项目预算[：:]\s*([\d,.]+)\s*[万]?元",
            r"控制价[：:]\s*([\d,.]+)\s*[万]?元",
        ],
        "project_number": [
            r"项目编号[：:]\s*(.+)",
            r"招标编号[：:]\s*(.+)",
            r"采购编号[：:]\s*(.+)",
        ]
    }
    
    # 常见章节类型映射
    SECTION_TYPE_MAP = {
        "投标函": "cover",
        "投标函附录": "cover",
        "报价": "pricing",
        "报价表": "pricing",
        "商务标": "pricing",
        "资质": "qualification",
        "资质证明": "qualification",
        "资格": "qualification",
        "业绩": "performance",
        "案例": "performance",
        "项目经验": "performance",
        "团队": "team",
        "人员": "team",
        "项目团队": "team",
        "服务方案": "proposal",
        "技术方案": "proposal",
        "实施方案": "proposal",
        "服务承诺": "proposal",
        "售后": "proposal",
    }
    
    # 需要提交的材料类型及关键词
    REQUIRED_MATERIALS = {
        "资质证照": {
            "keywords": ["营业执照", "执业许可证", "资质证书", "律师执业证", "律所执业证"],
            "section_type": "qualification",
            "description": "律所及律师资质证明文件"
        },
        "业绩证明": {
            "keywords": ["业绩", "合同复印件", "服务协议", "项目经验", "类似项目", "同类项目"],
            "section_type": "performance",
            "description": "法律服务业绩合同等证明材料"
        },
        "团队资料": {
            "keywords": ["项目负责人", "团队成员", "人员简历", "律师简历", "执业经验"],
            "section_type": "team",
            "description": "项目团队成员简历及资质"
        },
        "财务资料": {
            "keywords": ["审计报告", "财务报表", "纳税证明", "社保缴纳", "财务状况"],
            "section_type": "finance",
            "description": "财务审计报告及纳税证明"
        },
        "公司介绍": {
            "keywords": ["公司简介", "律所简介", "基本情况", "组织架构", "办公场所"],
            "section_type": "company",
            "description": "律所基本情况介绍"
        },
        "服务方案": {
            "keywords": ["服务方案", "技术方案", "实施方案", "服务承诺", "工作计划"],
            "section_type": "proposal",
            "description": "法律服务方案及承诺"
        },
        "获奖荣誉": {
            "keywords": ["获奖", "荣誉", "奖项", "排名", "评级"],
            "section_type": "award",
            "description": "法律行业获奖及排名证明"
        },
        "授权委托": {
            "keywords": ["授权委托书", "法定代表人身份证明", "身份证复印件"],
            "section_type": "authorization",
            "description": "法定代表人证明及授权委托书"
        }
    }
    
    def analyze_with_docx(self, file_path: str) -> Dict:
        """使用 python-docx 分析招标文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            分析结果
        """
        # 先检查文件扩展名
        import os
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.doc':
            # .doc 格式直接用旧版解析器
            return self._analyze_old_doc(file_path)
        
        try:
            from docx import Document
            
            # 检查文件是否是有效的 zip 文件 (docx 是 zip 格式)
            import zipfile
            if not zipfile.is_zipfile(file_path):
                # 可能是 .doc 格式，尝试用 antiword 或纯文本提取
                return self._analyze_old_doc(file_path)
            
            doc = Document(file_path)
            
            # 提取所有文本
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            text = "\n".join(full_text)
            
            # 提取项目信息
            project_info = self._extract_project_info(text)
            
            # 提取章节结构
            sections = self._extract_sections(doc)
            
            # 分析需要提交的材料
            required_materials = self._analyze_required_materials(text)
            
            # 合并章节和材料要求
            sections = self._merge_sections_with_materials(sections, required_materials)
            
            return {
                "success": True,
                "project_info": project_info,
                "sections": sections,
                "required_materials": required_materials,
                "text_preview": text[:2000],
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"分析招标文件失败: {e}")
            # 尝试用纯文本方式提取
            return self._analyze_as_text(file_path)
    
    def _analyze_old_doc(self, file_path: str) -> Dict:
        """分析旧版 .doc 格式文件"""
        try:
            import olefile
            
            # 使用 olefile 读取 .doc 文件
            ole = olefile.OleFileIO(file_path)
            
            # 尝试读取 WordDocument 流
            if ole.exists('WordDocument'):
                # 读取主文档流
                word_stream = ole.openstream('WordDocument')
                content = word_stream.read()
                
                # 提取文本（简单方法：查找可读文本）
                text = self._extract_text_from_doc_binary(content)
                
                if len(text) > 100:
                    ole.close()
                    return self._process_text(text)
            
            # 尝试读取其他流
            for stream_name in ole.listdir():
                stream_path = '/'.join(stream_name)
                if 'Table' not in stream_path and 'Data' not in stream_path:
                    try:
                        stream = ole.openstream(stream_path)
                        content = stream.read()
                        text = self._extract_text_from_doc_binary(content)
                        if len(text) > 200:
                            ole.close()
                            return self._process_text(text)
                    except:
                        continue
            
            ole.close()
            
            # 兜底：尝试用 antiword
            import subprocess
            result = subprocess.run(['antiword', '-m', 'UTF-8', file_path], capture_output=True, timeout=30)
            if result.returncode == 0:
                text = result.stdout
                if len(text) > 100:
                    return self._process_text(text)
            
            return {"success": False, "error": "无法解析 .doc 格式文件，建议转换为 .docx 或 .pdf 格式后重试"}
            
        except Exception as e:
            logger.error(f"解析 .doc 文件失败: {e}")
            return {"success": False, "error": f"解析 .doc 文件失败: {str(e)}"}
    
    def _extract_text_from_doc_binary(self, data: bytes) -> str:
        """从 .doc 二进制数据中提取文本"""
        import re
        
        # 尝试 UTF-16 LE 解码（Word 常用编码）
        try:
            text = data.decode('utf-16-le', errors='ignore')
            # 过滤不可读字符
            text = re.sub(r'[^\x20-\x7e\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\n\r]', ' ', text)
            text = re.sub(r'\s+', ' ', text)
            if len(text) > 100:
                return text
        except:
            pass
        
        # 尝试 GBK 解码
        try:
            text = data.decode('gbk', errors='ignore')
            text = re.sub(r'[^\x20-\x7e\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\n\r]', ' ', text)
            text = re.sub(r'\s+', ' ', text)
            if len(text) > 100:
                return text
        except:
            pass
        
        # 最后尝试 ASCII
        text = data.decode('ascii', errors='ignore')
        text = re.sub(r'[^\x20-\x7e\n\r]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text
    
    def _analyze_as_text(self, file_path: str) -> Dict:
        """将文件作为纯文本分析（最后的兜底方案）"""
        try:
            # 尝试读取为二进制并提取可读文本
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # 提取可读的 UTF-8 文本
            text = content.decode('utf-8', errors='ignore')
            
            # 过滤掉不可读字符
            import re
            text = re.sub(r'[^\x20-\x7e\u4e00-\u9fff\n]', ' ', text)
            text = re.sub(r'\s+', ' ', text)
            
            if len(text) < 50:
                return {"success": False, "error": "无法从文件中提取文本内容"}
            
            return self._process_text(text)
            
        except Exception as e:
            return {"success": False, "error": f"文件解析失败: {str(e)}"}
    
    def _process_text(self, text: str) -> Dict:
        """处理提取的文本"""
        project_info = self._extract_project_info(text)
        sections = self._extract_sections_from_text(text)
        required_materials = self._analyze_required_materials(text)
        sections = self._merge_sections_with_materials(sections, required_materials)
        
        return {
            "success": True,
            "project_info": project_info,
            "sections": sections,
            "required_materials": required_materials,
            "text_preview": text[:2000],
            "paragraph_count": 0,
            "table_count": 0
        }
    
    def analyze_with_lexitool(self, file_path: str) -> Dict:
        """使用 lexitool 分析招标文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            分析结果
        """
        try:
            import sys
            sys.path.insert(0, '/opt/lexitool')
            from lexitool.api import stats, doc_stats
            
            # 获取文档统计信息
            doc_info = stats(file_path)
            
            # 使用 python-docx 提取内容
            return self.analyze_with_docx(file_path)
            
        except Exception as e:
            logger.warning(f"lexitool 分析失败，回退到 python-docx: {e}")
            return self.analyze_with_docx(file_path)
    
    def _extract_project_info(self, text: str) -> Dict:
        """从文本中提取项目信息"""
        info = {}
        
        for field, patterns in self.FIELD_PATTERNS.items():
            for pattern in patterns:
                match = re.search(pattern, text)
                if match:
                    value = match.group(1).strip()
                    # 清理值
                    value = value.rstrip("。．.")
                    if value:
                        info[field] = value
                        break
        
        # 解析日期
        if "deadline" in info:
            try:
                deadline_str = info["deadline"]
                # 尝试多种日期格式
                for fmt in ["%Y年%m月%d日", "%Y-%m-%d", "%Y/%m/%d"]:
                    try:
                        info["deadline_parsed"] = datetime.strptime(deadline_str, fmt).isoformat()
                        break
                    except:
                        continue
            except:
                pass
        
        return info
    
    def _extract_sections(self, doc) -> List[Dict]:
        """从文档中提取章节结构"""
        sections = []
        seen_titles = set()
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检查是否是标题（通过样式或格式判断）
            is_heading = False
            heading_level = 0
            
            # 通过样式判断
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if 'heading' in style_name or '标题' in style_name:
                    is_heading = True
                    # 提取级别
                    level_match = re.search(r'(\d+)', style_name)
                    if level_match:
                        heading_level = int(level_match.group(1))
            
            # 通过格式判断（加粗、字号大等）
            if not is_heading:
                for run in para.runs:
                    if run.bold and len(text) < 50:
                        is_heading = True
                        break
            
            # 通过编号模式判断
            if not is_heading:
                # 匹配 "一、" "1." "第一章" 等模式
                if re.match(r'^[一二三四五六七八九十]+[、．.]', text) or \
                   re.match(r'^第[一二三四五六七八九十百]+[章节部分]', text) or \
                   re.match(r'^\d+[.、．]', text):
                    is_heading = True
            
            if is_heading and len(text) < 80 and text not in seen_titles:
                seen_titles.add(text)
                
                # 判断章节类型
                section_type = self._guess_section_type(text)
                
                sections.append({
                    "title": text,
                    "level": heading_level,
                    "section_type": section_type,
                    "order": len(sections) + 1
                })
        
        return sections
    
    def _guess_section_type(self, title: str) -> str:
        """根据标题猜测章节类型"""
        title_lower = title.lower()
        
        for keyword, section_type in self.SECTION_TYPE_MAP.items():
            if keyword in title_lower:
                return section_type
        
        return "other"
    
    def _analyze_required_materials(self, text: str) -> List[Dict]:
        """分析招标文件中要求提交的材料
        
        Args:
            text: 文档全文
        
        Returns:
            需要提交的材料列表
        """
        text_lower = text.lower()
        found_materials = []
        
        for material_name, config in self.REQUIRED_MATERIALS.items():
            matched_keywords = []
            for keyword in config["keywords"]:
                if keyword in text_lower:
                    matched_keywords.append(keyword)
            
            if matched_keywords:
                found_materials.append({
                    "name": material_name,
                    "section_type": config["section_type"],
                    "description": config["description"],
                    "matched_keywords": matched_keywords,
                    "confidence": min(0.5 + len(matched_keywords) * 0.15, 0.95)
                })
        
        # 按置信度排序
        found_materials.sort(key=lambda x: x["confidence"], reverse=True)
        return found_materials
    
    def _merge_sections_with_materials(self, sections: List[Dict], materials: List[Dict]) -> List[Dict]:
        """将章节结构与材料要求合并
        
        如果文档中已有明确的章节标题，保留原标题；
        如果材料要求在文档中没有对应的章节，自动补充。
        """
        # 已有的章节类型
        existing_types = set(s.get("section_type") for s in sections)
        
        # 补充文档中没有但材料要求中有的章节
        for material in materials:
            if material["section_type"] not in existing_types:
                sections.append({
                    "title": material["name"],
                    "level": 0,
                    "section_type": material["section_type"],
                    "order": len(sections) + 1,
                    "is_required": True,
                    "requirement": material["description"]
                })
        
        # 标记已有章节中哪些是必须的
        required_types = set(m["section_type"] for m in materials)
        for section in sections:
            if section.get("section_type") in required_types:
                section["is_required"] = True
        
        return sections
    
    def analyze_pdf(self, file_path: str) -> Dict:
        """分析 PDF 格式的招标文件"""
        try:
            import fitz
            
            doc = fitz.open(file_path)
            full_text = []
            
            for page in doc:
                full_text.append(page.get_text())
            
            text = "\n".join(full_text)
            doc.close()
            
            # 提取项目信息
            project_info = self._extract_project_info(text)
            
            # PDF 无法直接提取章节结构，尝试从文本中识别
            sections = self._extract_sections_from_text(text)
            
            # 分析需要提交的材料
            required_materials = self._analyze_required_materials(text)
            
            # 合并章节和材料要求
            sections = self._merge_sections_with_materials(sections, required_materials)
            
            return {
                "success": True,
                "project_info": project_info,
                "sections": sections,
                "required_materials": required_materials,
                "text_preview": text[:2000],
                "page_count": len(doc) if doc else 0
            }
            
        except Exception as e:
            logger.error(f"分析 PDF 招标文件失败: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_sections_from_text(self, text: str) -> List[Dict]:
        """从纯文本中提取章节结构"""
        sections = []
        seen_titles = set()
        
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line or len(line) > 80:
                continue
            
            # 匹配章节标题模式
            is_section = False
            section_type = "other"
            
            # 一、二、三、... 格式
            if re.match(r'^[一二三四五六七八九十]+[、．.]', line):
                is_section = True
            # 第一章/节/部分 格式
            elif re.match(r'^第[一二三四五六七八九十百]+[章节部分]', line):
                is_section = True
            # 1. 2. 3. 格式（且内容较短）
            elif re.match(r'^\d+[.、．]\s*\S', line) and len(line) < 30:
                is_section = True
            
            if is_section and line not in seen_titles:
                seen_titles.add(line)
                section_type = self._guess_section_type(line)
                sections.append({
                    "title": line,
                    "level": 0,
                    "section_type": section_type,
                    "order": len(sections) + 1
                })
        
        return sections


# 全局实例
tender_analyzer = TenderAnalyzer()
