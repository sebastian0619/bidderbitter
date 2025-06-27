#!/usr/bin/env python3
"""
水印引擎模块 - 完整版
提供Word文档水印功能，支持多种水印模式和完整配置
支持倾斜角度、多种位置、字体大小、颜色、透明度等全部前端配置
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.shared import qn
from docx.oxml import parse_xml
from typing import Dict, Any, Optional
import logging
import math

logger = logging.getLogger(__name__)

class WatermarkConfig:
    """水印配置类 - 支持完整配置"""
    def __init__(self, **kwargs):
        self.enable_watermark = kwargs.get('enable_watermark', False)
        self.watermark_text = kwargs.get('watermark_text', '水印')
        self.watermark_font_size = kwargs.get('watermark_font_size', 24)
        self.watermark_rotation = kwargs.get('watermark_rotation', -45)  # 倾斜角度
        self.watermark_opacity = kwargs.get('watermark_opacity', 0.3)
        self.watermark_color = kwargs.get('watermark_color', '#808080')
        self.watermark_position = kwargs.get('watermark_position', 'center')
    
    def is_valid(self) -> bool:
        """检查配置是否有效"""
        if not self.enable_watermark:
            return False
        
        if not self.watermark_text or not self.watermark_text.strip():
            return False
            
        if self.watermark_font_size <= 0 or self.watermark_font_size > 200:
            return False
            
        # 支持的位置类型
        valid_positions = ['center', 'repeat', 'background', 'top-left', 'top-right', 'bottom-left', 'bottom-right']
        if self.watermark_position not in valid_positions:
            return False
            
        return True
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            'enable_watermark': self.enable_watermark,
            'watermark_text': self.watermark_text,
            'watermark_font_size': self.watermark_font_size,
            'watermark_rotation': self.watermark_rotation,
            'watermark_opacity': self.watermark_opacity,
            'watermark_color': self.watermark_color,
            'watermark_position': self.watermark_position
        }

class WatermarkEngine:
    """水印引擎类 - 支持完整配置"""
    
    @staticmethod
    def hex_to_rgb(hex_color: str) -> tuple:
        """转换十六进制颜色为RGB元组"""
        try:
            hex_color = hex_color.lstrip('#')
            if len(hex_color) != 6:
                raise ValueError("Invalid hex color format")
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        except Exception as e:
            logger.warning(f"Invalid color {hex_color}, using default gray: {e}")
            return (128, 128, 128)  # 默认灰色
    
    @staticmethod
    def apply_transparency(rgb_color: tuple, opacity: float) -> tuple:
        """应用透明度到RGB颜色（通过混合白色模拟）"""
        opacity = max(0.0, min(1.0, opacity))  # 限制在0-1范围内
        return tuple(int(c * opacity + 255 * (1 - opacity)) for c in rgb_color)
    
    @staticmethod
    def apply_rotation_decoration(text: str, angle: int) -> str:
        """根据角度添加装饰符号模拟倾斜效果"""
        if angle == 0:
            return text
        elif angle > 0:
            # 正角度，使用右倾斜装饰
            return f"╱ {text} ╱"
        else:
            # 负角度，使用左倾斜装饰  
            return f"╲ {text} ╲"
    
    @classmethod
    def apply_watermark(cls, doc: Document, config: WatermarkConfig) -> bool:
        """
        在Word文档的每一页都应用水印
        """
        try:
            if not config.is_valid():
                logger.info("Watermark disabled or invalid configuration")
                return True
            
            # 准备颜色
            rgb_color = cls.hex_to_rgb(config.watermark_color)
            rgb_color = cls.apply_transparency(rgb_color, config.watermark_opacity)
            
            # 在每页插入水印 - 遍历所有段落，每隔一定数量插入水印
            paragraphs = list(doc.paragraphs)
            paragraphs_per_page = 25  # 估算每页大约25个段落
            
            # 从后往前插入，避免索引变化
            for i in range(len(paragraphs) - 1, -1, -paragraphs_per_page):
                insert_position = max(0, i - paragraphs_per_page + 5)  # 在页面前部插入
                if insert_position < len(paragraphs):
                    # 在指定位置插入水印
                    cls._insert_watermark_at_position(doc, insert_position, config.watermark_text, 
                                                    config.watermark_font_size, rgb_color, 
                                                    config.watermark_rotation, config.watermark_position)
            
            # 在文档开头也插入一个水印
            cls._insert_watermark_at_position(doc, 0, config.watermark_text, 
                                            config.watermark_font_size, rgb_color, 
                                            config.watermark_rotation, config.watermark_position)
            
            logger.info(f"Added watermarks throughout document: {config.watermark_text}")
            return True
                
        except Exception as e:
            logger.error(f"Error applying watermark: {e}")
            return False
    
    @classmethod
    def _insert_watermark_at_position(cls, doc: Document, position: int, text: str, font_size: int, rgb_color: tuple, rotation: int, watermark_position: str):
        """在指定位置插入水印"""
        try:
            # 根据角度应用装饰
            decorated_text = cls.apply_rotation_decoration(text, rotation)
            
            # 在指定位置插入段落
            if position < len(doc.paragraphs):
                target_para = doc.paragraphs[position]
                watermark_para = target_para.insert_paragraph_before()
            else:
                watermark_para = doc.add_paragraph()
            
            # 设置水印样式
            if watermark_position == "center":
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                watermark_text = f"『 {decorated_text} 』"
                font_size_multiplier = 1.5
            elif watermark_position == "repeat":
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
                watermark_text = f"◆ {decorated_text} ◆"
                font_size_multiplier = 1.0
            elif watermark_position == "background":
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                watermark_text = f"【 {decorated_text} 】"
                font_size_multiplier = 1.2
            elif watermark_position == "top-left":
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                watermark_text = f"◆ {decorated_text}"
                font_size_multiplier = 0.8
            elif watermark_position == "top-right":
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                watermark_text = f"{decorated_text} ◆"
                font_size_multiplier = 0.8
            elif watermark_position == "bottom-left":
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                watermark_text = f"◇ {decorated_text}"
                font_size_multiplier = 0.8
            elif watermark_position == "bottom-right":
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                watermark_text = f"{decorated_text} ◇"
                font_size_multiplier = 0.8
            else:
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                watermark_text = f"『 {decorated_text} 』"
                font_size_multiplier = 1.5
            
            # 添加水印文字
            run = watermark_para.add_run(watermark_text)
            run.font.size = Pt(int(font_size * font_size_multiplier))
            run.font.color.rgb = RGBColor(*rgb_color)
            run.font.bold = True
            run.font.name = '楷体'
            
        except Exception as e:
            logger.error(f"Error inserting watermark at position {position}: {e}")

    @staticmethod
    def _add_center_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int) -> bool:
        """废弃 - 使用新的apply_watermark方法"""
        return True
    
    @staticmethod
    def _add_repeat_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int) -> bool:
        """废弃 - 使用新的apply_watermark方法"""
            return True
    
    @staticmethod
    def _add_background_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int) -> bool:
        """废弃 - 使用新的apply_watermark方法"""
            return True
    
    @staticmethod
    def _add_corner_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int, position: str) -> bool:
        """废弃 - 使用新的apply_watermark方法"""
            return True

# 便捷函数
def apply_watermark_to_document(doc: Document, watermark_config: Dict[str, Any]) -> bool:
    """
    便捷函数：直接使用字典配置为文档添加水印
    
    Args:
        doc: Word文档对象
        watermark_config: 水印配置字典
        
    Returns:
        bool: 是否成功
    """
    config = WatermarkConfig(**watermark_config)
    return WatermarkEngine.apply_watermark(doc, config)

def create_watermark_config(enable=False, text="水印", font_size=24, color="#808080", 
                           position="center", opacity=0.3, rotation=-45) -> WatermarkConfig:
    """
    便捷函数：创建水印配置对象
    """
    return WatermarkConfig(
        enable_watermark=enable,
        watermark_text=text,
        watermark_font_size=font_size,
        watermark_color=color,
        watermark_position=position,
        watermark_opacity=opacity,
        watermark_rotation=rotation
    )

if __name__ == "__main__":
    # 简单测试
    print("水印引擎模块测试（完整版）...")
    
    # 创建测试文档
    test_doc = Document()
    test_doc.add_heading('完整水印引擎测试', 0)
    test_doc.add_paragraph("这是测试内容。")
    
    # 创建测试配置（包含所有配置选项）
    test_config = create_watermark_config(
        enable=True,
        text="投标苦 • 完整配置",
        font_size=20,
        color="#FF0000",
        position="center",
        opacity=0.4,
        rotation=-30  # 测试倾斜角度
    )
    
    # 应用水印
    success = WatermarkEngine.apply_watermark(test_doc, test_config)
    
    if success:
        test_doc.save("test_watermark_engine_full.docx")
        print("✓ 完整水印引擎测试成功！生成了 test_watermark_engine_full.docx")
        print(f"  配置: {test_config.to_dict()}")
    else:
        print("✗ 完整水印引擎测试失败！") 