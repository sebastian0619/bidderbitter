import os
import logging
from typing import Dict, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import fitz  # PyMuPDF for PDF to image conversion
from io import BytesIO

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        self.output_dir = "/app/generated_docs"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_award_document(self, awards: List[Dict], selected_filters: Dict) -> str:
        """生成获奖文档"""
        try:
            # 创建新文档
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.18)
                section.right_margin = Cm(3.18)
            
            # 添加标题
            title = doc.add_heading('获奖情况', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加筛选条件说明
            self._add_filter_summary(doc, selected_filters)
            
            # 按年份和厂牌分组显示获奖信息
            grouped_awards = self._group_awards(awards)
            
            for year in sorted(grouped_awards.keys(), reverse=True):
                # 年份标题
                year_heading = doc.add_heading(f'{year}年获奖情况', 1)
                year_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for brand in grouped_awards[year]:
                    # 厂牌子标题
                    brand_heading = doc.add_heading(f'{brand}', 2)
                    
                    for award in grouped_awards[year][brand]:
                        # 添加获奖信息
                        self._add_award_content(doc, award)
                        
                        # 添加页面分隔符（每个奖项后）
                        doc.add_page_break()
            
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"awards_report_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # 保存文档
            doc.save(filepath)
            
            logger.info(f"获奖文档生成成功: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"生成获奖文档失败: {str(e)}")
            raise e
    
    def generate_performance_document(self, performances: List[Dict], selected_filters: Dict) -> str:
        """生成业绩文档"""
        try:
            # 创建新文档
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.18)
                section.right_margin = Cm(3.18)
            
            # 添加标题
            title = doc.add_heading('业绩情况', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加筛选条件说明
            self._add_filter_summary(doc, selected_filters)
            
            # 添加业绩汇总表
            self._add_performance_summary_table(doc, performances)
            
            # 按年份分组显示详细业绩
            grouped_performances = self._group_performances_by_year(performances)
            
            for year in sorted(grouped_performances.keys(), reverse=True):
                # 年份标题
                year_heading = doc.add_heading(f'{year}年业绩情况', 1)
                
                for performance in grouped_performances[year]:
                    # 添加业绩详情
                    self._add_performance_content(doc, performance)
            
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"performance_report_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # 保存文档
            doc.save(filepath)
            
            logger.info(f"业绩文档生成成功: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"生成业绩文档失败: {str(e)}")
            raise e
    
    def generate_combined_document(self, awards: List[Dict], performances: List[Dict], selected_filters: Dict) -> str:
        """生成综合文档（获奖+业绩）"""
        try:
            # 创建新文档
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.18)
                section.right_margin = Cm(3.18)
            
            # 添加总标题
            title = doc.add_heading('投标资料汇总', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加筛选条件说明
            self._add_filter_summary(doc, selected_filters)
            
            # 第一部分：获奖情况
            if awards:
                doc.add_heading('第一部分 获奖情况', 1)
                grouped_awards = self._group_awards(awards)
                
                for year in sorted(grouped_awards.keys(), reverse=True):
                    year_heading = doc.add_heading(f'{year}年获奖情况', 2)
                    
                    for brand in grouped_awards[year]:
                        brand_heading = doc.add_heading(f'{brand}', 3)
                        
                        for award in grouped_awards[year][brand]:
                            self._add_award_content(doc, award)
                            doc.add_page_break()
            
            # 第二部分：业绩情况
            if performances:
                doc.add_heading('第二部分 业绩情况', 1)
                
                # 业绩汇总表
                self._add_performance_summary_table(doc, performances)
                
                # 详细业绩
                grouped_performances = self._group_performances_by_year(performances)
                
                for year in sorted(grouped_performances.keys(), reverse=True):
                    year_heading = doc.add_heading(f'{year}年业绩情况', 2)
                    
                    for performance in grouped_performances[year]:
                        self._add_performance_content(doc, performance)
            
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"combined_report_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # 保存文档
            doc.save(filepath)
            
            logger.info(f"综合文档生成成功: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"生成综合文档失败: {str(e)}")
            raise e
    
    def _add_filter_summary(self, doc: Document, filters: Dict):
        """添加筛选条件汇总"""
        if not filters:
            return
        
        doc.add_heading('筛选条件', 2)
        
        filter_info = []
        if filters.get('brands'):
            filter_info.append(f"厂牌：{', '.join(filters['brands'])}")
        if filters.get('business_fields'):
            filter_info.append(f"业务领域：{', '.join(filters['business_fields'])}")
        if filters.get('years'):
            filter_info.append(f"年份：{', '.join(map(str, filters['years']))}")
        
        if filter_info:
            para = doc.add_paragraph('\n'.join(filter_info))
            para.style = 'List Bullet'
        
        doc.add_paragraph()  # 空行
    
    def _group_awards(self, awards: List[Dict]) -> Dict:
        """按年份和厂牌分组获奖信息"""
        grouped = {}
        
        for award in awards:
            year = award.get('year', 0)
            brand = award.get('brand', '未知厂牌')
            
            if year not in grouped:
                grouped[year] = {}
            if brand not in grouped[year]:
                grouped[year][brand] = []
            
            grouped[year][brand].append(award)
        
        return grouped
    
    def _group_performances_by_year(self, performances: List[Dict]) -> Dict:
        """按年份分组业绩信息"""
        grouped = {}
        
        for performance in performances:
            year = performance.get('year', 0)
            
            if year not in grouped:
                grouped[year] = []
            
            grouped[year].append(performance)
        
        return grouped
    
    def _add_award_content(self, doc: Document, award: Dict):
        """添加单个获奖内容"""
        # 获奖标题
        title_text = f"{award.get('brand', '')} - {award.get('year', '')}年度 - {award.get('title', '')}"
        award_title = doc.add_paragraph(title_text)
        award_title.style = 'Heading 3'
        
        # 获奖描述
        if award.get('description'):
            doc.add_paragraph(award['description'])
        
        # 业务类型
        if award.get('business_type'):
            business_para = doc.add_paragraph(f"业务类型：{award['business_type']}")
            business_para.style = 'List Bullet'
        
        # 添加截图
        self._add_award_screenshots(doc, award)
    
    def _add_award_screenshots(self, doc: Document, award: Dict):
        """添加获奖截图"""
        screenshots = award.get('screenshot_pages', [])
        
        if not screenshots:
            return
        
        for i, screenshot_info in enumerate(screenshots):
            screenshot_path = screenshot_info.get('path')
            
            if screenshot_path and os.path.exists(screenshot_path):
                try:
                    # 优化图片尺寸
                    optimized_path = self._optimize_image_for_word(screenshot_path)
                    
                    # 添加图片
                    paragraph = doc.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(optimized_path, width=Inches(6.5))  # A4纸宽度约6.5英寸
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 如果有多页，添加页面分隔符
                    if i < len(screenshots) - 1:
                        doc.add_page_break()
                        
                except Exception as e:
                    logger.error(f"添加截图失败: {str(e)}")
                    doc.add_paragraph(f"截图加载失败: {screenshot_path}")
    
    def _add_performance_summary_table(self, doc: Document, performances: List[Dict]):
        """添加业绩汇总表"""
        if not performances:
            return
        
        doc.add_heading('业绩汇总表', 2)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        headers = ['序号', '客户名称', '项目名称', '项目类型', '业务领域', '年份', '合同金额']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 设置表头格式
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加数据行
        for i, performance in enumerate(performances, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = performance.get('client_name', '')
            row_cells[2].text = performance.get('project_name', '')
            row_cells[3].text = performance.get('project_type', '')
            row_cells[4].text = performance.get('business_field', '')
            row_cells[5].text = str(performance.get('year', ''))
            
            # 合同金额格式化
            amount = performance.get('contract_amount')
            currency = performance.get('currency', 'CNY')
            if amount:
                row_cells[6].text = f"{amount:,.2f} {currency}"
            else:
                row_cells[6].text = '未填写'
        
        doc.add_paragraph()  # 空行
    
    def _add_performance_content(self, doc: Document, performance: Dict):
        """添加单个业绩内容"""
        # 业绩标题
        title_text = f"{performance.get('client_name', '')} - {performance.get('project_name', '')}"
        perf_title = doc.add_paragraph(title_text)
        perf_title.style = 'Heading 3'
        
        # 业绩基本信息
        info_items = []
        if performance.get('project_type'):
            info_items.append(f"项目类型：{performance['project_type']}")
        if performance.get('business_field'):
            info_items.append(f"业务领域：{performance['business_field']}")
        if performance.get('start_date') and performance.get('end_date'):
            info_items.append(f"服务期间：{performance['start_date']} 至 {performance['end_date']}")
        if performance.get('contract_amount'):
            currency = performance.get('currency', 'CNY')
            info_items.append(f"合同金额：{performance['contract_amount']:,.2f} {currency}")
        
        for item in info_items:
            para = doc.add_paragraph(item)
            para.style = 'List Bullet'
        
        # 项目描述
        if performance.get('description'):
            doc.add_paragraph(performance['description'])
        
        # 添加合同文件
        self._add_performance_documents(doc, performance)
        
        doc.add_page_break()
    
    def _add_performance_documents(self, doc: Document, performance: Dict):
        """添加业绩相关文档"""
        files = performance.get('files', [])
        
        for file_info in files:
            file_path = file_info.get('file_path')
            file_type = file_info.get('file_type')
            
            if not file_path or not os.path.exists(file_path):
                continue
            
            if file_type == 'contract' and file_path.lower().endswith('.pdf'):
                # PDF转图片插入
                self._add_pdf_as_images(doc, file_path)
            elif file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                # 直接插入图片
                self._add_image_to_doc(doc, file_path)
    
    def _add_pdf_as_images(self, doc: Document, pdf_path: str):
        """将PDF转换为图片并插入文档"""
        try:
            pdf_document = fitz.open(pdf_path)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # 渲染页面为图片
                mat = fitz.Matrix(2.0, 2.0)  # 提高分辨率
                pix = page.get_pixmap(matrix=mat)
                
                # 转换为PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data))
                
                # 保存临时图片
                temp_path = f"/tmp/pdf_page_{page_num}.png"
                img.save(temp_path)
                
                # 插入到文档
                self._add_image_to_doc(doc, temp_path)
                
                # 删除临时文件
                os.remove(temp_path)
                
                # 如果不是最后一页，添加分页符
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            
            pdf_document.close()
            
        except Exception as e:
            logger.error(f"PDF转图片失败: {str(e)}")
            doc.add_paragraph(f"PDF文档处理失败: {pdf_path}")
    
    def _add_image_to_doc(self, doc: Document, image_path: str):
        """向文档添加图片"""
        try:
            # 优化图片
            optimized_path = self._optimize_image_for_word(image_path)
            
            paragraph = doc.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(optimized_path, width=Inches(6.5))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            logger.error(f"添加图片失败: {str(e)}")
            doc.add_paragraph(f"图片加载失败: {image_path}")
    
    def _optimize_image_for_word(self, image_path: str) -> str:
        """优化图片以适应Word文档"""
        try:
            with Image.open(image_path) as img:
                # A4纸宽度约为1654像素 (6.5英寸 * 254 DPI)
                max_width = 1654
                
                if img.width > max_width:
                    # 计算缩放比例
                    scale_ratio = max_width / img.width
                    new_height = int(img.height * scale_ratio)
                    
                    # 重新调整大小
                    resized_img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                    
                    # 生成优化后的文件名
                    optimized_path = image_path.replace('.png', '_word_optimized.png')
                    resized_img.save(optimized_path, 'PNG', optimize=True, quality=85)
                    
                    return optimized_path
                else:
                    return image_path
                    
        except Exception as e:
            logger.error(f"图片优化失败: {str(e)}")
            return image_path

# 创建全局文档生成器实例
document_generator = DocumentGenerator() 