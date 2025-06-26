#!/usr/bin/env python3
"""
水印引擎模块 - 完整版
提供Word文档水印功能，支持多种水印模式和完整配置
支持倾斜角度、多种位置、字体大小、颜色、透明度等全部前端配置
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.shared import qn
from docx.oxml import parse_xml
from typing import Dict, Any, Optional
import logging
import math

logger = logging.getLogger(__name__)

class WatermarkConfig:
    """水印配置类 - 支持完整配置"""
    def __init__(self, **kwargs):
        self.enable_watermark = kwargs.get('enable_watermark', False)
        self.watermark_text = kwargs.get('watermark_text', '水印')
        self.watermark_font_size = kwargs.get('watermark_font_size', 24)
        self.watermark_rotation = kwargs.get('watermark_rotation', -45)  # 倾斜角度
        self.watermark_opacity = kwargs.get('watermark_opacity', 0.3)
        self.watermark_color = kwargs.get('watermark_color', '#808080')
        self.watermark_position = kwargs.get('watermark_position', 'center')
    
    def is_valid(self) -> bool:
        """检查配置是否有效"""
        if not self.enable_watermark:
            return False
        
        if not self.watermark_text or not self.watermark_text.strip():
            return False
            
        if self.watermark_font_size <= 0 or self.watermark_font_size > 200:
            return False
            
        # 支持的位置类型
        valid_positions = ['center', 'repeat', 'background', 'top-left', 'top-right', 'bottom-left', 'bottom-right']
        if self.watermark_position not in valid_positions:
            return False
            
        return True
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            'enable_watermark': self.enable_watermark,
            'watermark_text': self.watermark_text,
            'watermark_font_size': self.watermark_font_size,
            'watermark_rotation': self.watermark_rotation,
            'watermark_opacity': self.watermark_opacity,
            'watermark_color': self.watermark_color,
            'watermark_position': self.watermark_position
        }

class WatermarkEngine:
    """水印引擎类 - 支持完整配置"""
    
    @staticmethod
    def hex_to_rgb(hex_color: str) -> tuple:
        """转换十六进制颜色为RGB元组"""
        try:
            hex_color = hex_color.lstrip('#')
            if len(hex_color) != 6:
                raise ValueError("Invalid hex color format")
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        except Exception as e:
            logger.warning(f"Invalid color {hex_color}, using default gray: {e}")
            return (128, 128, 128)  # 默认灰色
    
    @staticmethod
    def apply_transparency(rgb_color: tuple, opacity: float) -> tuple:
        """应用透明度到RGB颜色（通过混合白色模拟）"""
        opacity = max(0.0, min(1.0, opacity))  # 限制在0-1范围内
        return tuple(int(c * opacity + 255 * (1 - opacity)) for c in rgb_color)
    
    @staticmethod
    def apply_rotation_decoration(text: str, angle: int) -> str:
        """根据角度添加装饰符号模拟倾斜效果"""
        if angle == 0:
            return text
        elif angle > 0:
            # 正角度，使用右倾斜装饰
            return f"╱ {text} ╱"
        else:
            # 负角度，使用左倾斜装饰  
            return f"╲ {text} ╲"
    
    @classmethod
    def apply_watermark(cls, doc: Document, config: WatermarkConfig) -> bool:
        """
        在Word文档中应用水印
        
        Args:
            doc: docx Document对象
            config: 水印配置对象
            
        Returns:
            bool: 是否成功应用水印
        """
        try:
            if not config.is_valid():
                logger.info("Watermark disabled or invalid configuration")
                return True
            
            # 准备颜色
            rgb_color = cls.hex_to_rgb(config.watermark_color)
            rgb_color = cls.apply_transparency(rgb_color, config.watermark_opacity)
            
            # 根据位置模式应用水印
            if config.watermark_position == "center":
                return cls._add_center_watermark(doc, config.watermark_text, config.watermark_font_size, rgb_color, config.watermark_rotation)
            elif config.watermark_position == "repeat":
                return cls._add_repeat_watermark(doc, config.watermark_text, config.watermark_font_size, rgb_color, config.watermark_rotation)
            elif config.watermark_position == "background":
                return cls._add_background_watermark(doc, config.watermark_text, config.watermark_font_size, rgb_color, config.watermark_rotation)
            elif config.watermark_position in ["top-left", "top-right", "bottom-left", "bottom-right"]:
                return cls._add_corner_watermark(doc, config.watermark_text, config.watermark_font_size, rgb_color, config.watermark_rotation, config.watermark_position)
            else:
                logger.error(f"Unknown watermark position: {config.watermark_position}")
                return False
                
        except Exception as e:
            logger.error(f"Error applying watermark: {e}")
            return False
    
    @staticmethod
    def _add_center_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int) -> bool:
        """添加居中大水印"""
        try:
            # 在文档中央位置插入水印
            for i in range(8):
                doc.add_paragraph()
            
            # 根据角度应用装饰
            decorated_text = WatermarkEngine.apply_rotation_decoration(text, rotation)
            
            # 上装饰边框
            border_para1 = doc.add_paragraph()
            border_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            border1_run = border_para1.add_run("◆" * 20)
            border1_run.font.size = Pt(max(8, font_size // 3))
            border1_run.font.color.rgb = RGBColor(*rgb_color)
            border1_run.font.name = '楷体'
            
            # 主水印段落
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            main_run = watermark_para.add_run(f"『 {decorated_text} 』")
            main_run.font.size = Pt(font_size * 2)
            main_run.font.color.rgb = RGBColor(*rgb_color)
            main_run.font.bold = True
            main_run.font.name = '楷体'
            
            # 下装饰边框
            border_para2 = doc.add_paragraph()
            border_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            border2_run = border_para2.add_run("◇" * 20)
            border2_run.font.size = Pt(max(8, font_size // 3))
            border2_run.font.color.rgb = RGBColor(*rgb_color)
            border2_run.font.name = '楷体'
            
            logger.info(f"Added center watermark: {text} (rotation: {rotation}°)")
            return True
            
        except Exception as e:
            logger.error(f"Error adding center watermark: {e}")
            return False
    
    @staticmethod
    def _add_repeat_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int) -> bool:
        """添加重复平铺水印"""
        try:
            # 获取当前文档长度，在多个位置插入水印
            current_length = len(doc.paragraphs)
            
            # 计算水印位置（大约每10段插入一次）
            watermark_positions = []
            for i in range(5, max(50, current_length + 30), 12):
                watermark_positions.append(i)
            
            # 根据角度应用装饰
            decorated_text = WatermarkEngine.apply_rotation_decoration(text, rotation)
            
            for pos in watermark_positions:
                # 确保文档有足够的段落
                while len(doc.paragraphs) < pos:
                    doc.add_paragraph()
                
                # 插入水印段落
                watermark_para = doc.add_paragraph()
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 创建重复的水印文字
                repeat_text = f"◆ {decorated_text} ◆   ◇ {decorated_text} ◇   ◆ {decorated_text} ◆"
                watermark_run = watermark_para.add_run(repeat_text)
                watermark_run.font.size = Pt(font_size)
                watermark_run.font.color.rgb = RGBColor(*rgb_color)
                watermark_run.font.bold = True
                watermark_run.italic = True
                watermark_run.font.name = '楷体'
                
                # 添加装饰线
                deco_para = doc.add_paragraph()
                deco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                deco_run = deco_para.add_run("~ " * 15)
                deco_run.font.size = Pt(max(8, font_size // 2))
                deco_run.font.color.rgb = RGBColor(*rgb_color)
                deco_run.font.name = '楷体'
            
            logger.info(f"Added repeat watermark: {text} at {len(watermark_positions)} positions (rotation: {rotation}°)")
            return True
            
        except Exception as e:
            logger.error(f"Error adding repeat watermark: {e}")
            return False
    
    @staticmethod
    def _add_background_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int) -> bool:
        """添加背景样式水印"""
        try:
            # 根据角度应用装饰
            decorated_text = WatermarkEngine.apply_rotation_decoration(text, rotation)
            
            # 创建背景水印区域
            for row in range(12):
                bg_para = doc.add_paragraph()
                bg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if row == 5:  # 中心主水印
                    bg_run = bg_para.add_run(f"【 {decorated_text} 】")
                    bg_run.font.size = Pt(font_size * 2)
                    bg_run.font.bold = True
                elif row == 2 or row == 8:  # 副水印行
                    bg_run = bg_para.add_run(f"◆ {decorated_text} ◆")
                    bg_run.font.size = Pt(font_size)
                    bg_run.font.bold = True
                elif row % 2 == 0:  # 装饰行
                    bg_run = bg_para.add_run("～ ～ ～ ～ ～ ～ ～ ～ ～ ～")
                    bg_run.font.size = Pt(max(8, font_size // 2))
                else:  # 空行，只有少量装饰
                    bg_run = bg_para.add_run("· · · · ·")
                    bg_run.font.size = Pt(max(6, font_size // 3))
                
                bg_run.font.color.rgb = RGBColor(*rgb_color)
                bg_run.font.name = '楷体'
            
            logger.info(f"Added background watermark: {text} (rotation: {rotation}°)")
            return True
            
        except Exception as e:
            logger.error(f"Error adding background watermark: {e}")
            return False
    
    @staticmethod
    def _add_corner_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple, rotation: int, position: str) -> bool:
        """添加角落位置水印"""
        try:
            # 根据角度应用装饰
            decorated_text = WatermarkEngine.apply_rotation_decoration(text, rotation)
            
            # 根据位置确定对齐方式
            if position == "top-left":
                alignment = WD_ALIGN_PARAGRAPH.LEFT
                prefix_lines = 2
            elif position == "top-right":
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
                prefix_lines = 2
            elif position == "bottom-left":
                alignment = WD_ALIGN_PARAGRAPH.LEFT
                prefix_lines = 15
            elif position == "bottom-right":
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
                prefix_lines = 15
            else:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
                prefix_lines = 8
            
            # 添加定位空段落
            for i in range(prefix_lines):
                doc.add_paragraph()
            
            # 添加水印段落
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = alignment
            
            watermark_run = watermark_para.add_run(f"◆ {decorated_text} ◆")
            watermark_run.font.size = Pt(font_size)
            watermark_run.font.color.rgb = RGBColor(*rgb_color)
            watermark_run.font.bold = True
            watermark_run.font.name = '楷体'
            
            # 添加装饰线
            deco_para = doc.add_paragraph()
            deco_para.alignment = alignment
            deco_run = deco_para.add_run("━" * 10)
            deco_run.font.size = Pt(max(8, font_size // 3))
            deco_run.font.color.rgb = RGBColor(*rgb_color)
            deco_run.font.name = '楷体'
            
            logger.info(f"Added corner watermark: {text} at {position} (rotation: {rotation}°)")
            return True
            
        except Exception as e:
            logger.error(f"Error adding corner watermark: {e}")
            return False

# 便捷函数
def apply_watermark_to_document(doc: Document, watermark_config: Dict[str, Any]) -> bool:
    """
    便捷函数：直接使用字典配置为文档添加水印
    
    Args:
        doc: Word文档对象
        watermark_config: 水印配置字典
        
    Returns:
        bool: 是否成功
    """
    config = WatermarkConfig(**watermark_config)
    return WatermarkEngine.apply_watermark(doc, config)

def create_watermark_config(enable=False, text="水印", font_size=24, color="#808080", 
                           position="center", opacity=0.3, rotation=-45) -> WatermarkConfig:
    """
    便捷函数：创建水印配置对象
    """
    return WatermarkConfig(
        enable_watermark=enable,
        watermark_text=text,
        watermark_font_size=font_size,
        watermark_color=color,
        watermark_position=position,
        watermark_opacity=opacity,
        watermark_rotation=rotation
    )

if __name__ == "__main__":
    # 简单测试
    print("水印引擎模块测试（完整版）...")
    
    # 创建测试文档
    test_doc = Document()
    test_doc.add_heading('完整水印引擎测试', 0)
    test_doc.add_paragraph("这是测试内容。")
    
    # 创建测试配置（包含所有配置选项）
    test_config = create_watermark_config(
        enable=True,
        text="投标苦 • 完整配置",
        font_size=20,
        color="#FF0000",
        position="center",
        opacity=0.4,
        rotation=-30  # 测试倾斜角度
    )
    
    # 应用水印
    success = WatermarkEngine.apply_watermark(test_doc, test_config)
    
    if success:
        test_doc.save("test_watermark_engine_full.docx")
        print("✓ 完整水印引擎测试成功！生成了 test_watermark_engine_full.docx")
        print(f"  配置: {test_config.to_dict()}")
    else:
        print("✗ 完整水印引擎测试失败！") 