#!/usr/bin/env python3
"""
高级水印功能测试脚本
模拟正式代码中的水印配置对象和处理逻辑
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import os
import json
from typing import Dict, Any

class WatermarkConfig:
    """水印配置类，模拟前端传来的配置"""
    def __init__(self, **kwargs):
        self.enable_watermark = kwargs.get('enable_watermark', False)
        self.watermark_text = kwargs.get('watermark_text', '水印')
        self.watermark_font_size = kwargs.get('watermark_font_size', 24)
        self.watermark_rotation = kwargs.get('watermark_rotation', -45)
        self.watermark_opacity = kwargs.get('watermark_opacity', 0.3)
        self.watermark_color = kwargs.get('watermark_color', '#808080')
        self.watermark_position = kwargs.get('watermark_position', 'center')
    
    def to_dict(self):
        return {
            'enable_watermark': self.enable_watermark,
            'watermark_text': self.watermark_text,
            'watermark_font_size': self.watermark_font_size,
            'watermark_rotation': self.watermark_rotation,
            'watermark_opacity': self.watermark_opacity,
            'watermark_color': self.watermark_color,
            'watermark_position': self.watermark_position
        }

def hex_to_rgb(hex_color: str) -> tuple:
    """转换十六进制颜色为RGB"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def apply_word_watermark(doc: Document, watermark_config: WatermarkConfig) -> None:
    """
    在Word文档中应用水印
    这个函数模拟正式代码中的水印处理逻辑
    """
    if not watermark_config.enable_watermark or not watermark_config.watermark_text:
        return
    
    text = watermark_config.watermark_text
    font_size = watermark_config.watermark_font_size
    color_hex = watermark_config.watermark_color
    position = watermark_config.watermark_position
    opacity = watermark_config.watermark_opacity
    
    # 转换颜色
    try:
        rgb_color = hex_to_rgb(color_hex)
        # 应用透明度（简化处理，调整RGB值）
        rgb_color = tuple(int(c * opacity + 255 * (1 - opacity)) for c in rgb_color)
    except:
        rgb_color = (128, 128, 128)  # 默认灰色
    
    if position == "center":
        # 居中大水印模式
        _add_center_watermark(doc, text, font_size, rgb_color)
    elif position == "repeat":
        # 平铺重复模式
        _add_repeat_watermark(doc, text, font_size, rgb_color)
    elif position == "background":
        # 背景模式
        _add_background_watermark(doc, text, font_size, rgb_color)

def _add_center_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple):
    """添加居中大水印"""
    # 在文档中央插入大号水印
    for i in range(10):  # 在文档中部位置
        doc.add_paragraph()
    
    # 创建水印段落
    watermark_para = doc.add_paragraph()
    watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 上装饰线
    border_para1 = doc.add_paragraph()
    border_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border1_run = border_para1.add_run("◆" * 15)
    border1_run.font.size = Pt(font_size // 3)
    border1_run.font.color.rgb = RGBColor(*rgb_color)
    
    # 主水印文字
    main_run = watermark_para.add_run(f"『{text}』")
    main_run.font.size = Pt(font_size * 2)
    main_run.font.color.rgb = RGBColor(*rgb_color)
    main_run.font.bold = True
    
    # 下装饰线
    border_para2 = doc.add_paragraph()
    border_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border2_run = border_para2.add_run("◇" * 15)
    border2_run.font.size = Pt(font_size // 3)
    border2_run.font.color.rgb = RGBColor(*rgb_color)

def _add_repeat_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple):
    """添加重复平铺水印"""
    # 在多个位置插入水印
    watermark_positions = [5, 15, 25, 35, 45]  # 不同的段落位置
    
    for pos in watermark_positions:
        # 确保文档有足够的段落
        while len(doc.paragraphs) < pos:
            doc.add_paragraph()
        
        # 在指定位置插入水印
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        watermark_run = watermark_para.add_run(f"◆ {text} ◆   ◇ {text} ◇   ◆ {text} ◆")
        watermark_run.font.size = Pt(font_size)
        watermark_run.font.color.rgb = RGBColor(*rgb_color)
        watermark_run.font.bold = True
        watermark_run.italic = True

def _add_background_watermark(doc: Document, text: str, font_size: int, rgb_color: tuple):
    """添加背景样式水印"""
    # 在文档开始处添加大面积背景
    original_paras = len(doc.paragraphs)
    
    # 创建背景水印层
    for row in range(8):
        bg_para = doc.add_paragraph()
        bg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if row == 3:  # 中间行放主要水印
            bg_run = bg_para.add_run(f"【 {text} 】")
            bg_run.font.size = Pt(font_size * 2)
            bg_run.font.bold = True
        elif row == 1 or row == 5:  # 副水印行
            bg_run = bg_para.add_run(f"◆ {text} ◆")
            bg_run.font.size = Pt(font_size)
        else:  # 装饰行
            bg_run = bg_para.add_run("～ ～ ～ ～ ～ ～ ～ ～ ～ ～")
            bg_run.font.size = Pt(font_size // 2)
        
        bg_run.font.color.rgb = RGBColor(*rgb_color)

def test_watermark_scenarios():
    """测试各种水印场景"""
    
    # 测试场景1：默认配置
    config1 = WatermarkConfig(
        enable_watermark=True,
        watermark_text="投标苦",
        watermark_font_size=24,
        watermark_color="#FF0000",
        watermark_position="center"
    )
    
    doc1 = Document()
    doc1.add_heading('场景1：默认居中红色水印', 0)
    doc1.add_paragraph("这是测试内容。水印应该显示在文档中央。")
    apply_word_watermark(doc1, config1)
    doc1.add_paragraph("这是水印后的内容。")
    
    output1 = "test_scenario1_default.docx"
    doc1.save(output1)
    print(f"场景1完成: {output1}")
    
    # 测试场景2：平铺模式
    config2 = WatermarkConfig(
        enable_watermark=True,
        watermark_text="CONFIDENTIAL",
        watermark_font_size=18,
        watermark_color="#0080FF",
        watermark_position="repeat",
        watermark_opacity=0.5
    )
    
    doc2 = Document()
    doc2.add_heading('场景2：平铺模式蓝色水印', 0)
    for i in range(50):  # 添加大量内容测试平铺效果
        doc2.add_paragraph(f"这是第 {i+1} 段测试内容。水印应该分布在多个位置。")
    
    apply_word_watermark(doc2, config2)
    
    output2 = "test_scenario2_repeat.docx"
    doc2.save(output2)
    print(f"场景2完成: {output2}")
    
    # 测试场景3：背景模式
    config3 = WatermarkConfig(
        enable_watermark=True,
        watermark_text="内部资料",
        watermark_font_size=20,
        watermark_color="#808080",
        watermark_position="background",
        watermark_opacity=0.2
    )
    
    doc3 = Document()
    doc3.add_heading('场景3：背景模式灰色水印', 0)
    apply_word_watermark(doc3, config3)
    
    for i in range(20):
        doc3.add_paragraph(f"正文内容第 {i+1} 段。水印应该作为背景显示。")
    
    output3 = "test_scenario3_background.docx"
    doc3.save(output3)
    print(f"场景3完成: {output3}")
    
    # 测试场景4：禁用水印
    config4 = WatermarkConfig(
        enable_watermark=False,
        watermark_text="不应该显示的水印"
    )
    
    doc4 = Document()
    doc4.add_heading('场景4：禁用水印测试', 0)
    doc4.add_paragraph("这个文档不应该有任何水印。")
    apply_word_watermark(doc4, config4)
    doc4.add_paragraph("确认没有水印显示。")
    
    output4 = "test_scenario4_disabled.docx"
    doc4.save(output4)
    print(f"场景4完成: {output4}")
    
    # 打印配置信息
    print("\n配置信息:")
    scenarios = [
        ("场景1 - 默认配置", config1),
        ("场景2 - 平铺模式", config2),
        ("场景3 - 背景模式", config3),
        ("场景4 - 禁用水印", config4)
    ]
    
    for name, config in scenarios:
        print(f"\n{name}:")
        config_dict = config.to_dict()
        for key, value in config_dict.items():
            print(f"  {key}: {value}")
    
    return [output1, output2, output3, output4]

def main():
    """主函数"""
    print("开始高级水印功能测试...")
    print("=" * 60)
    
    try:
        output_files = test_watermark_scenarios()
        
        print("\n" + "=" * 60)
        print("所有场景测试完成！")
        print("\n生成的文件:")
        
        for i, filename in enumerate(output_files, 1):
            if os.path.exists(filename):
                size = os.path.getsize(filename)
                print(f"{i}. {filename} (大小: {size} 字节)")
            else:
                print(f"{i}. {filename} ❌ 创建失败")
        
        print(f"\n请打开这些Word文档验证水印效果是否符合预期！")
        print("特别注意：")
        print("- 场景1应该显示居中的红色'投标苦'水印")
        print("- 场景2应该显示多处重复的蓝色'CONFIDENTIAL'水印")
        print("- 场景3应该显示背景样式的灰色'内部资料'水印")
        print("- 场景4应该没有任何水印")
        
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 