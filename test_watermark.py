#!/usr/bin/env python3
"""
水印功能测试脚本
测试在Word文档中添加各种类型的水印
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import os
import sys

def test_simple_watermark():
    """测试简单文本水印"""
    print("测试1: 简单文本水印")
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('水印测试文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加水印提示
    watermark_para = doc.add_paragraph()
    watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = watermark_para.add_run("【此文档包含水印: 测试水印】")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
    run.font.bold = True
    
    # 添加普通内容
    doc.add_paragraph("这是一段普通的文档内容。")
    doc.add_paragraph("水印应该在上方清晰可见。")
    
    # 添加水印段落
    for i in range(3):
        doc.add_paragraph()  # 空行
        
    watermark_center = doc.add_paragraph()
    watermark_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_run = watermark_center.add_run("★★★ 测试水印 ★★★")
    center_run.font.size = Pt(36)
    center_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
    center_run.font.bold = True
    
    for i in range(3):
        doc.add_paragraph()  # 空行
    
    # 添加更多内容
    doc.add_paragraph("这是水印后的内容。")
    doc.add_paragraph("水印应该在中间位置显示。")
    
    # 保存文档
    output_file = "test_simple_watermark.docx"
    doc.save(output_file)
    print(f"已保存: {output_file}")
    return output_file

def test_repeat_watermark():
    """测试重复平铺水印"""
    print("测试2: 重复平铺水印")
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('平铺水印测试', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加内容和水印交替出现
    for page in range(3):  # 3页内容
        if page > 0:
            doc.add_page_break()
            
        doc.add_paragraph(f"第 {page + 1} 页内容开始")
        
        # 每页添加多个水印
        for watermark_pos in range(4):
            # 添加一些内容
            doc.add_paragraph(f"这是第 {page + 1} 页的第 {watermark_pos + 1} 段内容。")
            doc.add_paragraph("中间穿插水印，确保每页都能看到。")
            
            # 添加水印
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            watermark_run = watermark_para.add_run("◆ 平铺水印 ◆")
            watermark_run.font.size = Pt(20)
            watermark_run.font.color.rgb = RGBColor(100, 150, 200)  # 蓝色
            watermark_run.font.bold = True
            watermark_run.italic = True
            
            # 装饰线
            deco_para = doc.add_paragraph()
            deco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            deco_run = deco_para.add_run("~ ~ ~ ~ ~ ~ ~ ~ ~ ~")
            deco_run.font.size = Pt(12)
            deco_run.font.color.rgb = RGBColor(100, 150, 200)
    
    # 保存文档
    output_file = "test_repeat_watermark.docx"
    doc.save(output_file)
    print(f"已保存: {output_file}")
    return output_file

def test_background_style_watermark():
    """测试背景样式水印（在页面背景位置）"""
    print("测试3: 背景样式水印")
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('背景样式水印测试', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 在文档开始处添加大面积水印背景
    for i in range(10):
        watermark_line = doc.add_paragraph()
        watermark_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if i == 5:  # 中间行放主水印
            line_run = watermark_line.add_run("【 背景水印测试 】")
            line_run.font.size = Pt(48)
            line_run.font.bold = True
        else:
            # 其他行放装饰
            line_run = watermark_line.add_run("◇ " * 15)
            line_run.font.size = Pt(14)
        
        line_run.font.color.rgb = RGBColor(200, 200, 200)  # 浅灰色
    
    # 添加普通内容（会在水印上方显示）
    doc.add_paragraph()
    doc.add_paragraph("这些内容应该显示在水印背景上方。")
    doc.add_paragraph("水印应该像背景一样分布在页面中。")
    
    # 保存文档
    output_file = "test_background_watermark.docx"
    doc.save(output_file)
    print(f"已保存: {output_file}")
    return output_file

def test_configurable_watermark(text="自定义水印", font_size=24, color_rgb=(255, 100, 50), position="center"):
    """测试可配置水印"""
    print(f"测试4: 可配置水印 - 文字: {text}, 大小: {font_size}, 位置: {position}")
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('可配置水印测试', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 根据位置配置水印
    if position == "center":
        # 居中大水印
        for i in range(8):
            doc.add_paragraph()  # 定位到中央
            
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 上边框
        border1 = doc.add_paragraph()
        border1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        border1_run = border1.add_run("═" * 20)
        border1_run.font.size = Pt(font_size // 2)
        border1_run.font.color.rgb = RGBColor(*color_rgb)
        
        # 主水印
        main_run = watermark_para.add_run(text)
        main_run.font.size = Pt(font_size * 2)
        main_run.font.color.rgb = RGBColor(*color_rgb)
        main_run.font.bold = True
        
        # 下边框
        border2 = doc.add_paragraph()
        border2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        border2_run = border2.add_run("═" * 20)
        border2_run.font.size = Pt(font_size // 2)
        border2_run.font.color.rgb = RGBColor(*color_rgb)
        
    elif position == "repeat":
        # 重复水印
        for i in range(6):
            for j in range(3):
                line_para = doc.add_paragraph()
                line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                line_run = line_para.add_run(f"   {text}   " * 3)
                line_run.font.size = Pt(font_size // 2)
                line_run.font.color.rgb = RGBColor(*color_rgb)
            doc.add_paragraph()  # 间隔
    
    # 添加说明文字
    doc.add_paragraph()
    info_para = doc.add_paragraph(f"水印配置: 文字={text}, 大小={font_size}, 颜色=RGB{color_rgb}, 位置={position}")
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    output_file = f"test_configurable_watermark_{position}.docx"
    doc.save(output_file)
    print(f"已保存: {output_file}")
    return output_file

def main():
    """运行所有测试"""
    print("开始水印功能测试...")
    print("=" * 50)
    
    created_files = []
    
    try:
        # 测试1: 简单水印
        file1 = test_simple_watermark()
        created_files.append(file1)
        print()
        
        # 测试2: 重复水印
        file2 = test_repeat_watermark()
        created_files.append(file2)
        print()
        
        # 测试3: 背景样式水印
        file3 = test_background_style_watermark()
        created_files.append(file3)
        print()
        
        # 测试4: 可配置水印 - 居中
        file4 = test_configurable_watermark("投标苦水印", 32, (255, 0, 0), "center")
        created_files.append(file4)
        print()
        
        # 测试5: 可配置水印 - 平铺
        file5 = test_configurable_watermark("CONFIDENTIAL", 18, (128, 128, 128), "repeat")
        created_files.append(file5)
        print()
        
        print("=" * 50)
        print("所有测试完成！")
        print("生成的测试文件:")
        for i, file in enumerate(created_files, 1):
            print(f"{i}. {file}")
            if os.path.exists(file):
                size = os.path.getsize(file)
                print(f"   文件大小: {size} 字节")
            else:
                print("   ❌ 文件创建失败")
        
        print("\n请打开这些Word文档查看水印效果！")
        
    except Exception as e:
        print(f"测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 