#!/usr/bin/env python3
"""
水印引擎集成示例
展示如何将水印功能集成到现有的文档处理API中
"""

from watermark_engine import WatermarkEngine, WatermarkConfig, apply_watermark_to_document
from docx import Document
import json

def simulate_api_request_data():
    """模拟前端发送的API请求数据"""
    return {
        # 文档转换基本参数
        "document_title": "投标文档示例",
        "include_images": True,
        
        # 水印配置参数（模拟前端发送的数据）
        "enable_watermark": True,
        "watermark_text": "投标苦测试",
        "watermark_font_size": 28,
        "watermark_rotation": -45,
        "watermark_opacity": 0.4,
        "watermark_color": "#FF5722",
        "watermark_position": "center"
    }

def extract_watermark_config(request_data: dict) -> dict:
    """从请求数据中提取水印配置"""
    watermark_keys = [
        'enable_watermark',
        'watermark_text', 
        'watermark_font_size',
        'watermark_rotation',
        'watermark_opacity',
        'watermark_color',
        'watermark_position'
    ]
    
    watermark_config = {}
    for key in watermark_keys:
        if key in request_data:
            watermark_config[key] = request_data[key]
    
    return watermark_config

def create_sample_document(title: str) -> Document:
    """创建示例文档内容"""
    doc = Document()
    
    # 添加标题
    main_title = doc.add_heading(title, 0)
    main_title.alignment = 1  # 居中对齐
    
    # 添加章节内容
    doc.add_heading('第一章 项目概述', level=1)
    doc.add_paragraph("这是项目概述的内容。我们的团队具有丰富的经验...")
    doc.add_paragraph("项目将按照既定的时间表和质量标准执行...")
    
    doc.add_heading('第二章 技术方案', level=1)  
    doc.add_paragraph("我们采用最新的技术框架和方法论...")
    doc.add_paragraph("技术栈包括但不限于以下几个方面：")
    
    # 添加列表
    doc.add_paragraph("• 前端技术：Vue.js, React", style='List Bullet')
    doc.add_paragraph("• 后端技术：FastAPI, Django", style='List Bullet')
    doc.add_paragraph("• 数据库：PostgreSQL, Redis", style='List Bullet')
    
    doc.add_heading('第三章 项目时间表', level=1)
    doc.add_paragraph("项目分为以下几个阶段：")
    doc.add_paragraph("第一阶段：需求分析和设计（2-3周）")
    doc.add_paragraph("第二阶段：核心功能开发（4-6周）")
    doc.add_paragraph("第三阶段：测试和部署（1-2周）")
    
    return doc

def simulate_document_conversion_api(request_data: dict) -> dict:
    """
    模拟文档转换API的处理流程
    这里展示如何在现有流程中集成水印功能
    """
    try:
        print("开始文档转换处理...")
        print(f"请求数据: {json.dumps(request_data, indent=2, ensure_ascii=False)}")
        
        # 1. 创建基础文档（模拟原有的文档生成逻辑）
        doc_title = request_data.get('document_title', '未命名文档')
        doc = create_sample_document(doc_title)
        print("✓ 基础文档创建完成")
        
        # 2. 处理其他文档内容（模拟原有逻辑）
        if request_data.get('include_images', False):
            doc.add_paragraph("注：图片内容将在此处插入")
            print("✓ 图片处理标记已添加")
        
        # 3. 【新增】水印处理逻辑
        watermark_config = extract_watermark_config(request_data)
        print(f"提取的水印配置: {watermark_config}")
        
        if watermark_config.get('enable_watermark', False):
            print("开始应用水印...")
            
            # 使用便捷函数应用水印
            watermark_success = apply_watermark_to_document(doc, watermark_config)
            
            if watermark_success:
                print("✓ 水印应用成功")
            else:
                print("⚠ 水印应用失败，继续处理文档")
        else:
            print("ℹ 水印功能未启用")
        
        # 4. 保存文档
        output_filename = f"integrated_example_{doc_title.replace(' ', '_')}.docx"
        doc.save(output_filename)
        print(f"✓ 文档保存完成: {output_filename}")
        
        # 5. 返回处理结果
        return {
            "success": True,
            "message": "文档转换成功",
            "output_file": output_filename,
            "watermark_applied": watermark_config.get('enable_watermark', False),
            "watermark_config": watermark_config
        }
        
    except Exception as e:
        print(f"✗ 文档转换失败: {e}")
        return {
            "success": False,
            "message": f"转换失败: {str(e)}",
            "output_file": None,
            "watermark_applied": False
        }

def test_different_watermark_modes():
    """测试不同的水印模式"""
    
    test_cases = [
        {
            "name": "居中红色水印",
            "config": {
                "document_title": "居中红色水印测试",
                "enable_watermark": True,
                "watermark_text": "投标苦 • 机密文档",
                "watermark_font_size": 32,
                "watermark_color": "#D32F2F",
                "watermark_position": "center",
                "watermark_opacity": 0.6
            }
        },
        {
            "name": "平铺蓝色水印", 
            "config": {
                "document_title": "平铺蓝色水印测试",
                "enable_watermark": True,
                "watermark_text": "CONFIDENTIAL",
                "watermark_font_size": 20,
                "watermark_color": "#1976D2",
                "watermark_position": "repeat",
                "watermark_opacity": 0.4
            }
        },
        {
            "name": "背景灰色水印",
            "config": {
                "document_title": "背景灰色水印测试", 
                "enable_watermark": True,
                "watermark_text": "内部使用",
                "watermark_font_size": 24,
                "watermark_color": "#616161",
                "watermark_position": "background",
                "watermark_opacity": 0.3
            }
        },
        {
            "name": "无水印文档",
            "config": {
                "document_title": "无水印测试文档",
                "enable_watermark": False,
                "watermark_text": "不应显示的水印"
            }
        }
    ]
    
    results = []
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\n{'='*60}")
        print(f"测试用例 {i}: {test_case['name']}")
        print('='*60)
        
        result = simulate_document_conversion_api(test_case['config'])
        results.append({
            "test_name": test_case['name'],
            "result": result
        })
        
        print(f"结果: {'✓ 成功' if result['success'] else '✗ 失败'}")
    
    return results

def main():
    """主函数"""
    print("水印引擎集成示例")
    print("="*80)
    
    # 1. 基本集成测试
    print("\n1. 基本集成测试")
    print("-"*40)
    sample_request = simulate_api_request_data()
    basic_result = simulate_document_conversion_api(sample_request)
    
    # 2. 多种水印模式测试
    print("\n\n2. 多种水印模式测试")
    print("-"*40)
    test_results = test_different_watermark_modes()
    
    # 3. 结果总结
    print("\n\n3. 测试结果总结")
    print("-"*40)
    
    total_tests = len(test_results) + 1  # 包括基本测试
    successful_tests = sum(1 for r in test_results if r['result']['success']) + (1 if basic_result['success'] else 0)
    
    print(f"总测试数量: {total_tests}")
    print(f"成功测试: {successful_tests}")
    print(f"失败测试: {total_tests - successful_tests}")
    print(f"成功率: {successful_tests/total_tests*100:.1f}%")
    
    print("\n生成的文件:")
    if basic_result['success']:
        print(f"• {basic_result['output_file']}")
    
    for test_result in test_results:
        if test_result['result']['success']:
            print(f"• {test_result['result']['output_file']}")
    
    print("\n集成建议:")
    print("1. 在现有的文档转换API中添加水印配置参数")
    print("2. 在文档生成完成后、保存前应用水印")
    print("3. 添加水印应用的错误处理，确保水印失败不影响主流程")
    print("4. 考虑在转换历史中记录水印配置信息")
    print("5. 为不同的文档类型设置默认水印模板")

if __name__ == "__main__":
    main() 