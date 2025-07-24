#!/usr/bin/env python3
"""
生成测试Word文件用于验证文件拼接功能
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

def create_test_word_file(filename, title, content):
    """创建测试Word文件"""
    doc = Document()
    
    # 添加标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加内容
    for i, paragraph_text in enumerate(content):
        p = doc.add_paragraph(paragraph_text)
        if i == 0:  # 第一段加粗
            p.runs[0].bold = True
    
    # 添加表格
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格标题
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "描述"
    table.cell(1, 0).text = "测试项目1"
    table.cell(1, 1).text = "这是第一个测试项目"
    table.cell(2, 0).text = "测试项目2"
    table.cell(2, 1).text = "这是第二个测试项目"
    
    # 保存文件
    doc.save(filename)
    print(f"已创建测试文件: {filename}")

if __name__ == "__main__":
    # 创建第一个测试文件
    create_test_word_file(
        "test_doc1.docx",
        "测试文档一",
        [
            "这是第一个测试Word文档的内容。",
            "这个文档包含了一些基本的文本内容。",
            "我们将在文件拼接功能中测试这个文档。"
        ]
    )
    
    # 创建第二个测试文件
    create_test_word_file(
        "test_doc2.docx",
        "测试文档二",
        [
            "这是第二个测试Word文档的内容。",
            "这个文档也包含了一些基本的文本内容。",
            "我们将测试多个Word文档的拼接功能。"
        ]
    )
    
    print("测试文件创建完成！") 